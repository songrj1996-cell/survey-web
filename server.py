"""调研分析平台 Web 后端 v2

新增：
- 本地题型推断 + 用户确认（Step 2）
- Prompt 管理（可编辑 + 版本历史）
- 历史记录（最近 20 条）
- Word 下载修复（RFC 5987）
"""

import asyncio
import base64
import csv
import hashlib
import html
import io
import json
import os
import re
import shutil
import socket
import subprocess
import tempfile
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import AsyncGenerator
from urllib.parse import quote
from urllib.request import urlopen

import openpyxl
import markdown as markdown_lib
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    FileResponse,
    JSONResponse,
    RedirectResponse,
    StreamingResponse,
)
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

import secrets

import annotate
import comment_analysis
import crosstab_parser
import survey_plan
import survey_stats
from app.integrations import feishu_client as feishu_export
from app.integrations.dify_client import chat as dify_chat  # noqa: F401 (kept for compatibility)
from app.schemas.requests import (
    AdminUserPatch,
    AdminUserRequest,
    AnnotateConfirmAIRequest,
    AnnotateConfirmRequest,
    AppSettingsPatch,
    ColumnConfirmRequest,
    HistoryQARequest,
    HistoryTitleUpdateByIdRequest,
    HistoryTitleUpdateRequest,
    PlanConfirmRequest,
    PromptUpdateRequest,
    QARequest,
    UiTextUpdateRequest,
)
from app.storage.prompts import (
    DEFAULT_PROMPTS,
    _get_planner_extra,
    _get_writer_requirements,
    _load_prompts,
    _save_prompts,
)
from app.storage.settings import (
    DEFAULT_APP_SETTINGS,
    _load_app_settings,
    _save_app_settings,
)
from app.storage.audit_store import (
    _append_audit_log,
    _load_audit_logs,
    _save_audit_logs,
)
from app.storage.history import (
    _ensure_history_report_numbers,
    _history_no_value,
    _load_history,
    _next_history_report_no,
    _save_history,
)
from app.storage.whitelist import (
    _PERMS_SCHEMA_VERSION,
    _load_whitelist,
    _migrate_whitelist_perms,
    _save_whitelist,
)
from app.storage.ui_texts import (
    DEFAULT_UI_TEXTS,
    _load_ui_texts,
    _save_ui_texts,
)
from app.core.parsing import _parse_csv, _parse_excel, _parse_file
from app.core.responses import sse_event
from app.core.text import _short_text
from app.storage.sessions import (
    SESSION_TTL,
    _COMMENT_UPLOAD_DIR,
    _SESSION_DIR,
    _session_path,
    _sweep_old_sessions,
    _write_session,
    get_session,
    new_session,
    save_session,
)
from app.storage.logins import (
    _load_web_logins,
    _save_web_logins,
    _sync_web_logins_from_disk,
    web_logins,
)
from app.core.security import (
    ALL_PERMS,
    _admin_user_rows,
    _assign_session_owner,
    _current_login,
    _email,
    _find_history_for_login,
    _forbidden_response,
    _get_user_perms,
    _history_owner_key,
    _is_admin,
    _is_public_path,
    _login_allowed,
    _login_denied_reason,
    _login_url,
    _open_id,
    _owner_from_login,
    _require_admin,
    _safe_next_path,
    _trim_history_for_owner,
    _unauthorized_response,
    _visible_to_owner,
    _wants_api_response,
    _whitelist_match,
)
from app.core.audit import (
    AUDIT_FEATURES,
    AUDIT_FEATURE_LABELS,
    _audit_log_from_login,
    _audit_user,
    _client_ip,
    audit_log,
)

# 配置常量、Dify/飞书参数、数据路径、阈值、默认文案统一来自 app/core/config。
# 过渡期 server.py 仍以 import * 取回这些名字;step3 server.py 瘦身后此行移除。
from app.core.config import *  # noqa: F401,F403,E402


def _inject_disclaimer(md: str, mode: str = "") -> str:
    """在第一行 `# 标题` 之后插入免责声明引用行；幂等；无 H1 则插到最前。

    通用声明 REPORT_DISCLAIMER 三种模式都插。第二条声明按 mode 选择：
      - mode == "crosstab"（倍市得）：不插第二条，并清除历史报告里残留的定性/评论声明
      - mode == "comment"（评论舆情）：插 COMMENT_DISCLAIMER，清除残留的定性声明
      - 其它（定性，默认）：插 QUALITATIVE_DISCLAIMER，清除残留的评论声明
    """
    if not md:
        return md

    # 选定本模式应有的“第二条声明”
    if mode == "crosstab":
        second = None
    elif mode == "comment":
        second = COMMENT_DISCLAIMER
    else:
        second = QUALITATIVE_DISCLAIMER

    # 清除不属于当前模式的历史第二条声明（兼容历史报告 / 模式切换）
    for stale in (QUALITATIVE_DISCLAIMER, COMMENT_DISCLAIMER):
        if stale is not second and stale in md:
            md = md.replace("\n" + stale, "").replace(stale + "\n", "")

    has_report = REPORT_DISCLAIMER in md
    has_second = bool(second) and second in md

    # 已经齐备就不再插
    if has_report and (second is None or has_second):
        return md

    lines = md.split("\n")

    # 已有 REPORT_DISCLAIMER 但缺第二条声明
    if has_report and second and not has_second:
        for i, ln in enumerate(lines):
            if ln.strip() == REPORT_DISCLAIMER:
                lines.insert(i + 1, second)
                return "\n".join(lines)

    for i, ln in enumerate(lines):
        if ln.startswith("# ") and not ln.startswith("## "):
            lines.insert(i + 1, "")
            insert_at = i + 2
            if not has_report:
                lines.insert(insert_at, REPORT_DISCLAIMER)
                insert_at += 1
            if second and not has_second:
                lines.insert(insert_at, second)
            return "\n".join(lines)

    # 没有 H1：插到最前
    prefix = []
    if not has_report:
        prefix.append(REPORT_DISCLAIMER)
    if second and not has_second:
        prefix.append(second)
    return "\n".join(prefix) + "\n\n" + md


def _strip_core_markers(md: str) -> str:
    """移除核心结论包裹标记行（仅飞书导出用于定位高亮块，其它导出不应出现）。"""
    if not md:
        return md
    return "\n".join(ln for ln in md.split("\n") if ln.strip() not in (CORE_START, CORE_END))


def _prep_export_md(md: str, mode: str = "") -> str:
    """通用导出前处理：补免责声明（幂等）+ 去掉核心结论标记。"""
    return _strip_core_markers(_inject_disclaimer(md, mode=mode))




# ============================================================
# Session 管理（文件持久化，支持多 worker）
#
# 每个 session 存为 data/sessions/<uuid>.json。
# 写入走 tmp + os.replace 保证原子性，多进程安全。
# annotate_sessions 仍用内存（标注会话生命周期短，不跨请求保活）。
# ============================================================

# ============================================================
# 历史记录
# ============================================================

def _qa_user_count(entry: dict) -> int:
    return sum(1 for m in entry.get("qa_messages", []) if m.get("role") == "user")


def _sanitize_report_title(title: str) -> str:
    cleaned = re.sub(r"[\r\n\t]+", " ", str(title or "")).strip()
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        raise HTTPException(status_code=400, detail="报告名称不能为空")
    return cleaned[:120]


def _replace_report_h1(report_md: str, title: str) -> str:
    report_md = report_md or ""
    if re.search(r"^#\s+.+?$", report_md, re.MULTILINE):
        return re.sub(r"^#\s+.+?$", f"# {title}", report_md, count=1, flags=re.MULTILINE)
    return f"# {title}\n\n{report_md.lstrip()}"


def _comment_report_title(sess: dict) -> str:
    post_title = re.sub(r"\s+", " ", str(sess.get("comment_post_title") or "").strip())
    if post_title:
        suffix = "·舆情简报"
        max_post_len = max(1, 120 - len(suffix))
        if len(post_title) > max_post_len:
            post_title = post_title[: max_post_len - 1] + "…"
        return f"{post_title}{suffix}"
    return "评论分析·舆情简报"


def save_to_history(session_id: str, sess: dict) -> None:
    report_md = sess.get("report_md", "")
    if not report_md:
        return
    history = _load_history()
    _ensure_history_report_numbers(history, save=False)
    old_entry = next((h for h in history if h.get("id") == session_id), None)
    qa_messages = sess.get("qa_messages")
    if qa_messages is None and old_entry:
        qa_messages = old_entry.get("qa_messages", [])
    if sess.get("mode") == "comment":
        title = sess.get("comment_report_title") or _comment_report_title(sess)
    else:
        title_m = re.search(r"^#\s+(.+?)$", report_md, re.MULTILINE)
        title = title_m.group(1).strip() if title_m else "未命名报告"
    owner = {
        "owner_key": sess.get("owner_key") or (old_entry or {}).get("owner_key", ""),
        "owner_email": sess.get("owner_email") or (old_entry or {}).get("owner_email", ""),
        "owner_open_id": sess.get("owner_open_id") or (old_entry or {}).get("owner_open_id", ""),
        "owner_name": sess.get("owner_name") or (old_entry or {}).get("owner_name", ""),
    }
    entry = {
        "id": session_id,
        "report_no": old_entry.get("report_no") if old_entry else _next_history_report_no(history),
        "filename": sess.get("filename", "unknown"),
        "title": title,
        "created_at": old_entry.get("created_at") if old_entry else datetime.now().isoformat(),
        "report_md": report_md,
        "plan": sess.get("plan"),
        "stats_md": sess.get("stats_md"),
        "analyst_conv_id": sess.get("analyst_conv_id", ""),
        "analyst_app": sess.get("analyst_app", ""),
        "qa_messages": qa_messages or [],
        "rows_fed": True,  # 历史 QA 跳过投喂 rows（对话已包含上下文）
        "mode": sess.get("mode", ""),  # 保存模式，导出时据此选择免责声明（定性/crosstab/comment）
        "row_count": max(0, len(sess.get("rows") or []) - 1) or (old_entry or {}).get("row_count", 0),
        **owner,
    }
    if sess.get("mode") == "comment":
        entry.update({
            "comment_file_hash": sess.get("comment_file_hash", ""),
            "comment_source_filename": sess.get("filename", "unknown"),
            "comment_post_title": sess.get("comment_post_title", ""),
            "comment_report_title": title,
            "comment_result": sess.get("comment_result"),
            "comment_sample_meta": sess.get("comment_sample_meta", {}),
            "comment_relevance_stats": sess.get("comment_relevance_stats", {}),
            "comment_selected_raw_comments": sess.get("comment_selected_raw_comments", []),
            "comment_valid_count": sess.get("comment_valid_count", 0),
            "comment_sample_count": sess.get("comment_sample_count", 0),
            "comment_scan_rows": sess.get("comment_scan_rows", 0),
            "comment_nonempty_count": sess.get("comment_nonempty_count", 0),
        })
    history = [h for h in history if h["id"] != session_id]
    history.insert(0, entry)
    history = _trim_history_for_owner(history, owner.get("owner_key", ""))
    _save_history(history)


def save_annotate_to_history(sid: str, sess: dict, result_path: str, download_name: str) -> None:
    history = _load_history()
    _ensure_history_report_numbers(history, save=False)
    old_entry = next((h for h in history if h.get("id") == sid), None)
    filename = sess.get("filename", "annotated")
    stem = re.sub(r"\.(csv|xlsx|xls)$", "", filename, flags=re.IGNORECASE).strip() or "标注结果"
    owner = {
        "owner_key": sess.get("owner_key") or (old_entry or {}).get("owner_key", ""),
        "owner_email": sess.get("owner_email") or (old_entry or {}).get("owner_email", ""),
        "owner_open_id": sess.get("owner_open_id") or (old_entry or {}).get("owner_open_id", ""),
        "owner_name": sess.get("owner_name") or (old_entry or {}).get("owner_name", ""),
    }
    tasks = sess.get("tasks", {}) or {}
    entry = {
        "id": sid,
        "report_no": old_entry.get("report_no") if old_entry else _next_history_report_no(history),
        "filename": filename,
        "title": (old_entry or {}).get("title") or f"数据标注结果 - {stem}",
        "created_at": old_entry.get("created_at") if old_entry else datetime.now().isoformat(),
        "report_md": "",
        "plan": None,
        "stats_md": "",
        "analyst_conv_id": "",
        "analyst_app": "",
        "qa_messages": [],
        "rows_fed": False,
        "mode": "annotate",
        "row_count": max(0, len(sess.get("rows") or []) - 1),
        "annotate_result_path": result_path,
        "annotate_download_name": download_name,
        "annotate_ai_count": len(sess.get("ai_results") or []),
        "annotate_confirmed_ai_count": len(sess.get("confirmed_ai_ids") or []),
        "annotate_quality_count": len(sess.get("quality_results") or []),
        "annotate_tasks": {
            "ai_detect": bool(tasks.get("ai_detect")),
            "quality": bool(tasks.get("quality")),
        },
        **owner,
    }
    history = [h for h in history if h.get("id") != sid]
    history.insert(0, entry)
    history = _trim_history_for_owner(history, owner.get("owner_key", ""))
    _save_history(history)


def _annotate_download_filename(filename: str) -> str:
    stem = re.sub(r"\.(csv|xlsx|xls)$", "", filename or "annotated", flags=re.IGNORECASE)
    safe = re.sub(r'[\\/:*?"<>|]', "_", stem).strip() or "annotated"
    return f"{safe}_标注结果.xlsx"


def _annotate_result_path(sid: str) -> Path:
    safe_sid = re.sub(r"[^A-Za-z0-9_-]", "_", sid)
    return ANNOTATE_RESULT_DIR / f"{safe_sid}.xlsx"


def _annotate_incomplete_detail(sess: dict) -> str:
    missing_ai = sess.get("missing_ai_ids", []) or []
    missing_q = sess.get("missing_quality_ids", []) or []
    parts = []
    if missing_ai:
        ids_preview = ", ".join(missing_ai[:5]) + ("…" if len(missing_ai) > 5 else "")
        parts.append(f"AI 检测漏返 {len(missing_ai)} 行（ID：{ids_preview}）")
    if missing_q:
        ids_preview = ", ".join(missing_q[:5]) + ("…" if len(missing_q) > 5 else "")
        parts.append(f"质量打标漏返 {len(missing_q)} 行（ID：{ids_preview}）")
    return "；".join(parts)


def _build_annotate_excel_from_session(sess: dict) -> tuple[bytes, str]:
    rows = sess.get("rows")
    headers = sess.get("headers")
    if not rows:
        raise HTTPException(status_code=400, detail="会话中没有数据")
    incomplete = _annotate_incomplete_detail(sess)
    if incomplete:
        raise HTTPException(status_code=400, detail=f"结果不完整，无法下载：{incomplete}。请返回重试对应任务。")
    filename = sess.get("filename", "annotated")
    excel_bytes = annotate.generate_annotated_excel(
        rows,
        headers,
        sess.get("ai_results", []),
        set(sess.get("confirmed_ai_ids", [])),
        sess.get("quality_results", []),
        sess.get("open_text_cols", []),
        sess.get("id_col", 1),
        sess.get("tasks", {}),
    )
    return excel_bytes, _annotate_download_filename(filename)


async def _save_annotate_result_history(sid: str, sess: dict, request: Request) -> None:
    if _annotate_incomplete_detail(sess):
        return
    login = await _current_login(request)
    _assign_session_owner(sess, login)
    loop = asyncio.get_event_loop()
    excel_bytes, download_name = await loop.run_in_executor(
        None,
        _build_annotate_excel_from_session,
        sess,
    )
    ANNOTATE_RESULT_DIR.mkdir(parents=True, exist_ok=True)
    result_path = _annotate_result_path(sid)
    result_path.write_bytes(excel_bytes)
    save_annotate_to_history(sid, sess, str(result_path), download_name)


def _find_comment_duplicate_report(file_hash: str, login: dict | None) -> dict | None:
    if not file_hash:
        return None
    history = _ensure_history_report_numbers(_load_history())
    for entry in history:
        if (
            entry.get("mode") == "comment"
            and entry.get("comment_file_hash") == file_hash
            and _visible_to_owner(entry, login)
        ):
            return {
                "id": entry.get("id", ""),
                "report_no": entry.get("report_no", ""),
                "title": entry.get("title", "评论分析报告"),
                "filename": entry.get("filename", ""),
                "created_at": entry.get("created_at", ""),
                "valid_count": entry.get("comment_valid_count", 0),
                "sample_count": entry.get("comment_sample_count", 0),
            }
    return None


def _history_effective_row_count(entry: dict) -> int:
    try:
        row_count = int(entry.get("row_count") or 0)
    except (TypeError, ValueError):
        row_count = 0
    if row_count > 0:
        return row_count
    text = f"{entry.get('stats_md') or ''}\n{entry.get('report_md') or ''}"
    for pattern in (
        r"有效样本\(总计\):总体=(\d+)",
        r"有效样本（总计）:总体=(\d+)",
        r"有效样本[^\n]*总体=(\d+)",
        r"(\d+)\s*名受访者",
    ):
        m = re.search(pattern, text)
        if m:
            return int(m.group(1))
    return 0


# ============================================================
# 文件解析
# ============================================================

# ============================================================
# 本地题型推断
# ============================================================

ROLE_LABEL_MAP = {
    "id":            "用户ID",
    "mlbbid":        "MLBB ID",
    "profile_dim":   "画像维度",
    "single_choice": "单选题",
    "multi_choice":  "多选题",
    "scale":         "量表题",
    "matrix_scale":  "矩阵打分",
    "matrix_multi":  "矩阵多选",
    "open_text":     "开放题",
    "ignore":        "忽略此列",
}

LABEL_ROLE_MAP = {v: k for k, v in ROLE_LABEL_MAP.items()}


def _heuristic_type(header: str, values: list[str]) -> str:
    h = header.lower().strip()

    # ── ID / 忽略 ──
    if any(kw in h for kw in ["_id", "userid", "player", "玩家id", "用户id", "编号"]):
        return "id"
    if re.match(r"^(id|uid|uuid)$", h):
        return "id"
    if any(kw in h for kw in ["时间", "timestamp", "submit", "提交", "date", "日期"]):
        return "ignore"

    non_empty = [v.strip() for v in values if v.strip()]
    if not non_empty:
        return "open_text"
    total = len(non_empty)

    # ── 数值型 → 量表 ──
    nums = []
    for v in non_empty:
        try:
            nums.append(float(v))
        except ValueError:
            pass
    if len(nums) / total > 0.85:
        mn, mx = min(nums), max(nums)
        if mx - mn <= 15 and mn >= 0:
            return "scale"

    # ── 长文本优先判断（避免后续规则误伤）──
    avg_len = sum(len(v) for v in non_empty) / total
    if avg_len > 25:
        return "open_text"

    # ── 多选（分隔符检测：只认短片段列表，排除句子中的标点）──
    delimiters = [",", "，", ";", "；", "、", "|"]
    def _is_list(v: str) -> bool:
        for d in delimiters:
            if d in v:
                parts = [p.strip() for p in v.split(d) if p.strip()]
                if len(parts) >= 2 and all(len(p) < 30 for p in parts):
                    return True
        return False
    delim_count = sum(1 for v in non_empty if _is_list(v))
    if delim_count / total > 0.25:
        return "multi_choice"

    # ── 唯一值数量 ──
    unique_vals = set(non_empty)
    n_unique = len(unique_vals)
    ratio = n_unique / total

    if n_unique <= 2:
        return "single_choice"
    if n_unique <= 8:
        return "single_choice"
    if ratio < 0.25:
        return "single_choice"

    return "single_choice"


# ============================================================
# Google Form 矩阵题分组 + LLM 题型识别
# ============================================================

# `主问题 [子项]` 或 `主问题（子项）`：Google Form 矩阵题导出格式
_MATRIX_HEADER_RE = re.compile(r"^(.*?)\s*[\[\(（](.+?)[\]\)）]\s*$")


def _group_googleform_matrix(headers: list[str]) -> list[dict]:
    """识别 `主问题 [子项]` 多列 → 合并为逻辑题。

    返回逻辑题分组（保持原列顺序，矩阵组落在其首列位置）：
      {"type": "matrix"|"single", "title", "member_indexes": [...], "row_labels": [...]}
    """
    parsed = []  # (idx, prefix|None, sub|None)
    for i, h in enumerate(headers):
        hs = (h or "").strip()
        m = _MATRIX_HEADER_RE.match(hs)
        if m and m.group(1).strip():
            parsed.append((i, m.group(1).strip(), m.group(2).strip()))
        else:
            parsed.append((i, None, None))

    prefix_cols: dict[str, list[tuple[int, str]]] = {}
    for idx, pref, sub in parsed:
        if pref is not None:
            prefix_cols.setdefault(pref, []).append((idx, sub))
    matrix_prefixes = {p for p, cols in prefix_cols.items() if len(cols) >= 2}

    groups: list[dict] = []
    emitted: set[int] = set()
    for idx, pref, _sub in parsed:
        if idx in emitted:
            continue
        if pref in matrix_prefixes:
            cols = prefix_cols[pref]
            groups.append({
                "type": "matrix",
                "title": pref,
                "member_indexes": [c[0] for c in cols],
                "row_labels": [c[1] for c in cols],
            })
            emitted.update(c[0] for c in cols)
        else:
            groups.append({
                "type": "single",
                "title": (headers[idx] or "").strip() or f"列{idx}",
                "member_indexes": [idx],
                "row_labels": [],
            })
            emitted.add(idx)
    return groups


def _detect_open_text_cols(rows: list, headers: list) -> list[int]:
    """检测主观题列，复用 _heuristic_type + 矩阵题过滤，与分析流程保持一致。"""
    if len(rows) <= 1:
        return []
    body = rows[1:]

    # 矩阵题的所有子列均为单选，先排除
    matrix_idxs: set[int] = set()
    for g in _group_googleform_matrix(headers):
        if g["type"] == "matrix":
            matrix_idxs.update(g["member_indexes"])

    result = []
    for i, header in enumerate(headers):
        if i in matrix_idxs:
            continue
        vals = [str(r[i]) if i < len(r) else "" for r in body]
        if _heuristic_type(header, vals) == "open_text":
            result.append(i)
    return result


def _col_samples(body: list[list], idx: int, n: int = 20) -> list[str]:
    out: list[str] = []
    for r in body:
        v = str(r[idx]) if idx < len(r) else ""
        if v.strip():
            out.append(v.strip())
        if len(out) >= n:
            break
    return out


_COLUMN_DETECT_SCHEMA_HINT = """\
请只输出一段 ```json``` 围栏，schema 如下（不要附加任何解释文字）：
{
  "questions": [
    {
      "name_zh": "中文题名（把英文/原文题目翻译成简洁中文）",
      "role": "single_choice|multi_choice|scale|profile_dim|open_text|id|mlbbid|matrix_scale|matrix_multi|ignore",
      "column_indexes": [列号...],            // 普通题1个；矩阵题为该题所有子项列号
      "delimiter": "，",                       // 仅 multi_choice：选项分隔符（兜底）
      "options": ["选项A","选项B"],            // 选项题清单：优先用合并后的中文标准值
      "scale_min": 1, "scale_max": 5,          // scale / matrix_scale：量程
      "rows": ["子项1","子项2"],               // matrix_*：与 column_indexes 顺序一一对应的行标签
      "value_aliases": {"中文标准值": ["原始变体1","Mythic","Mítica"]},  // 见下「同义归并」
      "low_confidence": false  // 若对该题型判断不确定（样本稀少/题名模糊），设为 true
    }
  ]
}
角色判断要点：
- 玩家ID/编号 → id；明确是 MLBB 游戏ID → mlbbid；提交时间等 → ignore
- 年龄段/段位/地区等用于分群的 → profile_dim
- 数值评分（如 1–5、1–10）→ scale
- 含多个选项、可多选 → multi_choice，并尽量给出 options 清单；若多数回答为逗号分隔的短词（如 "Classic, Ranked"），即使少数回答带括号补充说明（如 "Arcade (Such as Tide Siege)"），括号内容视为该选项的描述、不影响题型判断，仍应判断为 multi_choice
- 长文本主观回答 → open_text
- 标注【疑似矩阵题】的，按子项判断 matrix_scale（每子项打分）或 matrix_multi（每子项可多选），rows 用给出的子项标签

同义归并（value_aliases，重要）：
- 仅对 single_choice / profile_dim / multi_choice / matrix_multi 这类「选项题」给出。
- 我已在每列附上「去重取值」。请把**语义相同但写法/语种不同**的取值（如 神话/Mythic/Mítica，中国/China/CN）归并到**同一个中文标准值**：key=中文标准值，value=该标准值对应的所有原始变体（含中文标准值本身可不必重复列出）。
- options 使用这些合并后的中文标准值；中文标准值可以不直接出现在原始数据里，但必须能由 value_aliases 中的真实取值支撑。
- 只有确属同义才合并；拿不准就不要合并。没有任何同义可并的列，可以省略 value_aliases 或给 {}。
- **程度不同的选项禁止合并**：如「有点长」与「太长了」、「还好」与「很好」，虽然方向相同但程度档位不同，不是同义，必须分开。同义归并仅限写法/语种不同、语义程度完全一致的取值。
- **同一道题内必须完整归并**：若已把某语种的某个表达归入了某标准值（如把「Lebih dari 3 tahun」归入「3年以上」），则该题内同语种所有语义相同的取值也必须全部归入，不得部分遗漏（如已识别印尼语「tahun」=年，则含相同数字区间的「1~3 tahun」也须归入「1~3年」）。
- 这样统计表里这些取值会被合并计数，避免同义被拆开。\
"""


def _col_distinct(body: list[list], idx: int, n: int = 60) -> tuple[list[tuple[str, int]], int]:
    """返回某列去重取值 [(值, 频次)]（按频次降序，上限 n）+ 去重总数。"""
    from collections import Counter
    cnt: Counter = Counter()
    for r in body:
        v = (str(r[idx]) if idx < len(r) else "").strip()
        if v:
            cnt[v] += 1
    total_distinct = len(cnt)
    return cnt.most_common(n), total_distinct


def _fmt_distinct(body: list[list], idx: int, n: int = 60) -> str:
    items, total_distinct = _col_distinct(body, idx, n)
    parts = [f"{v}（{c}）" for v, c in items]
    s = " | ".join(parts)
    if total_distinct > len(items):
        s += f" …（共 {total_distinct} 种不同取值，已截断前 {len(items)}）"
    return s or "（空）"


CHOICE_ROLES = {"single_choice", "profile_dim", "multi_choice", "matrix_multi"}


def _norm_option_key(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().casefold()


def _split_option_cell(value: str, delimiter: str | None = None) -> list[str]:
    text = str(value or "").strip()
    if not text:
        return []
    delims = [delimiter] if delimiter else []
    delims += ["，", ",", ";", "；", "\n", "\r\n", "|"]
    for d in delims:
        if d and d in text:
            return [p.strip() for p in text.split(d) if p.strip()]
    return [text]


def _real_options_for_question(rows: list[list], q: dict) -> list[str]:
    body = rows[1:]
    role = q.get("role")
    cis = q.get("column_indexes") or []
    delimiter = q.get("delimiter")
    out: list[str] = []
    seen: set[str] = set()
    for idx in cis:
        for r in body:
            raw = str(r[idx]) if idx < len(r) else ""
            parts = _split_option_cell(raw, delimiter) if role in ("multi_choice", "matrix_multi") else [raw.strip()]
            for part in parts:
                key = _norm_option_key(part)
                if key and key not in seen:
                    seen.add(key)
                    out.append(part)
    return out


def _sanitize_choice_options(rows: list[list], questions: list[dict]) -> list[dict]:
    """Ground choice options in real cell values while preserving canonical aliases.

    The detector may translate and merge multilingual values into a canonical label
    that does not literally appear in the raw data. Keep those labels only when
    value_aliases proves they are backed by real cell values.
    """
    for q in questions:
        role = q.get("role")
        if role not in CHOICE_ROLES:
            continue
        real_options = _real_options_for_question(rows, q)
        real_by_key = {_norm_option_key(o): o for o in real_options}
        if not real_by_key:
            continue

        raw_aliases = q.get("value_aliases") if isinstance(q.get("value_aliases"), dict) else {}
        cleaned_options: list[str] = []
        cleaned_originals: list[str] = []
        cleaned_aliases: dict[str, list[str]] = {}
        seen_canonical: set[str] = set()
        covered_real: set[str] = set()

        def real_values_for(canonical: str, aliases: list | tuple | None = None) -> list[str]:
            values = [canonical]
            if aliases:
                values.extend(aliases)
            out: list[str] = []
            seen_real: set[str] = set()
            for value in values:
                key = _norm_option_key(str(value))
                real = real_by_key.get(key)
                if real is not None and key not in seen_real:
                    seen_real.add(key)
                    out.append(real)
            return out

        def add_canonical(canonical: str, aliases: list | tuple | None = None) -> None:
            canonical = str(canonical or "").strip()
            ckey = _norm_option_key(canonical)
            if not canonical or ckey in seen_canonical:
                return
            matched = [v for v in real_values_for(canonical, aliases) if _norm_option_key(v) not in covered_real]
            if not matched:
                return
            seen_canonical.add(ckey)
            cleaned_options.append(canonical)
            cleaned_originals.append(matched[0])
            for value in matched:
                covered_real.add(_norm_option_key(value))
            if len(matched) > 1 or _norm_option_key(canonical) != _norm_option_key(matched[0]):
                cleaned_aliases[canonical] = matched

        for opt in q.get("options") or []:
            opt_text = str(opt or "").strip()
            add_canonical(opt_text, raw_aliases.get(opt_text))

        for canonical, aliases in raw_aliases.items():
            add_canonical(str(canonical), aliases if isinstance(aliases, (list, tuple)) else [])

        for opt in real_options:
            add_canonical(opt, [])

        if cleaned_options:
            q["options"] = cleaned_options
            q["options_original"] = cleaned_originals
            if cleaned_aliases:
                q["value_aliases"] = cleaned_aliases
            else:
                q.pop("value_aliases", None)
    return questions


def _build_column_detect_query(rows: list[list], groups: list[dict]) -> str:
    body = rows[1:]
    blocks: list[str] = []
    for g in groups:
        if g["type"] == "matrix":
            lines = [
                f"【疑似矩阵题】主问题: {g['title']}"
                f"（{len(g['member_indexes'])} 个子项，column_indexes={g['member_indexes']}）"
            ]
            for k, idx in enumerate(g["member_indexes"]):
                sub = g["row_labels"][k]
                lines.append(f"  · 子项[{sub}] 列{idx} 去重取值: {_fmt_distinct(body, idx, 40)}")
            blocks.append("\n".join(lines))
        else:
            idx = g["member_indexes"][0]
            blocks.append(f"列{idx} 表头: {g['title']}\n  去重取值（值后括号是出现次数）: {_fmt_distinct(body, idx, 60)}")

    total = max(0, len(rows) - 1)
    return (
        f"<columns>\n总数据行数（不含表头）: {total}\n\n"
        + "\n\n".join(blocks)
        + "\n</columns>\n\n"
        + "选项边界（严格执行）：options 必须由该列「去重取值」里的真实单元格取值或多选拆分值支撑；不得从题干/表头中抽取选项。若 New Medal 等词只出现在题干里、没有出现在该列取值里，不得写入 options。若多语言取值语义相同，options 请写合并后的中文标准值，并在 value_aliases 中列出支撑它的真实取值。\n\n"
        + _COLUMN_DETECT_SCHEMA_HINT
    )


def _heuristic_questions(rows: list[list], groups: list[dict]) -> list[dict]:
    """LLM 解析失败时的本地兜底：用启发式 + 矩阵分组拼出 questions。"""
    headers = rows[0]
    body = rows[1:]
    out: list[dict] = []
    for g in groups:
        if g["type"] == "matrix":
            # 子项样本是否多为数值 → matrix_scale，否则 matrix_multi
            numeric = 0
            checked = 0
            for idx in g["member_indexes"]:
                for v in _col_samples(body, idx, 10):
                    checked += 1
                    try:
                        float(v)
                        numeric += 1
                    except ValueError:
                        pass
            role = "matrix_scale" if checked and numeric / checked > 0.7 else "matrix_multi"
            q = {
                "name_zh": g["title"],
                "role": role,
                "column_indexes": g["member_indexes"],
                "rows": g["row_labels"],
            }
            if role == "matrix_scale":
                q["scale_min"], q["scale_max"] = 1, 5
            else:
                q["delimiter"] = "，"
            out.append(q)
        else:
            idx = g["member_indexes"][0]
            vals = [str(r[idx]) if idx < len(r) else "" for r in body]
            role = _heuristic_type(headers[idx], vals)
            q = {"name_zh": g["title"], "role": role, "column_indexes": [idx]}
            if role == "multi_choice":
                q["delimiter"] = "，"
            if role == "scale":
                q["scale_min"], q["scale_max"] = 1, 5
            out.append(q)
    return out


def _enrich_questions(questions: list[dict], headers: list[str], groups: list[dict]) -> list[dict]:
    """补全前端展示需要的字段（name_zh 兜底、矩阵 rows 兜底）。"""
    matrix_rows_by_first: dict[int, list[str]] = {}
    for g in groups:
        if g["type"] == "matrix":
            matrix_rows_by_first[g["member_indexes"][0]] = g["row_labels"]
    for q in questions:
        cis = q.get("column_indexes") or []
        if not q.get("name_zh"):
            first = cis[0] if cis else None
            q["name_zh"] = (headers[first].strip() if first is not None and first < len(headers) else "") or "未命名题目"
        # 矩阵题缺 rows → 用本地分组补
        if q.get("role") in ("matrix_scale", "matrix_multi") and not q.get("rows"):
            if cis and cis[0] in matrix_rows_by_first:
                q["rows"] = matrix_rows_by_first[cis[0]]
    return questions


# ============================================================
# Planner / Writer 构建
# ============================================================

def _build_planner_sample(rows: list[list], sample_n: int = 5) -> str:
    if not rows:
        return ""
    headers = rows[0]
    sample = rows[1: 1 + sample_n]

    def esc(s):
        return ("" if s is None else str(s)).replace("|", "\\|").replace("\n", "<br>")

    md = "| " + " | ".join(esc(h) for h in headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for r in sample:
        cells = [r[i] if i < len(r) else "" for i in range(len(headers))]
        md += "| " + " | ".join(esc(c) for c in cells) + " |\n"

    total_data_rows = max(0, len(rows) - 1)
    return (
        f"<sample>\n"
        f"总数据行数（不含表头）: {total_data_rows}\n"
        f"以下展示表头 + 前 {len(sample)} 行样本：\n\n"
        f"{md}\n"
        f"</sample>"
    )


def _build_planner_query_with_confirmed(rows: list[list], confirmed_columns: list[dict]) -> str:
    """构建给 Planner 的完整 query，含用户确认的题型（逻辑题，矩阵题跨多列）。"""
    sample_md = _build_planner_sample(rows)

    confirmed_lines = []
    for q in confirmed_columns:
        # 兼容旧结构（confirmed_type/index）与新结构（role/name_zh/column_indexes）
        role = q.get("role") or q.get("confirmed_type") or "single_choice"
        name = q.get("name_zh") or q.get("name") or "?"
        cis = q.get("column_indexes") or ([q["index"]] if "index" in q else [])
        label = ROLE_LABEL_MAP.get(role, role)
        extra = ""
        if role in ("single_choice", "profile_dim", "multi_choice", "matrix_multi") and q.get("options"):
            opts = "、".join(str(o) for o in q["options"][:12])
            extra += f"，选项: {opts}"
        if role in ("multi_choice",) and q.get("delimiter"):
            extra += f"，分隔符: 「{q['delimiter']}」"
        if role in ("scale", "matrix_scale") and q.get("scale_min") is not None:
            extra += f"，量程: {q.get('scale_min')}–{q.get('scale_max')}"

        if role in ("matrix_scale", "matrix_multi"):
            rows_lbl = "、".join(str(r) for r in (q.get("rows") or []))
            confirmed_lines.append(
                f"- 矩阵题「{name}」({label})，子项行: {rows_lbl}；"
                f"对应列号 {cis}（这些列同属一道题，**必须整体归入同一个 part**）{extra}"
            )
        else:
            idx = cis[0] if cis else q.get("index", 0)
            confirmed_lines.append(f"- 列{idx}「{name}」: {label}{extra}")

    confirmed_block = "<confirmed_column_types>\n" + "\n".join(confirmed_lines) + "\n</confirmed_column_types>"
    extra_instructions = _get_planner_extra()

    # 检测是否存在画像维度列，生成对应的画像约束指令
    profile_dims = [q for q in confirmed_columns if (q.get("role") or q.get("confirmed_type")) == "profile_dim"]
    if not profile_dims:
        profile_constraint = (
            "\n⚠️ 画像约束（严格执行）：本问卷中用户**没有将任何题目标注为画像维度**。\n"
            "- cross_tabs 数组**必须为空** []\n"
            "- open_questions **不得**建议将任何题目用作用户画像或分组维度\n"
            "- 报告不应包含任何「用户画像」/「人群结构」分析章节\n"
        )
    else:
        dim_names = "、".join(
            f"「{q.get('name_zh') or q.get('name') or '?'}」" for q in profile_dims
        )
        profile_constraint = (
            f"\n画像维度约束：本问卷的画像维度列为 {dim_names}。"
            f"cross_tabs 的 profile_index **只能**使用上述列对应的列号，不得使用其他单选题做交叉分析。\n"
        )

    return (
        f"{sample_md}\n\n"
        f"{confirmed_block}\n\n"
        f"重要：以上题型和选项已由用户在界面中逐一确认；选择题选项必须以 <confirmed_column_types> 中的「选项」为权威，不得根据题干、表头或样本重新猜测选项，也不得围绕已确认选项再次提问。\n"
        f"注意：以上题型已由用户在界面中逐一确认，**不得**在 open_questions 中再次对题型进行发问。"
        f"选项的归并方式（哪个原始值归入哪个标准选项）同样已由用户在界面中逐一确认，**不得**在 open_questions 中就选项归并或分拆方式再次提问。"
        f"矩阵题的多个列号务必整体归入同一个 part。"
        f"{profile_constraint}\n"
        f"{extra_instructions}"
    )


def _build_plan_revision_query(plan: dict, headers: list[str], confirmed_columns: list[dict], user_text: str) -> str:
    header_lines = "\n".join(f"- 列{i}: {h}" for i, h in enumerate(headers))
    confirmed_json = json.dumps(confirmed_columns or [], ensure_ascii=False, indent=2)
    plan_json = json.dumps(plan or {}, ensure_ascii=False, indent=2)
    return (
        "你正在修订一份问卷分析方案。请根据用户的修改意见，在当前方案基础上输出一份完整的新 plan JSON。\n\n"
        "严格要求：\n"
        "1. 只能输出一个完整 JSON 对象，不要输出解释、确认语、Markdown 文本或 ```json 围栏外的内容。\n"
        "2. JSON 必须包含 columns、parts、cross_tabs、open_questions 字段，并通过既有 schema 校验。\n"
        "3. columns 必须保留用户已确认的题型、列号、选项、矩阵题分组等权威信息；不要重新猜测题型或选项。\n"
        "4. parts 必须使用实际存在的列号；矩阵题成员列必须整体归入同一个 part。\n"
        "5. 若用户意见只要求调整章节/分析重点，只改 parts、cross_tabs 或 open_questions，不要无故改 columns。\n\n"
        f"<headers>\n{header_lines}\n</headers>\n\n"
        f"<confirmed_columns_json>\n{confirmed_json}\n</confirmed_columns_json>\n\n"
        f"<current_plan_json>\n{plan_json}\n</current_plan_json>\n\n"
        f"<user_revision_request>\n{user_text.strip()}\n</user_revision_request>\n\n"
        "请现在返回修订后的完整 JSON 对象。"
    )


def _build_crosstab_planner_query(
    questionnaire_text: str,
    available_questions: list[str],
    open_question_names: list[str],
) -> str:
    """跑数表模式：给章节策划 planner 的初始 query（任务/输出格式由 Dify 应用 system prompt 定义）。"""
    q_text = (questionnaire_text or "").strip()
    if len(q_text) > 12000:
        q_text = q_text[:12000] + "\n…（问卷过长，已截断）"
    avail = "\n".join(f"- {q}" for q in available_questions) or "（无）"
    opens = "\n".join(f"- {q}" for q in open_question_names) or "（无）"
    return (
        f"<questionnaire>\n{q_text}\n</questionnaire>\n\n"
        f"<available_questions>\n{avail}\n</available_questions>\n\n"
        f"<open_questions_list>\n{opens}\n</open_questions_list>\n\n"
        "请基于以上规划报告章节大纲，按 system prompt 约定的 JSON 格式输出。"
    )


def _build_crosstab_plan_revision_query(
    questionnaire_text: str,
    current_parts: list[dict],
    user_text: str,
) -> str:
    """跑数表模式：章节大纲的修订 query。"""
    q_text = (questionnaire_text or "").strip()
    if len(q_text) > 12000:
        q_text = q_text[:12000] + "\n…（问卷过长，已截断）"
    outline = json.dumps(current_parts or [], ensure_ascii=False, indent=2)
    return (
        f"<questionnaire>\n{q_text}\n</questionnaire>\n\n"
        f"<current_outline>\n{outline}\n</current_outline>\n\n"
        f"<user_request>\n{user_text.strip()}\n</user_request>\n\n"
        "请在当前大纲基础上按用户意见调整，按 system prompt 约定的 JSON 格式输出。"
    )


def _render_crosstab_plan_card(plan: dict) -> str:
    """跑数表模式：把章节大纲 + 待确认问题渲染成给用户/历史看的 markdown。"""
    lines = ["## 报告章节大纲"]
    for i, p in enumerate(plan.get("parts", []), 1):
        scope = p.get("scope", "")
        lines.append(f"{i}. **{p['name']}**" + (f" — {scope}" if scope else ""))
    oqs = plan.get("open_questions") or []
    if oqs:
        lines.append("")
        lines.append("## 待确认问题")
        for q in oqs:
            lines.append(f"- {q}")
    return "\n".join(lines)


def _json_loads_loose(raw: str) -> tuple[dict | None, str]:
    raw = (raw or "").strip()
    if not raw:
        return None, "empty"
    candidates = [raw]
    fenced = re.search(r"```(?:json)?\s*(.*?)```", raw, re.S | re.I)
    if fenced:
        candidates.insert(0, fenced.group(1).strip())
    brace_start = raw.find("{")
    brace_end = raw.rfind("}")
    if 0 <= brace_start < brace_end:
        candidates.append(raw[brace_start:brace_end + 1])

    for text in candidates:
        try:
            obj = json.loads(text)
            if isinstance(obj, dict):
                return obj, ""
            return None, f"json root is {type(obj).__name__}"
        except Exception as e:
            last_err = str(e)
    return None, last_err[:180] if "last_err" in locals() else "invalid json"


def _cluster_diag_column(col_idx: int, col_name: str, total: int, batches: int) -> dict:
    return {
        "col_index": col_idx,
        "col_name": col_name,
        "total": total,
        "batches": batches,
        "phase_a": [],
        "phase_b": {},
        "phase_c": [],
        "status": "running",
        "reason": "",
        "themes": 0,
        "classifications": 0,
        "assignments": 0,
    }


async def _batch_qualitative_analysis(
    open_text: dict,
    plan: dict,
    headers: list,
    session_id: str,
):
    """大样本定性分析四阶段批处理。

    异步生成器，yield ("progress", msg) 或 ("result", clustered_themes)。
    clustered_themes 结构：
    {
        col_idx: {
            "col_name": str,
            "total": int,
            "themes": [{"id","name","description","count","percentage",
                        "positive_count","positive_pct","positive_summary",
                        "negative_count","negative_pct","negative_summary",
                        "quotes": [str]}],
            "other_themes": [{"name","count","percentage"}]
        }
    }
    """
    import json as _json
    from app.integrations.dify_client import workflow_run, STOP_SIGNAL

    clustered_themes: dict = {}
    diagnostics: dict[str, dict] = {}

    for col_idx, entries in open_text.items():
        col = next((c for c in plan["columns"] if c["index"] == col_idx), None)
        col_name = (col and col.get("name")) or (
            headers[col_idx] if col_idx < len(headers) else f"列{col_idx}"
        )
        total = len(entries)
        yield ("progress", f"【{col_name}】开始分析（共 {total} 条）")

        # ── Phase A：分批提取主题候选 ──────────────────────────────────────
        batches = [entries[i:i + BATCH_SIZE] for i in range(0, total, BATCH_SIZE)]
        all_candidates: list[dict] = []
        diag = _cluster_diag_column(col_idx, col_name, total, len(batches))
        diagnostics[str(col_idx)] = diag

        for bi, batch in enumerate(batches, 1):
            yield ("progress", f"【{col_name}】提取主题（批次 {bi}/{len(batches)}）")
            responses_text = "\n".join(
                f"[{i}] {e.get('text', '')}" for i, e in enumerate(batch)
            )
            raw = await workflow_run(
                inputs={
                    "question": col_name,
                    "responses": responses_text,
                    "count": len(batch),
                },
                api_key=DIFY_THEME_EXTRACT_KEY,
                log_prefix=f"A col={col_idx} batch={bi}",
            )
            phase_a = {"batch": bi, "raw_len": len(raw or ""), "parsed": False, "themes": 0, "error": ""}
            if raw == STOP_SIGNAL:
                phase_a["error"] = "Dify returned 400 / STOP_SIGNAL"
                diag["phase_a"].append(phase_a)
                diag["status"] = "failed"
                diag["reason"] = phase_a["error"]
                yield ("progress", f"【{col_name}】主题提取遇到错误，跳过该列")
                break
            parsed, err = _json_loads_loose(raw)
            if parsed:
                themes = parsed.get("themes", [])
                if isinstance(themes, list):
                    all_candidates.extend(themes)
                    phase_a["parsed"] = True
                    phase_a["themes"] = len(themes)
                else:
                    phase_a["error"] = "`themes` is not list"
            else:
                phase_a["error"] = err
                yield ("progress", f"【{col_name}】主题提取结果解析失败（批次 {bi}），继续处理后续批次")
            diag["phase_a"].append(phase_a)

        if not all_candidates:
            diag["status"] = "failed"
            diag["reason"] = diag["reason"] or "主题提取未返回 themes"
            yield ("progress", f"【{col_name}】没有提取到主题，后续报告将尝试使用原文兜底")
            continue

        # ── Phase B：合并去重 ──────────────────────────────────────────────
        yield ("progress", f"【{col_name}】合并主题候选（共 {len(all_candidates)} 个）")
        candidates_text = "\n".join(
            f"- {t.get('name', '')}：{t.get('description', '')}"
            for t in all_candidates
        )
        raw_b = await workflow_run(
            inputs={
                "question": col_name,
                "theme_candidates": candidates_text,
                "total_responses": total,
            },
            api_key=DIFY_THEME_MERGE_KEY,
            log_prefix=f"B col={col_idx}",
        )
        diag["phase_b"] = {"raw_len": len(raw_b or ""), "parsed": False, "themes": 0, "error": ""}
        if raw_b == STOP_SIGNAL or not raw_b:
            diag["phase_b"]["error"] = "Dify returned STOP_SIGNAL" if raw_b == STOP_SIGNAL else "empty response"
            diag["status"] = "failed"
            diag["reason"] = f"主题合并失败：{diag['phase_b']['error']}"
            yield ("progress", f"【{col_name}】主题合并失败，跳过该列")
            continue
        merged, err = _json_loads_loose(raw_b)
        if not merged:
            diag["phase_b"]["error"] = err
            diag["status"] = "failed"
            diag["reason"] = f"主题合并结果解析失败：{err}"
            yield ("progress", f"【{col_name}】主题合并结果解析失败，跳过该列")
            continue
        final_themes = merged.get("themes", [])
        if isinstance(final_themes, list):
            diag["phase_b"]["parsed"] = True
            diag["phase_b"]["themes"] = len(final_themes)
            diag["themes"] = len(final_themes)
        else:
            final_themes = []
            diag["phase_b"]["error"] = "`themes` is not list"

        if not final_themes:
            diag["status"] = "failed"
            diag["reason"] = diag["reason"] or "主题合并未返回 themes"
            yield ("progress", f"【{col_name}】主题合并为空，后续报告将尝试使用原文兜底")
            continue

        theme_list_text = _json.dumps(
            [{"id": t["id"], "name": t["name"], "description": t["description"]}
             for t in final_themes],
            ensure_ascii=False,
        )

        # ── Phase C：回跑分类 ──────────────────────────────────────────────
        # counts[theme_id] = {"total": int, "pos": int, "neg": int, "neutral": int, "mixed": int}
        counts: dict[str, dict] = {t["id"]: {"total": 0, "pos": 0, "neg": 0, "neutral": 0, "mixed": 0}
                                    for t in final_themes}
        counts["other"] = {"total": 0, "pos": 0, "neg": 0, "neutral": 0, "mixed": 0}
        # quotes_pool[theme_id] = list of (sentiment, text)
        quotes_pool: dict[str, list] = {t["id"]: [] for t in final_themes}

        for bi, batch in enumerate(batches, 1):
            yield ("progress", f"【{col_name}】分类回复（批次 {bi}/{len(batches)}）")
            responses_text = "\n".join(
                f"[{i}] {e.get('text', '')}" for i, e in enumerate(batch)
            )
            raw_c = await workflow_run(
                inputs={
                    "question": col_name,
                    "theme_list": theme_list_text,
                    "responses": responses_text,
                },
                api_key=DIFY_CLASSIFY_KEY,
                log_prefix=f"C col={col_idx} batch={bi}",
            )
            phase_c = {"batch": bi, "raw_len": len(raw_c or ""), "parsed": False, "classifications": 0, "assignments": 0, "error": ""}
            if raw_c == STOP_SIGNAL:
                phase_c["error"] = "Dify returned 400 / STOP_SIGNAL"
                diag["phase_c"].append(phase_c)
                yield ("progress", f"【{col_name}】分类回复遇到错误（批次 {bi}），继续处理后续批次")
                continue
            cls_data, err = _json_loads_loose(raw_c)
            if not cls_data:
                phase_c["error"] = err
                diag["phase_c"].append(phase_c)
                yield ("progress", f"【{col_name}】分类结果解析失败（批次 {bi}），继续处理后续批次")
                continue

            classifications = cls_data.get("classifications", [])
            if not isinstance(classifications, list):
                phase_c["error"] = "`classifications` is not list"
                diag["phase_c"].append(phase_c)
                continue
            phase_c["parsed"] = True
            phase_c["classifications"] = len(classifications)
            diag["classifications"] += len(classifications)

            for item in classifications:
                try:
                    resp_idx = int(str(item.get("response_id", "")).strip("[]"))
                    original_text = batch[resp_idx].get("text", "") if resp_idx < len(batch) else ""
                except (ValueError, IndexError):
                    continue

                assignments = item.get("assignments", [])
                if isinstance(assignments, list):
                    phase_c["assignments"] += len(assignments)
                for assign in assignments:
                    tid = assign.get("theme_id", "other")
                    sentiment = assign.get("sentiment", "neutral")
                    if tid not in counts:
                        tid = "other"
                    counts[tid]["total"] += 1
                    if sentiment == "positive":
                        counts[tid]["pos"] += 1
                    elif sentiment == "negative":
                        counts[tid]["neg"] += 1
                    elif sentiment == "mixed":
                        counts[tid]["mixed"] += 1
                    else:
                        counts[tid]["neutral"] += 1
                    if tid != "other" and len(quotes_pool[tid]) < 10 and original_text:
                        quotes_pool[tid].append((sentiment, original_text))
            diag["assignments"] += phase_c["assignments"]
            diag["phase_c"].append(phase_c)

        # ── 统计汇总 ──────────────────────────────────────────────────────
        total_mentions = sum(v["total"] for v in counts.values())
        if total_mentions == 0:
            diag["status"] = "failed"
            diag["reason"] = "分类阶段未产生任何主题归属"
            yield ("progress", f"【{col_name}】分类未产生有效归属，后续报告将尝试使用原文兜底")
            continue

        themes_out = []
        other_themes_out = []

        for t in final_themes:
            tid = t["id"]
            c = counts[tid]
            cnt = c["total"]
            pct = round(cnt / total_mentions * 100, 1)
            pos_cnt = c["pos"]
            neg_cnt = c["neg"]
            pos_pct = round(pos_cnt / cnt * 100, 1) if cnt else 0.0
            neg_pct = round(neg_cnt / cnt * 100, 1) if cnt else 0.0

            # 代表性引用：每种情感最多取 1-2 条
            pool = quotes_pool.get(tid, [])
            pos_q = [txt for sent, txt in pool if sent == "positive"][:2]
            neg_q = [txt for sent, txt in pool if sent == "negative"][:2]
            neu_q = [txt for sent, txt in pool if sent not in ("positive", "negative")][:2]
            quotes = (pos_q + neg_q + neu_q)[:6]
            # 不足 3 条时从 pool 中补充未使用的原文，保证 Writer 有足够素材
            if len(quotes) < 3:
                used = set(quotes)
                extras = [txt for _, txt in pool if txt not in used]
                quotes = quotes + extras[:max(0, 3 - len(quotes))]

            entry = {
                "id": tid,
                "name": t["name"],
                "description": t.get("description", ""),
                "count": cnt,
                "percentage": pct,
                "positive_count": pos_cnt,
                "positive_pct": pos_pct,
                "positive_summary": t.get("positive_summary") or "",
                "negative_count": neg_cnt,
                "negative_pct": neg_pct,
                "negative_summary": t.get("negative_summary") or "",
                "quotes": quotes,
            }
            if pct < OTHER_THEME_PCT:
                other_themes_out.append({"name": t["name"], "count": cnt, "percentage": pct})
            else:
                themes_out.append(entry)

        themes_out.sort(key=lambda x: x["count"], reverse=True)
        other_themes_out.sort(key=lambda x: x["count"], reverse=True)

        clustered_themes[col_idx] = {
            "col_name": col_name,
            "total": total,
            "themes": themes_out,
            "other_themes": other_themes_out,
        }
        diag["status"] = "ok"
        diag["reason"] = ""
        yield ("progress", f"【{col_name}】分析完成，识别 {len(themes_out)} 个主要主题")

    yield ("diagnostics", diagnostics)
    yield ("result", clustered_themes)


def _extract_satisfaction_stats(stats_md: str) -> str:
    """Extract ## sections whose title contains '满意度' from stats markdown."""
    lines = stats_md.split("\n")
    sections: list[list[str]] = []
    current: list[str] = []
    capturing = False

    for line in lines:
        if line.startswith("## "):
            if capturing and current:
                sections.append(current)
            current = [line]
            capturing = "满意度" in line
        elif capturing:
            current.append(line)

    if capturing and current:
        sections.append(current)

    return "\n\n".join("\n".join(s) for s in sections)


def _entry_identity(entry: dict) -> str:
    parts = []
    ids = entry.get("ids") or {}
    profile = entry.get("profile") or {}
    for k, v in ids.items():
        if str(v).strip():
            parts.append(f"{k}={v}")
    for k, v in profile.items():
        if str(v).strip():
            parts.append(f"{k}={v}")
    return "；".join(parts)


def _clip_text(text: str, limit: int = 420) -> str:
    text = str(text or "").strip().replace("\r", " ").replace("\n", " ")
    return text if len(text) <= limit else text[:limit].rstrip() + "..."


def _sample_open_entries(entries: list[dict], limit: int = 60) -> list[dict]:
    if len(entries) <= limit:
        return entries
    if limit <= 1:
        return entries[:1]
    step = (len(entries) - 1) / (limit - 1)
    idxs = []
    seen = set()
    for i in range(limit):
        idx = round(i * step)
        if idx not in seen:
            seen.add(idx)
            idxs.append(idx)
    return [entries[i] for i in idxs]


def _build_open_text_fallback_md(
    open_text: dict | None,
    clustered_themes: dict,
    plan: dict,
    headers: list[str],
) -> str:
    """Build a deterministic raw-text fallback for open columns without themes."""
    if not open_text:
        return ""

    clustered_keys = {str(k) for k in (clustered_themes or {}).keys()}
    blocks = []
    total_chars = 0
    max_chars = 45000

    for raw_idx, entries in open_text.items():
        idx_key = str(raw_idx)
        if idx_key in clustered_keys:
            continue
        if not entries:
            continue
        try:
            col_idx = int(raw_idx)
        except (TypeError, ValueError):
            col_idx = raw_idx
        col = next((c for c in plan.get("columns", []) if c.get("index") == col_idx), None)
        name = (col and col.get("name")) or (
            headers[col_idx] if isinstance(col_idx, int) and col_idx < len(headers) else f"列{raw_idx}"
        )
        lines = [f"### {name}（列 {raw_idx}，共 {len(entries)} 条非空回答；以下为抽样原文）"]
        for i, entry in enumerate(_sample_open_entries(entries), 1):
            ident = _entry_identity(entry)
            prefix = f"{i}. "
            if ident:
                prefix += f"[{ident}] "
            lines.append(prefix + _clip_text(entry.get("text", "")))
        block = "\n".join(lines)
        if total_chars + len(block) > max_chars:
            blocks.append("### 其余开放题\n（原文较多，已达到兜底上下文上限，未继续展开。）")
            break
        blocks.append(block)
        total_chars += len(block)

    if not blocks:
        return ""
    return (
        "<open_text_fallback>\n"
        "以下开放题未能产出稳定聚类结果。请仅基于这些真实原文做定性归纳和代表性引用；"
        "不要编造精确主题占比或人数。若内容明显属于年龄、性别、地区等画像补充项，不要当作体验观点展开。\n\n"
        + "\n\n".join(blocks)
        + "\n</open_text_fallback>"
    )


def _build_large_sample_writer_query(
    stats_md: str,
    clustered_themes: dict,
    plan: dict,
    headers: list[str],
    open_text: dict | None = None,
) -> str:
    parts_lines = ["  Part 1 受访者画像（固定）"]
    for i, p in enumerate(plan["parts"], 2):
        if "column_indexes" in p:
            col_names = []
            for idx in p["column_indexes"]:
                col = next((c for c in plan["columns"] if c["index"] == idx), None)
                nm = (col and col.get("name")) or (headers[idx] if idx < len(headers) else f"列{idx}")
                rl = col["role"] if col else "?"
                col_names.append(f"{nm}({rl})")
            parts_lines.append(f"  Part {i} {p['name']}: {'; '.join(col_names)}")
        else:
            scope = p.get("scope", "")
            parts_lines.append(f"  Part {i} {p['name']}" + (f": {scope}" if scope else ""))
    plan_summary = "<plan>\n报告结构：\n" + "\n".join(parts_lines) + "\n</plan>"

    theme_blocks = []
    for col_idx, data in clustered_themes.items():
        col_name = data["col_name"]
        total = data["total"]
        lines = [f"### 问题：{col_name}（共 {total:,} 条有效回答）\n"]

        for i, t in enumerate(data["themes"], 1):
            lines.append(f"**主题{i}：{t['name']}**（提及 {t['count']:,} 人次，占 {t['percentage']}%）")
            if t["positive_summary"] or t["positive_count"]:
                lines.append(f"- 正面（{t['positive_count']:,} / {t['positive_pct']}%）：{t['positive_summary']}")
            if t["negative_summary"] or t["negative_count"]:
                lines.append(f"- 负面（{t['negative_count']:,} / {t['negative_pct']}%）：{t['negative_summary']}")
            if t["quotes"]:
                lines.append(f"- 代表原文（请在报告中完整引用，并附中文翻译）：")
                for q in t["quotes"]:
                    lines.append(f'  > "{q}"')
            lines.append("")

        if data["other_themes"]:
            other_parts = "、".join(
                f"{o['name']}（{o['percentage']}%）" for o in data["other_themes"]
            )
            lines.append(f"**其他声音**（合计占比较低）：{other_parts}")

        theme_blocks.append("\n".join(lines))

    open_text_md = (
        "<open_text_themes>\n" + "\n\n".join(theme_blocks) + "\n</open_text_themes>"
        if theme_blocks else "<open_text_themes>（无开放题聚类结果）</open_text_themes>"
    )
    fallback_md = _build_open_text_fallback_md(open_text, clustered_themes, plan, headers)

    satisfaction_md = _extract_satisfaction_stats(stats_md)
    priority_block = (
        f"<priority_metrics>\n{satisfaction_md}\n</priority_metrics>\n\n"
        if satisfaction_md else ""
    )
    requirements = _get_large_sample_writer_requirements(has_satisfaction=bool(satisfaction_md))

    return (
        "**任务**：基于以下大样本问卷分析结果撰写调研报告。\n\n"
        f"{plan_summary}\n\n"
        f"<stats>\n{stats_md}\n</stats>\n\n"
        f"{priority_block}"
        f"{open_text_md}\n\n"
        + (f"{fallback_md}\n\n" if fallback_md else "")
        + f"**要求**：\n{requirements}"
    )


def _get_large_sample_writer_requirements(has_satisfaction: bool = False) -> str:
    satisfaction_rule = (
        "   - **满意度优先原则**：`<priority_metrics>` 中已提取满意度数据，必须将其作为核心结论中最靠前的 1-2 条展示，须包含具体数字"
        if has_satisfaction else
        "   - **满意度优先原则**：若报告中存在任何与满意度评分/评价相关的数据（如好评率、满意度评分、认可度等），必须将其作为核心结论中最靠前的 1-2 条展示，且须包含具体数字"
    )
    return f"""一、报告结构（严格按此顺序，不得调换）
1. **## 核心结论**（必须是第一个二级章节）
   - 列出整份报告中最重要的 5-8 条发现，每条一行，格式：「**结论标题**：具体说明（含数字）」
   - 覆盖所有 Part 的关键洞察，让读者读完此节即可掌握全部重点
{satisfaction_rule}
2. **## Part 1 受访者画像**（固定为第一个 Part，紧接核心结论之后）
   - 画像分布数据用 Markdown 表格呈现（列：维度 / 选项 / 人数占比），不要纯文字罗列
   - 表格之后用 1-2 句话解读画像特征
3. 其余 Part 按方案顺序逐章展开
4. **## 行动建议**（最后一节，3-5 条，每条必须有对应数据依据）

二、结论驱动
- 以"多少人持有什么观点"为核心叙事框架
- 每个结论必须附具体数字（人数或占比），禁止使用"部分用户""少数玩家"等模糊表述

三、主观题原文展示（关键）
- 每个主题/观点至少引用 3 条代表性玩家原文
- 展示格式：先展示原始语言原文（用引号括起），下方紧跟中文翻译（若原文已是中文则免翻译）
  示例：
  > "She's very outdated compared to other mage heroes."（该英雄与其他法师相比显得十分过时。）
  > "Modelnya kurang dipoles."（模型精致度不足。）
  > "模型感觉太老了，需要 revamp。"
- 引用的原文要能支撑该主题的核心论断，优先选择信息量最丰富的
- 若某主题可用原文不足 3 条，则展示全部可用原文，不要编造或重复引用

四、语言风格
- 简洁直接，去掉冗长铺垫和过渡句
- 报告语言为中文；玩家原文保留原语种并附中文翻译"""


def _writer_parts_meta(plan: dict, headers: list[str]) -> list[dict]:
    """返回 [{'i','name','col_desc'}]，供分轮生成时逐 Part 取标题与列说明。"""
    meta = []
    for i, p in enumerate(plan["parts"], 1):
        col_names = []
        for idx in p["column_indexes"]:
            col = next((c for c in plan["columns"] if c["index"] == idx), None)
            name = (col and col.get("name")) or (headers[idx] if idx < len(headers) else f"列{idx}")
            role = col["role"] if col else "?"
            col_names.append(f"{name}({role})")
        meta.append({"i": i, "name": p["name"], "col_desc": "; ".join(col_names)})
    return meta


def _build_writer_context(stats_md: str, open_text: dict, plan: dict, headers: list[str]) -> tuple[str, str, str]:
    """构造 Writer 的完整上下文：(plan_summary, open_text_md, requirements)。
    plan_summary/open_text/stats 仅在多轮生成的第 1 轮发送一次，后续轮次复用会话历史。"""
    parts_meta = _writer_parts_meta(plan, headers)
    parts_lines = [f"  Part {m['i']} {m['name']}: {m['col_desc']}" for m in parts_meta]
    plan_summary = "<plan>\n报告结构：\n" + "\n".join(parts_lines) + "\n</plan>"

    open_text_blocks = []
    for col_idx, texts in open_text.items():
        col = next((c for c in plan["columns"] if c["index"] == col_idx), None)
        name = (col and col.get("name")) or (headers[col_idx] if col_idx < len(headers) else f"列{col_idx}")
        joined_lines = []
        for entry in texts:
            ids = entry.get("ids", {})
            mlbb_vals = [str(v).strip() for k, v in ids.items() if "mlbb" in str(k).casefold() and str(v).strip()]
            player_vals = [str(v).strip() for k, v in ids.items() if "mlbb" not in str(k).casefold() and str(v).strip()]
            id_parts = []
            if player_vals:
                id_parts.append(f"玩家ID={' / '.join(player_vals)}")
            if mlbb_vals:
                id_parts.append(f"MLBBID={' / '.join(mlbb_vals)}")
            profile_str = " / ".join(f"{k}={v}" for k, v in entry.get("profile", {}).items())
            prefix = " | ".join(filter(None, [" | ".join(id_parts), f"画像={profile_str}" if profile_str else ""]))
            text_val = entry.get("text", "")
            joined_lines.append(f"- {f'[{prefix}] ' if prefix else ''}{text_val}")
        joined = "\n".join(joined_lines)
        open_text_blocks.append(f"### {name}（列 {col_idx}, 共 {len(texts)} 条非空回答）\n{joined}")

    open_text_md = (
        "<open_text>\n" + "\n\n".join(open_text_blocks) + "\n</open_text>"
        if open_text_blocks else "<open_text>（本问卷没有开放题）</open_text>"
    )

    requirements = _get_writer_requirements()
    requirements += (
        "\n\n补充：引用玩家原文时必须沿用 `<open_text>` 前缀里的玩家身份信息。"
        "`玩家ID=...` 和 `MLBBID=...` 是两个独立身份字段，报告表格中要拆成 `玩家ID`、`MLBBID` 两列，单元格只放值。"
        "如果出现 `MLBBID=123456(57001)`，括号内是区服编号，必须作为 MLBBID 的一部分展示，"
        "不得拆到「画像信息」或其它列；只有 `<open_text>` 前缀里真的存在 `画像=...` 时，才展示画像信息。"
    )
    return plan_summary, open_text_md, requirements


def _build_writer_first_query(stats_md: str, open_text: dict, plan: dict, headers: list[str]) -> str:
    """多轮生成第 1 轮：发送全部上下文 + 要求，但本轮只让模型输出一级标题。"""
    plan_summary, open_text_md, requirements = _build_writer_context(stats_md, open_text, plan, headers)
    return (
        "**协作方式**：本次报告将**分多轮**生成。下面先给你全部数据（<plan> 报告结构、<stats> 确定性统计、"
        "<open_text> 全部开放题原文）和完整的写作要求。请通读并牢记——后续每一轮我会指定你写其中**某一个章节**，"
        "你要从这些数据里取材，但**每轮只写我当轮指定的部分，绝不提前写其它章节**。\n\n"
        f"{plan_summary}\n\n"
        f"<stats>\n{stats_md}\n</stats>\n\n"
        f"{open_text_md}\n\n"
        f"<report_spec>\n以下是整篇报告最终要满足的写作要求（供你理解全局，后续逐轮执行）：\n{requirements}\n</report_spec>\n\n"
        "**本轮任务（第 1 轮）**：**只**输出报告的一级标题（`# 一级标题`）。"
        "如果 <stats> 或 metadata 中存在「被排除」样本依据，可在标题下另起一行用一句话说明依据。"
        "除此之外**什么都不要写**——不要写核心结论、不要写任何 Part、不要写 Bug 模块、不要写本节总结。"
        "确认你已读完全部数据，本轮输出仅一级标题。"
    )


def _build_writer_part_query(part: dict) -> str:
    """多轮生成中的某个 Part 轮：仅指示写这一个 Part。原文已在会话历史中。"""
    return (
        f"**本轮任务**：现在**只**写 `## Part {part['i']} {part['name']}` 这一个章节的完整内容"
        f"（涉及列：{part['col_desc']}）。\n"
        "严格按 <report_spec> 里对 Part 的写法：紧接 `## Part` 标题后先写一段详尽的「本节总结」段落（连贯文字、不用列表），"
        "再按题目逐一展开；开放题归纳必须用 `### 题目名` 三级标题，其下用 `#### 正面观点`/`#### 负面观点`/`#### 中立 / 建议` 分组，"
        "每个观点用 `**观点：短标题**` + `提及情况：` + `代表性原话：`（小表格）的固定结构，ID 展示规则照 <report_spec>。\n"
        "**约束**：① 只输出这一个 Part，不要写其它 Part；② 不要写核心结论、不要写 Bug 模块；"
        "③ 不要重复前面已经写过的标题或章节；④ 所有数字、百分比必须逐字取自 <stats>，禁止重算或编造。"
    )


def _build_writer_bug_query() -> str:
    """多轮生成的 Bug 模块轮：需要则只输出该模块，否则只回 NONE。"""
    return (
        "**本轮任务**：现在通览 <open_text> 里的**全部**开放反馈，按 <report_spec> 第 8 条判断是否需要 "
        "`## Bug 或待确认问题` 模块（仅当确有疑似功能 bug、体验异常、规则不明、玩家无法判断是否设计如此的问题时才需要）。\n"
        "- 若需要：**只**输出该模块——以 `## Bug 或待确认问题` 开头，下接 Markdown 表格，字段固定为 "
        "`问题类型`、`待确认问题`、`玩家信息`、`玩家原文翻译`，不要输出任何其它章节或解释。\n"
        "- 若不需要：**只**回复一个词 `NONE`，不要输出任何其它内容、不要解释。"
    )


def _build_writer_core_query(parts_meta: list[dict], has_bug: bool) -> str:
    """多轮生成的核心结论轮：基于已生成全部章节回写核心结论模块（放在报告顶部）。"""
    part_titles = "、".join(f"Part {m['i']} {m['name']}" for m in parts_meta)
    bug_clause = (
        "正文包含 `## Bug 或待确认问题` 模块，因此核心结论**最后必须**追加 `### 待确认问题概述`，只概述问题类型、不展开原文。"
        if has_bug else
        "正文没有 Bug 模块，因此核心结论**不要**写任何「待确认问题」相关小节。"
    )
    return (
        "**本轮任务（最后一轮）**：基于你前面已经生成的全部章节，撰写整篇报告的「核心结论」模块。"
        "这个模块最终会被放到报告**最顶部**（一级标题之后、各 Part 之前），所以请独立、完整地写出来。\n"
        "严格按 <report_spec> 里『核心结论』部分的格式：用 `<!--CORE_START-->` 和 `<!--CORE_END-->` 两个标记"
        "**各自独占一行**包裹整段，内部依次写：`## 核心结论`（首行写明样本总数）、`### 总体判断`、"
        f"逐个写 `### Part X 章节名：关键发现`（必须引用真实 Part 名：{part_titles}）、按需的 `### 高信号少数观点与风险`。\n"
        f"{bug_clause}\n"
        "**约束**：① 只输出从 `<!--CORE_START-->` 到 `<!--CORE_END-->` 的内容，不要重复正文章节、不要再写一级标题；"
        "② 核心结论里不使用百分比、不使用精确人数，改用量级描述（样本总数可引用）；"
        "③ 引用的绝对数值必须与 <stats> 一致。"
    )


def _format_rows_for_qa(rows: list[list], plan: dict) -> str:
    QA_MAX = 60000
    if not rows or len(rows) <= 1:
        return "（无数据）"
    headers = rows[0]
    body = rows[1:]
    total = len(body)
    col_names = [(h or "").strip() or f"col_{i}" for i, h in enumerate(headers)]

    def row_obj(row):
        return {col_names[i]: (row[i] if i < len(row) else "") for i in range(len(col_names))}

    dump = "\n".join(json.dumps(row_obj(r), ensure_ascii=False) for r in body)
    if len(dump) > QA_MAX:
        pidxs = [c["index"] for c in plan.get("columns", []) if c.get("role") == "profile_dim"]
        sampled = _stratified_sample(body, pidxs, 100)
        note = (
            f"# 原始数据共 {total} 行，超出上下文上限，已按画像维度分层抽样到 {len(sampled)} 行。\n\n"
        )
        dump = note + "\n".join(json.dumps(row_obj(r), ensure_ascii=False) for r in sampled)
    return dump


def _stratified_sample(body: list[list], profile_indexes: list[int], target: int = 100):
    if not profile_indexes or len(body) <= target:
        return body[:target]

    def key(row):
        return tuple(row[i] if i < len(row) else "" for i in profile_indexes)

    buckets: dict = {}
    for r in body:
        buckets.setdefault(key(r), []).append(r)

    out: list = []
    total = len(body)
    for items in buckets.values():
        share = max(1, round(len(items) / total * target))
        out.extend(items[:share])
        if len(out) >= target:
            break
    return out[:target]


# ============================================================
# Markdown → Word
# ============================================================

def _add_formatted_run(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    result = []
    for line in lines:
        if re.match(r'^\|[\s\-|:]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        result.append(cells)
    return result


def markdown_to_docx(md_text: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(level=1); h.clear()
            _add_formatted_run(h, line[2:].strip())
        elif line.startswith('## ') and not line.startswith('### '):
            h = doc.add_heading(level=2); h.clear()
            _add_formatted_run(h, line[3:].strip())
        elif line.startswith('### '):
            h = doc.add_heading(level=3); h.clear()
            _add_formatted_run(h, line[4:].strip())
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i]); i += 1
            table_data = _parse_md_table(table_lines)
            if table_data:
                num_cols = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=num_cols)
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(table_data):
                    for ci in range(num_cols):
                        ct = row_data[ci].replace('\\|', '|') if ci < len(row_data) else ""
                        cell = tbl.cell(ri, ci); cell.text = ""
                        run = cell.paragraphs[0].add_run(ct)
                        if ri == 0: run.bold = True
            continue
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet'); p.clear()
            _add_formatted_run(p, line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number'); p.clear()
            _add_formatted_run(p, re.sub(r'^\d+\.\s', '', line).strip())
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_run(p, line[2:].strip())
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph()
        elif not line.strip():
            pass
        else:
            p = doc.add_paragraph(); _add_formatted_run(p, line.strip())

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# SSE 工具
# ============================================================

def _annotate_ai_log(message: str, **fields) -> None:
    payload = " ".join(f"{k}={v!r}" for k, v in fields.items())
    print(f"[annotate.ai_detect] {message}" + (f" {payload}" if payload else ""), flush=True)


async def sse_dify_stream(
    query: str,
    user_id: str,
    conversation_id: str,
    api_key: str,
) -> AsyncGenerator[tuple[str, str], None]:
    import httpx

    payload = {
        "inputs": {},
        "query": query,
        "response_mode": "streaming",
        "user": user_id,
        "conversation_id": conversation_id,
    }
    req_headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    final_conv_id = conversation_id

    async with httpx.AsyncClient(timeout=1800, follow_redirects=True) as client:
        async with client.stream(
            "POST", f"{DIFY_API_BASE}/chat-messages",
            headers=req_headers, json=payload,
        ) as resp:
            if resp.status_code >= 400:
                err = await resp.aread()
                raise RuntimeError(f"Dify {resp.status_code}: {err.decode('utf-8', errors='replace')}")
            async for line in resp.aiter_lines():
                if not line or not line.startswith("data:"):
                    continue
                raw = line[5:].strip()
                if not raw:
                    continue
                try:
                    data = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                event = data.get("event")
                if event in ("message", "agent_message"):
                    yield data.get("answer", ""), ""
                if data.get("conversation_id"):
                    final_conv_id = data["conversation_id"]
                if event == "error":
                    raise RuntimeError(f"Dify error: {data.get('code')} {data.get('message')}")
                if event in ("message_end", "workflow_finished", "agent_message_end"):
                    break

    yield "", final_conv_id


async def sse_dify_completion_stream(
    query: str,
    user_id: str,
    api_key: str,
    input_var: str = "survey_batch",
) -> AsyncGenerator[str, None]:
    """调用 Dify completion-messages 端点（文本生成应用）。"""
    import httpx

    inputs = {input_var: query}
    for fallback_key in ("query", "text", "content"):
        inputs.setdefault(fallback_key, query)

    payload = {
        "inputs": inputs,
        "response_mode": "streaming",
        "user": user_id,
    }
    req_headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    async with httpx.AsyncClient(timeout=1800, follow_redirects=True) as client:
        async with client.stream(
            "POST", f"{DIFY_API_BASE}/completion-messages",
            headers=req_headers, json=payload,
        ) as resp:
            if resp.status_code >= 400:
                err = await resp.aread()
                raise RuntimeError(f"Dify {resp.status_code}: {err.decode('utf-8', errors='replace')}")
            async for line in resp.aiter_lines():
                if not line or not line.startswith("data:"):
                    continue
                raw = line[5:].strip()
                if not raw:
                    continue
                try:
                    data = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                event = data.get("event")
                if event == "message":
                    yield data.get("answer", "")
                if event == "error":
                    raise RuntimeError(f"Dify error: {data.get('code')} {data.get('message')}")
                if event in ("message_end", "workflow_finished"):
                    break


def _looks_like_dify_endpoint_mismatch(err: Exception) -> bool:
    text = str(err).lower()
    return any(k in text for k in [
        "completion-messages",
        "chat-messages",
        "app mode",
        "app_mode",
        "not support",
        "not supported",
        "invalid_param",
        "endpoint",
        "completion app",
        "chat app",
    ])


async def call_dify_compatible(
    query: str,
    user_id: str,
    api_key: str,
    conversation_id: str = "",
) -> tuple[str, str, str, str]:
    """优先按 chat 应用调用；若疑似端点/应用类型不匹配，自动改用 completion 应用。"""
    try:
        chunks: list[str] = []
        final_conv = conversation_id
        async for chunk, conv_id in sse_dify_stream(query, user_id, conversation_id, api_key):
            if chunk:
                chunks.append(chunk)
            if conv_id:
                final_conv = conv_id
        return "".join(chunks), final_conv, "chat", ""
    except Exception as e:
        if not _looks_like_dify_endpoint_mismatch(e):
            raise
        fallback_reason = str(e)

    chunks: list[str] = []
    async for chunk in sse_dify_completion_stream(query, user_id, api_key):
        if chunk:
            chunks.append(chunk)
    return "".join(chunks), "", "completion", fallback_reason


# ============================================================
# 帖子评论舆情分析（单 Dify Workflow，mode 路由 + Python 并发编排）
# ============================================================

def _comment_dify_inputs(mode: str, **kw) -> dict:
    """构造评论分析 Workflow 的输入；START 节点变量需全部带上。"""
    return {
        "mode": mode,
        "post_title": kw.get("post_title", ""),
        "post_content": kw.get("post_content", ""),
        "comments_json": kw.get("comments_json", ""),
        "themes_json": kw.get("themes_json", ""),
    }


def _comment_selected_raw_comments_md(items: list[dict]) -> str:
    if not items:
        return ""
    lines = [
        "## 玩家评论原文精选",
        "",
        "以下评论来自清洗后字数较长的评论候选池，并经模型筛选为与帖子主题相关、表达较完整的玩家反馈。内容已翻译为简体中文。",
        "",
    ]
    rendered_count = 0
    for idx, item in enumerate(items, start=1):
        translation = str(item.get("translation") or "").strip()
        if not translation:
            continue
        rendered_count += 1
        compact_translation = " ".join(line.strip() for line in translation.splitlines() if line.strip())
        lines.append(f"{idx}. {compact_translation}")
        lines.append("")
    if rendered_count <= 0:
        return ""
    return "\n".join(lines).strip()


def _comment_report_without_raw_comments(md: str) -> str:
    text = (md or "").strip()
    marker = "## 玩家评论原文精选"
    pos = text.find(marker)
    if pos >= 0:
        return text[:pos].strip()
    return text


def _comment_append_selected_raw_comments(report_md: str, selected_raw_comments: list[dict]) -> str:
    selected_md = _comment_selected_raw_comments_md(selected_raw_comments)
    base = _comment_report_without_raw_comments(report_md)
    if not selected_md:
        return base
    return (base + "\n\n" + selected_md).strip()


async def _select_comment_raw_quotes(sess: dict) -> list[dict]:
    """从清洗后的长评论候选中精选最多 50 条原文评论。失败由调用方隔离。"""
    import json as _json
    from app.integrations.dify_client import workflow_run, STOP_SIGNAL

    if not DIFY_COMMENT_ANALYSIS_KEY:
        raise RuntimeError("未配置 DIFY_COMMENT_ANALYSIS_KEY")

    post_title = sess.get("comment_post_title", "") or ""
    post_content = sess.get("comment_post_content", "") or ""
    long_candidates = [
        str(x or "").strip()
        for x in (sess.get("comment_long_candidates") or [])
        if str(x or "").strip()
    ][:comment_analysis.LONG_CANDIDATE_MAX_N]
    if not long_candidates:
        return []

    key = DIFY_COMMENT_ANALYSIS_KEY
    sem = asyncio.Semaphore(COMMENT_QUOTE_SELECT_CONCURRENCY)

    def _batch_to_json(batch: list) -> str:
        return _json.dumps(batch, ensure_ascii=False)

    def _first_nonempty(item: dict, keys: list[str]) -> str:
        for key in keys:
            value = item.get(key)
            if value is None:
                continue
            text_value = str(value).strip()
            if text_value:
                return text_value
        return ""

    def _has_cjk(text: str) -> bool:
        return any("\u4e00" <= ch <= "\u9fff" for ch in text or "")

    def _normalize_quote_item(item, fallback_idx: int | None = None) -> dict | None:
        if not isinstance(item, dict):
            return None
        try:
            idx = int(item.get("idx", item.get("id", fallback_idx)))
        except (TypeError, ValueError):
            return None
        original_aliases = [
            "text",
            "original",
            "original_text",
            "originalText",
            "raw_text",
            "rawText",
            "comment",
            "comment_text",
            "commentText",
            "原文",
            "原始评论",
            "评论原文",
        ]
        translation_aliases = [
            "translation",
            "zh_translation",
            "translation_zh",
            "chinese_translation",
            "chineseTranslation",
            "translated_text",
            "translatedText",
            "translated_comment",
            "translatedComment",
            "comment_translation",
            "commentTranslation",
            "comment_zh",
            "commentZh",
            "text_zh",
            "textZh",
            "translated",
            "zh",
            "cn",
            "chinese",
            "Chinese translation",
            "Chinese Translation",
            "中文",
            "中文翻译",
            "中文译文",
            "评论翻译",
            "评论译文",
            "翻译",
            "译文",
            "简体中文",
        ]
        text = _first_nonempty(item, original_aliases)
        translation = _first_nonempty(item, translation_aliases)
        explicit_original = _first_nonempty(item, [k for k in original_aliases if k != "text"])
        raw_text = str(item.get("text") or "").strip()
        if not translation and raw_text and explicit_original and _has_cjk(raw_text):
            translation = raw_text
            text = explicit_original
        try:
            score = float(item.get("score", 0) or 0)
        except (TypeError, ValueError):
            score = 0.0
        return {
            "idx": idx,
            "text": text,
            "translation": translation,
            "score": score,
            "reason": str(item.get("reason") or ""),
        }

    async def _quote_batch(batch: list[dict]) -> list[dict]:
        async with sem:
            raw = await workflow_run(
                inputs=_comment_dify_inputs(
                    "quote_select_batch",
                    post_title=post_title,
                    post_content=post_content,
                    comments_json=_batch_to_json(batch),
                ),
                api_key=key,
                log_prefix="comment quote_select_batch",
            )
        if not raw or raw == STOP_SIGNAL:
            return []
        parsed, _e = comment_analysis.loads_loose(raw)
        if not isinstance(parsed, list):
            print("[comment quote_select_batch] parse failed or non-list output", flush=True)
            return []
        expected = [int(x["idx"]) for x in batch]
        expected_set = set(expected)
        out: list[dict] = []
        for i, item in enumerate(parsed):
            fallback = expected[i] if i < len(expected) else None
            normalized = _normalize_quote_item(item, fallback)
            if normalized and normalized["idx"] not in expected_set and 0 <= normalized["idx"] < len(expected):
                normalized["idx"] = expected[normalized["idx"]]
            if normalized and normalized["idx"] in expected_set:
                normalized["text"] = long_candidates[normalized["idx"]]
                normalized["translation"] = str(normalized.get("translation") or "").strip()
                out.append(normalized)
        if not out and parsed:
            print(
                f"[comment quote_select_batch] no valid idx parsed; parsed={len(parsed)} expected={len(expected)}",
                flush=True,
            )
        translated_count = sum(1 for x in out if x.get("translation"))
        print(
            f"[comment quote_select_batch] selected={len(out)} translated={translated_count}",
            flush=True,
        )
        if out and translated_count == 0:
            first_raw = parsed[0] if parsed else {}
            first_keys = list(first_raw.keys()) if isinstance(first_raw, dict) else []
            first_preview = _json.dumps(first_raw, ensure_ascii=False)[:500] if isinstance(first_raw, dict) else str(first_raw)[:500]
            print(
                f"[comment quote_select_batch] translation missing; first_keys={first_keys}; first_item={first_preview}",
                flush=True,
            )
        return sorted(out, key=lambda x: x.get("score", 0), reverse=True)[:comment_analysis.QUOTE_SELECT_BATCH_KEEP_N]

    def _fallback_selected_from_batch_candidates(
        batch_candidates: list[dict],
        require_translation: bool = False,
    ) -> list[dict]:
        selected: list[dict] = []
        seen_text: set[str] = set()
        for item in sorted(batch_candidates, key=lambda x: x.get("score", 0), reverse=True):
            text = str(item.get("text") or "").strip()
            translation = str(item.get("translation") or "").strip()
            if not text:
                continue
            if require_translation and not translation:
                continue
            key_text = text.lower()
            if key_text in seen_text:
                continue
            seen_text.add(key_text)
            selected.append({
                "text": text,
                "translation": translation,
            })
            if len(selected) >= comment_analysis.QUOTE_SELECT_FINAL_N:
                break
        return selected

    indexed_long = [{"idx": i, "text": text} for i, text in enumerate(long_candidates)]
    quote_batches = comment_analysis.make_batches(indexed_long, comment_analysis.QUOTE_SELECT_BATCH_SIZE)
    batch_candidates: list[dict] = []
    tasks = [asyncio.create_task(_quote_batch(b)) for b in quote_batches]
    try:
        for t in asyncio.as_completed(tasks):
            batch_candidates.extend(await t)
    except BaseException:
        for task in tasks:
            if not task.done():
                task.cancel()
        await asyncio.gather(*tasks, return_exceptions=True)
        raise

    by_idx: dict[int, dict] = {}
    for item in batch_candidates:
        idx = int(item.get("idx", -1))
        if idx < 0:
            continue
        old = by_idx.get(idx)
        if old is None or float(item.get("score", 0) or 0) > float(old.get("score", 0) or 0):
            by_idx[idx] = item
    batch_candidates = sorted(by_idx.values(), key=lambda x: x.get("score", 0), reverse=True)[
        :comment_analysis.QUOTE_SELECT_FINAL_POOL_N
    ]
    print(
        "[comment quote_select] "
        f"long_candidates={len(long_candidates)} "
        f"batches={len(quote_batches)} "
        f"batch_candidates={len(batch_candidates)} "
        f"batch_translated={sum(1 for x in batch_candidates if x.get('translation'))}",
        flush=True,
    )
    if not batch_candidates:
        return []

    async with sem:
        raw_final = await workflow_run(
            inputs=_comment_dify_inputs(
                "quote_select_final",
                post_title=post_title,
                post_content=post_content,
                comments_json=_json.dumps(batch_candidates, ensure_ascii=False),
            ),
            api_key=key,
            log_prefix="comment quote_select_final",
        )
    if not raw_final or raw_final == STOP_SIGNAL:
        return _fallback_selected_from_batch_candidates(batch_candidates)
    final_parsed, _e = comment_analysis.loads_loose(raw_final)
    if not isinstance(final_parsed, list):
        print("[comment quote_select_final] parse failed or non-list output; fallback to batch candidates", flush=True)
        return _fallback_selected_from_batch_candidates(batch_candidates)

    selected: list[dict] = []
    seen_text: set[str] = set()
    candidates_by_idx = {int(x["idx"]): x for x in batch_candidates if "idx" in x}
    for i, item in enumerate(final_parsed):
        fallback = int(batch_candidates[i]["idx"]) if i < len(batch_candidates) and "idx" in batch_candidates[i] else None
        normalized = _normalize_quote_item(item, fallback)
        if not normalized:
            continue
        idx = normalized["idx"]
        if idx not in candidates_by_idx and 0 <= idx < len(batch_candidates):
            idx = int(batch_candidates[idx]["idx"])
        text = str(candidates_by_idx.get(idx, {}).get("text") or normalized["text"] or "").strip()
        if not text:
            continue
        translation = normalized.get("translation") or ""
        if not translation and idx in candidates_by_idx:
            translation = str(candidates_by_idx[idx].get("translation") or "").strip()
        key_text = text.lower()
        if key_text in seen_text:
            continue
        seen_text.add(key_text)
        selected.append({
            "text": text,
            "translation": str(translation or "").strip(),
        })
        if len(selected) >= comment_analysis.QUOTE_SELECT_FINAL_N:
            break
    if not selected:
        print(
            f"[comment quote_select_final] selected empty from parsed={len(final_parsed)}; fallback to batch candidates",
            flush=True,
        )
        return _fallback_selected_from_batch_candidates(batch_candidates)
    any_translated_candidate = any(x.get("translation") for x in batch_candidates)
    selected_translated = sum(1 for x in selected if x.get("translation"))
    print(
        f"[comment quote_select_final] parsed={len(final_parsed)} selected={len(selected)} translated={selected_translated}",
        flush=True,
    )
    if selected_translated == 0 and any_translated_candidate:
        print(
            "[comment quote_select_final] final returned no translations; fallback to translated batch candidates",
            flush=True,
        )
        return _fallback_selected_from_batch_candidates(batch_candidates, require_translation=True)
    if any_translated_candidate and selected_translated < len(selected):
        selected = [x for x in selected if x.get("translation")]
        seen_text = {str(x.get("text") or "").strip().lower() for x in selected}
        for item in sorted(batch_candidates, key=lambda x: x.get("score", 0), reverse=True):
            text = str(item.get("text") or "").strip()
            translation = str(item.get("translation") or "").strip()
            key_text = text.lower()
            if not text or not translation or key_text in seen_text:
                continue
            selected.append({"text": text, "translation": translation})
            seen_text.add(key_text)
            if len(selected) >= comment_analysis.QUOTE_SELECT_FINAL_N:
                break
        print(
            f"[comment quote_select_final] filled translated selected={len(selected)}",
            flush=True,
        )
    return selected


async def _comment_analysis_pipeline(sess: dict):
    """评论舆情分析流水线。异步生成器，yield ("progress", msg) / ("result", payload)。

    流程：relevance(并发) → extract(并发) → merge → classify(并发) → 本地统计 → report。
    所有 Dify 调用走同一个 Workflow，靠 mode 路由；并发用 Semaphore 限流。
    """
    import json as _json
    from app.integrations.dify_client import workflow_run, STOP_SIGNAL

    if not DIFY_COMMENT_ANALYSIS_KEY:
        raise RuntimeError("未配置 DIFY_COMMENT_ANALYSIS_KEY")

    post_title = sess.get("comment_post_title", "") or ""
    post_content = sess.get("comment_post_content", "") or ""
    sample_pool: list[str] = list(sess.get("comment_sample", []) or [])
    if not post_title.strip():
        raise RuntimeError("请填写帖子主题后再开始评论分析")
    if not post_content.strip():
        raise RuntimeError("请填写帖子正文后再开始评论分析")
    if not sample_pool:
        raise RuntimeError("没有可分析的评论，请重新上传文件")

    key = DIFY_COMMENT_ANALYSIS_KEY
    sem = asyncio.Semaphore(COMMENT_ANALYSIS_CONCURRENCY)

    def _batch_to_json(batch: list) -> str:
        return _json.dumps(batch, ensure_ascii=False)

    # ── Phase 0：语义相关性筛选。先跑初始 5000；不足目标时从样本池补样 ─────
    sample_pool = sample_pool[:comment_analysis.SAMPLE_POOL_MAX_N]
    related_comments: list[str] = []
    relevance_results: list[dict] = []
    processed_sample_count = 0
    relevance_done = 0
    total_relevance_batches = 0
    relevance_rounds: list[dict] = []

    def _parse_bool(value) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            return value.strip().lower() in {"true", "yes", "1", "是", "相关", "direct", "implicit"}
        return bool(value)

    def _normalize_relevance_item(item, fallback_idx: int | None = None) -> dict | None:
        if not isinstance(item, dict):
            return None
        try:
            idx = int(item.get("idx", fallback_idx))
        except (TypeError, ValueError):
            return None
        if idx < 0 or idx >= len(sample_pool):
            return None
        relation = str(item.get("relation") or item.get("relevance") or "").strip().lower()
        if "is_related" in item:
            is_related = _parse_bool(item.get("is_related"))
        else:
            is_related = relation in {"direct", "implicit", "related", "相关", "直接相关", "隐含相关"}
        if relation in {"off_topic", "irrelevant", "unrelated", "not_related", "无关"}:
            is_related = False
        return {
            "idx": idx,
            "is_related": is_related,
            "relation": relation or ("direct" if is_related else "off_topic"),
            "reason": str(item.get("reason") or ""),
        }

    async def _judge_relevance(batch: list[dict], depth: int = 0):
        expected_idxs = [int(x["idx"]) for x in batch]
        expected_set = set(expected_idxs)
        last_err = ""
        for attempt in range(2):
            async with sem:
                raw = await workflow_run(
                    inputs=_comment_dify_inputs(
                        "relevance",
                        post_title=post_title,
                        post_content=post_content,
                        comments_json=_batch_to_json(batch),
                    ),
                    api_key=key,
                    log_prefix="comment relevance",
                )
            if raw == STOP_SIGNAL:
                raise RuntimeError("评论相关性筛选调用被 Dify 拒绝（HTTP 400），请检查 relevance 节点配置或输入")
            arr, _e = comment_analysis.loads_loose(raw) if raw else (None, "empty")
            if isinstance(arr, list):
                by_idx: dict[int, dict] = {}
                for i, item in enumerate(arr):
                    fallback_idx = expected_idxs[i] if i < len(expected_idxs) else None
                    normalized = _normalize_relevance_item(item, fallback_idx)
                    if normalized and normalized["idx"] in expected_set:
                        by_idx[normalized["idx"]] = normalized
                # relevance 是筛选任务：Dify 可以只返回相关评论的 idx。
                # 未返回的 idx 默认视为 off_topic，避免为了无关评论反复重试。
                return [
                    by_idx.get(idx) or {
                        "idx": idx,
                        "is_related": False,
                        "relation": "off_topic",
                        "reason": "relevance 节点未返回该评论，按无关处理",
                    }
                    for idx in expected_idxs
                ]
            got = len(arr) if isinstance(arr, list) else "非数组"
            last_err = f"返回 {got}，期望 idx 数 {len(batch)}"
            print(f"[comment relevance] 第 {attempt + 1} 次返回格式不符（{last_err}），重试中…", flush=True)
        if len(batch) > 10 and depth < 3:
            mid = len(batch) // 2
            print(
                f"[comment relevance] 批次仍不匹配，自动拆分为 {mid}+{len(batch) - mid} 条继续重试",
                flush=True,
            )
            left = await _judge_relevance(batch[:mid], depth + 1)
            right = await _judge_relevance(batch[mid:], depth + 1)
            return left + right
        raise RuntimeError(
            f"评论相关性筛选返回 idx 不完整（{last_err}），拆分重试后仍失败，已中止分析"
        )

    async def _run_relevance_slice(start_idx: int, end_idx: int, round_results: list[dict]):
        nonlocal processed_sample_count, relevance_done, total_relevance_batches
        indexed = [{"idx": i, "text": sample_pool[i]} for i in range(start_idx, end_idx)]
        batches = comment_analysis.make_batches(indexed, comment_analysis.RELEVANCE_BATCH_SIZE)
        if not batches:
            return
        total_relevance_batches += len(batches)
        yield ("progress", f"正在按帖子主题筛选相关评论（样本 {start_idx + 1}-{end_idx}，共 {len(batches)} 批）…")
        tasks = [asyncio.create_task(_judge_relevance(b)) for b in batches]
        try:
            for t in asyncio.as_completed(tasks):
                round_results.extend(await t)
                relevance_done += 1
                yield ("progress", f"相关性筛选进度 {relevance_done}/{total_relevance_batches} 批")
        except BaseException:
            for task in tasks:
                if not task.done():
                    task.cancel()
            await asyncio.gather(*tasks, return_exceptions=True)
            raise
        processed_sample_count = end_idx

    next_start = 0
    while next_start < len(sample_pool):
        next_end = min(next_start + comment_analysis.SAMPLE_MAX_N, len(sample_pool))
        before = len(related_comments)
        round_results: list[dict] = []
        async for item in _run_relevance_slice(next_start, next_end, round_results):
            yield item
        relevance_results.extend(round_results)
        relevance_by_idx = {int(x["idx"]): x for x in round_results if isinstance(x, dict)}
        related_comments.extend(
            sample_pool[i]
            for i in range(next_start, next_end)
            if relevance_by_idx.get(i, {}).get("is_related") is True
        )
        relevance_rounds.append({
            "start": next_start,
            "end": next_end,
            "sample_count": next_end - next_start,
            "related_count": len(related_comments) - before,
        })
        yield (
            "progress",
            f"本轮相关性筛选完成：新增相关评论 {len(related_comments) - before} 条，累计 {len(related_comments)} 条",
        )
        if len(related_comments) >= comment_analysis.RELATED_TARGET_N:
            break
        if next_end >= len(sample_pool):
            break
        yield (
            "progress",
            f"相关评论不足 {comment_analysis.RELATED_TARGET_N} 条，继续补抽下一批样本…",
        )
        next_start = next_end

    comments = related_comments
    related_count = len(comments)
    original_sample_count = processed_sample_count
    off_topic_count = original_sample_count - related_count
    relevance_stats = {
        "sample_count": original_sample_count,
        "sample_pool_count": len(sample_pool),
        "related_count": related_count,
        "off_topic_count": off_topic_count,
        "related_target": comment_analysis.RELATED_TARGET_N,
        "sample_cap": comment_analysis.SAMPLE_POOL_MAX_N,
        "rounds": relevance_rounds,
    }
    sess["comment_relevance_stats"] = relevance_stats
    meta = sess.get("comment_sample_meta") or {}
    meta.update({
        "topic_related_count": related_count,
        "off_topic_count": off_topic_count,
        "relevance_sample_count": original_sample_count,
        "related_target": comment_analysis.RELATED_TARGET_N,
    })
    sess["comment_sample_meta"] = meta
    yield ("progress", f"主题相关性筛选完成：使用样本 {original_sample_count} 条，保留 {related_count} 条，剔除 {off_topic_count} 条无关评论")
    if related_count <= 0:
        raise RuntimeError("未筛选出与帖子主题或正文相关的评论，无法生成可靠报告")

    extract_batches = comment_analysis.make_batches(comments, comment_analysis.BATCH_SIZE)
    n_extract_batch = len(extract_batches)

    # ── Phase 1：并发提取各批主题候选 ───────────────────────────────────
    yield ("progress", f"正在并发提取评论主题（共 {n_extract_batch} 批，并发 {COMMENT_ANALYSIS_CONCURRENCY}）…")
    extract_done = 0

    async def _extract(batch: list[str]):
        async with sem:
            raw = await workflow_run(
                inputs=_comment_dify_inputs(
                    "extract", post_title=post_title, post_content=post_content, comments_json=_batch_to_json(batch)
                ),
                api_key=key,
                log_prefix="comment extract",
            )
        if not raw or raw == STOP_SIGNAL:
            return []
        parsed, _e = comment_analysis.loads_loose(raw)
        return parsed if isinstance(parsed, list) else []

    extract_tasks = [asyncio.create_task(_extract(b)) for b in extract_batches]
    candidates: list[dict] = []
    for t in asyncio.as_completed(extract_tasks):
        part = await t
        candidates.extend(x for x in part if isinstance(x, dict))
        extract_done += 1
        yield ("progress", f"主题提取进度 {extract_done}/{n_extract_batch} 批")

    if not candidates:
        raise RuntimeError("主题提取未返回任何结果，请检查 Dify extract 节点配置")

    # ── Phase 2：合并去重主题（单次） ───────────────────────────────────
    yield ("progress", "正在汇总并合并去重主题…")
    raw_merge = await workflow_run(
        inputs=_comment_dify_inputs(
            "merge",
            post_title=post_title,
            post_content=post_content,
            themes_json=_json.dumps(candidates, ensure_ascii=False),
        ),
        api_key=key,
        log_prefix="comment merge",
    )
    merged, _e = comment_analysis.loads_loose(raw_merge)
    if not isinstance(merged, list) or not merged:
        raise RuntimeError("主题合并未返回有效结果，请检查 Dify merge 节点配置")

    # Dify 用 theme_id / theme_name，内部统计用 id / name，转一层
    final_themes: list[dict] = []
    for t in merged:
        if not isinstance(t, dict):
            continue
        tid = str(t.get("theme_id") or "").strip()
        name = str(t.get("theme_name") or "").strip()
        if not tid or not name:
            continue
        final_themes.append({
            "id": tid,
            "name": name,
            "description": str(t.get("description") or ""),
            "sentiment": str(t.get("sentiment") or "neutral"),
        })
    if not final_themes:
        raise RuntimeError("主题合并结果缺少 theme_id / theme_name 字段")

    # 给 classify 用的主题列表（id + name + description）
    themes_for_classify = _json.dumps(
        [{"theme_id": t["id"], "theme_name": t["name"], "description": t["description"]} for t in final_themes],
        ensure_ascii=False,
    )

    # ── Phase 3：并发分类各批评论 ───────────────────────────────────────
    indexed_comments = [{"idx": i, "text": text} for i, text in enumerate(comments)]
    classify_batches = comment_analysis.make_batches(indexed_comments, comment_analysis.CLASSIFY_BATCH_SIZE)
    n_classify_batch = len(classify_batches)
    yield ("progress", f"正在并发分类评论观点（共 {n_classify_batch} 批，每批 {comment_analysis.CLASSIFY_BATCH_SIZE} 条）…")
    classify_done = 0

    def _normalize_classify_item(item, fallback_idx: int | None = None) -> dict | None:
        if not isinstance(item, dict):
            return None
        try:
            idx = int(item.get("idx", fallback_idx))
        except (TypeError, ValueError):
            return None
        if idx < 0 or idx >= len(comments):
            return None
        tids = item.get("theme_ids")
        if not isinstance(tids, list) or not tids:
            single = item.get("theme_id")
            tids = [single] if single else ["other"]
        tids = [str(x) for x in tids if x]
        return {
            "idx": idx,
            "theme_ids": tids or ["other"],
            "sentiment": str(item.get("sentiment") or "neutral"),
            "is_quote_candidate": bool(item.get("is_quote_candidate")),
            "translation": str(item.get("translation") or ""),
            "original": comments[idx],
        }

    async def _classify(batch: list[dict], depth: int = 0):
        # classify 输入带 idx，输出也必须带回 idx；本地按 idx 回填。
        # 如果 Dify 返回部分结果，缺失 idx 会拆成更小批次补跑，避免整批作废。
        expected_idxs = [int(x["idx"]) for x in batch]
        expected_set = set(expected_idxs)
        last_err = ""
        for attempt in range(3):  # 首次 + 重试 2 次
            async with sem:
                raw = await workflow_run(
                    inputs=_comment_dify_inputs(
                        "classify",
                        post_title=post_title,
                        post_content=post_content,
                        comments_json=_batch_to_json(batch),
                        themes_json=themes_for_classify,
                    ),
                    api_key=key,
                    log_prefix="comment classify",
                )
            if raw == STOP_SIGNAL:
                # STOP_SIGNAL 语义是 Dify 返回 400（永久性请求错误），重试无益，
                # 且静默丢批会导致统计偏小却显示成功，故直接中止整个分析。
                raise RuntimeError("评论分类调用被 Dify 拒绝（HTTP 400），已中止分析，请检查 classify 节点配置或输入")
            arr, _e = comment_analysis.loads_loose(raw) if raw else (None, "empty")
            if isinstance(arr, list):
                by_idx: dict[int, dict] = {}
                for i, item in enumerate(arr):
                    fallback_idx = expected_idxs[i] if i < len(expected_idxs) else None
                    normalized = _normalize_classify_item(item, fallback_idx)
                    if normalized and normalized["idx"] in expected_set:
                        by_idx[normalized["idx"]] = normalized
                missing = [idx for idx in expected_idxs if idx not in by_idx]
                if not missing:
                    return [by_idx[idx] for idx in expected_idxs]
                last_err = f"返回 {len(by_idx)} 条有效 idx，缺失 {len(missing)} 条"
                print(
                    f"[comment classify] 第 {attempt + 1} 次返回不完整（{last_err}），重试中…",
                    flush=True,
                )
                continue
            got = len(arr) if isinstance(arr, list) else "非数组"
            last_err = f"返回 {got}，期望 idx 数 {len(batch)}"
            print(f"[comment classify] 第 {attempt + 1} 次返回格式不符（{last_err}），重试中…", flush=True)
        if len(batch) > 10 and depth < 3:
            mid = len(batch) // 2
            print(
                f"[comment classify] 批次仍不匹配，自动拆分为 {mid}+{len(batch) - mid} 条继续重试",
                flush=True,
            )
            left = await _classify(batch[:mid], depth + 1)
            right = await _classify(batch[mid:], depth + 1)
            return left + right
        raise RuntimeError(
            f"评论分类批次返回 idx 不完整（{last_err}），拆分重试后仍失败，已中止分析"
        )

    classify_tasks = [asyncio.create_task(_classify(b)) for b in classify_batches]
    classifications: list[dict] = []
    try:
        for t in asyncio.as_completed(classify_tasks):
            classifications.extend(await t)
            classify_done += 1
            yield ("progress", f"评论分类进度 {classify_done}/{n_classify_batch} 批")
    except BaseException:
        # 任一批失败/中止时，取消其余仍在跑的批，避免继续空打 Dify
        for task in classify_tasks:
            if not task.done():
                task.cancel()
        await asyncio.gather(*classify_tasks, return_exceptions=True)
        raise

    if not classifications:
        raise RuntimeError("评论分类未返回任何结果，请检查 Dify classify 节点配置")

    # ── Phase 4：本地统计聚合 ───────────────────────────────────────────
    yield ("progress", "正在统计占比、情感分布与代表引用…")
    stats = comment_analysis.aggregate(final_themes, classifications)

    # ── Phase 5：生成中文舆情简报 ───────────────────────────────────────
    yield ("progress", "正在生成中文舆情简报…")
    report_payload = {
        "post_title": post_title,
        "total_comments": len(classifications),
        "source_sample_count": original_sample_count,
        "off_topic_count": off_topic_count,
        "sentiment_overall": stats["sentiment_overall"],
        "themes": [
            {
                "name": t["name"],
                "description": t["description"],
                "percentage": t["percentage"],
                "count": t["count"],
                "positive_pct": t["positive_pct"],
                "negative_pct": t["negative_pct"],
                "neutral_pct": t["neutral_pct"],
                "quotes": t["quotes"],
            }
            for t in stats["themes"]
        ],
        "other_themes": stats["other_themes"],
    }
    report_md = await workflow_run(
        inputs=_comment_dify_inputs(
            "report", post_title=post_title, post_content=post_content, themes_json=_json.dumps(report_payload, ensure_ascii=False)
        ),
        api_key=key,
        log_prefix="comment report",
    )
    if report_md == STOP_SIGNAL:
        report_md = ""

    yield ("result", {
        "themes": stats["themes"],
        "other_themes": stats["other_themes"],
        "sentiment_overall": stats["sentiment_overall"],
        "total_classified": stats["total_classified"],
        "relevance_stats": relevance_stats,
        "selected_raw_comments": [],
        "report_md": (report_md or "").strip(),
    })
    return


# ============================================================
# FastAPI:app 实例 / 中间件 / 静态 / 首页·登录 均在 app/main.py。
# 过渡期 server.py 仍 import 该 app,并在下方继续注册尚未迁出的路由;
# 步骤3 迁完后本文件仅剩 `from app.main import app` 一行。
# ============================================================

from app.main import app  # noqa: E402

# ── 上传 ──────────────────────────────────────────────

@app.post("/api/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    content = await file.read()
    try:
        rows = _parse_file(file.filename or "upload.csv", content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not rows:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(rows) <= 1:
        raise HTTPException(status_code=400, detail="文件只有表头没有数据行")

    sid = new_session()
    sess = get_session(sid)
    sess["rows"] = rows
    sess["filename"] = file.filename or "upload.csv"
    _assign_session_owner(sess, await _current_login(request))
    save_session(sid, sess)

    result = {
        "session_id": sid,
        "filename": file.filename,
        "total_rows": len(rows) - 1,
        "headers": rows[0],
        "preview": rows[1: min(6, len(rows))],
    }
    await audit_log(
        request,
        "survey",
        "上传数据",
        f"文件：{file.filename or 'unknown'}；样本行数：{len(rows) - 1}",
        metadata={"session_id": sid, "rows": len(rows) - 1},
    )
    return result


# ── 评论舆情分析：上传 + 运行 ──────────────────────────────────────

@app.post("/api/comment-analysis/upload")
async def comment_analysis_upload(
    request: Request,
    file: UploadFile = File(...),
    post_title: str = Form(""),
    post_content: str = Form(""),
):
    """上传评论文件 → 保存临时文件 → 创建 session。重型预处理走 SSE。"""
    content = await file.read()
    filename = file.filename or "upload.csv"
    suffix = Path(filename).suffix.lower()
    if suffix not in {".csv", ".xlsx", ".xls"}:
        raise HTTPException(status_code=400, detail="仅支持 CSV / Excel（.csv / .xlsx / .xls）文件")
    if not content:
        raise HTTPException(status_code=400, detail="文件为空")
    post_title = (post_title or "").strip()
    post_content = (post_content or "").strip()
    if not post_title:
        raise HTTPException(status_code=400, detail="请填写帖子标题")
    if not post_content:
        raise HTTPException(status_code=400, detail="请填写帖子原文")
    file_hash = hashlib.sha256(content).hexdigest()
    login = await _current_login(request)
    duplicate_report = None
    if _load_app_settings().get("comment_duplicate_reminder_enabled", True):
        duplicate_report = _find_comment_duplicate_report(file_hash, login)

    sid = new_session()
    _COMMENT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    upload_path = _COMMENT_UPLOAD_DIR / f"{sid}{suffix}"
    with open(upload_path, "wb") as f:
        f.write(content)

    print(
        f"[comment-upload] saved sid={sid} file={filename} size={len(content)} path={upload_path}",
        flush=True,
    )

    sess = get_session(sid)
    sess["kind"] = "comment"
    sess["mode"] = "comment"  # 导出时据此插评论分析专用免责声明，不插问卷定性声明
    sess["filename"] = filename
    sess["comment_file_hash"] = file_hash
    sess["comment_upload_path"] = str(upload_path)
    sess["comment_upload_size"] = len(content)
    sess["comment_post_title"] = post_title
    sess["comment_post_content"] = post_content
    sess["comment_preprocess_done"] = False
    _assign_session_owner(sess, login)
    save_session(sid, sess)
    await audit_log(
        request,
        "comment",
        "上传评论文件",
        f"文件：{filename}；大小：{len(content)} bytes",
        metadata={"session_id": sid, "size": len(content)},
    )
    return {
        "session_id": sid,
        "filename": filename,
        "size": len(content),
        "preprocess_required": True,
        "duplicate_report": duplicate_report,
    }


@app.get("/api/comment-analysis/preprocess/{session_id}")
async def comment_analysis_preprocess(session_id: str, request: Request):
    """SSE：流式解析、清洗、抽样评论文件，并保存抽样结果到 session。"""
    sess = get_session(session_id)
    if sess.get("kind") != "comment":
        raise HTTPException(status_code=400, detail="该会话不是评论分析任务")
    upload_path = sess.get("comment_upload_path")
    filename = sess.get("filename") or "upload.csv"
    if not upload_path or not os.path.exists(upload_path):
        raise HTTPException(status_code=400, detail="评论文件不存在或已过期，请重新上传")

    async def generate():
        try:
            if sess.get("comment_preprocess_done") and sess.get("comment_sample"):
                yield sse_event({
                    "type": "comment_preprocess_done",
                    "message": "预处理已完成",
                    "scan_rows": sess.get("comment_scan_rows", 0),
                    "nonempty_count": sess.get("comment_nonempty_count", 0),
                    "valid_count": sess.get("comment_valid_count", 0),
                    "sample_count": sess.get("comment_sample_count", 0),
                    "scan_capped": sess.get("comment_scan_capped", False),
                    "warning": sess.get("comment_scan_warning", ""),
                    "sample_meta": sess.get("comment_sample_meta", {}),
                })
                return

            start = time.time()
            print(f"[comment-preprocess] start sid={session_id} file={filename}", flush=True)
            yield sse_event({"type": "progress", "message": "上传完成，开始预处理评论文件…"})
            final_result = None
            for item in comment_analysis.preprocess_comment_file(upload_path, filename):
                if item.get("kind") == "progress":
                    payload = {"type": "progress", **{k: v for k, v in item.items() if k != "kind"}}
                    yield sse_event(payload)
                    await asyncio.sleep(0)
                elif item.get("kind") == "done":
                    final_result = item["result"]

            if not final_result:
                yield sse_event({"type": "error", "message": "预处理未产出结果"})
                return

            sample = final_result.pop("sample")
            long_candidates = final_result.pop("long_candidates", [])
            sess.update({
                "comment_preprocess_done": True,
                "comment_col_name": final_result.get("comment_column", ""),
                "comment_column_name": final_result.get("comment_column", ""),
                "comment_column_index": final_result.get("comment_column_index", 0),
                "comment_column_detection": final_result.get("comment_column_detection", {}),
                "comment_scan_rows": final_result.get("scan_rows", 0),
                "comment_nonempty_count": final_result.get("nonempty_count", 0),
                "comment_valid_count": final_result.get("valid_count", 0),
                "comment_sample_count": final_result.get("sample_count", 0),
                "comment_sample_pool_count": final_result.get("sample_pool_count", final_result.get("sample_count", 0)),
                "comment_scan_capped": final_result.get("scan_capped", False),
                "comment_scan_warning": final_result.get("warning", ""),
                "comment_filter_stats": final_result.get("filter_stats", {}),
                "comment_sample": sample,
                "comment_long_candidates": long_candidates,
                "comment_long_candidate_count": final_result.get("long_candidate_count", len(long_candidates)),
                "comment_sample_meta": {
                    "scan_rows": final_result.get("scan_rows", 0),
                    "nonempty_count": final_result.get("nonempty_count", 0),
                    "valid_count": final_result.get("valid_count", 0),
                    "sample_count": final_result.get("sample_count", 0),
                    "sample_pool_count": final_result.get("sample_pool_count", final_result.get("sample_count", 0)),
                    "initial_sample_max_n": final_result.get("initial_sample_max_n", 0),
                    "sample_pool_max_n": final_result.get("sample_pool_max_n", 0),
                    "long_candidate_count": final_result.get("long_candidate_count", len(long_candidates)),
                    "long_candidate_max_n": final_result.get("long_candidate_max_n", 0),
                    "scan_capped": final_result.get("scan_capped", False),
                    "warning": final_result.get("warning", ""),
                    "column_name": final_result.get("comment_column", ""),
                    "column_index": final_result.get("comment_column_index", 0),
                    "column_detection": final_result.get("comment_column_detection", {}),
                },
            })
            save_session(session_id, sess)
            print(
                "[comment-preprocess] done "
                f"sid={session_id} scan={sess['comment_scan_rows']} valid={sess['comment_valid_count']} "
                f"sample={sess['comment_sample_count']} elapsed={time.time() - start:.2f}s",
                flush=True,
            )
            await audit_log(
                request,
                "comment",
                "预处理评论数据",
                f"文件：{filename}；扫描 {sess['comment_scan_rows']} 行；有效 {sess['comment_valid_count']} 条；抽样 {sess['comment_sample_count']} 条",
                metadata={
                    "session_id": session_id,
                    "scan_rows": sess["comment_scan_rows"],
                    "valid": sess["comment_valid_count"],
                    "sample": sess["comment_sample_count"],
                    "scan_capped": sess["comment_scan_capped"],
                },
            )
            client_payload = {k: v for k, v in final_result.items() if k not in {"preview_comments"}}
            yield sse_event({
                "type": "comment_preprocess_done",
                "message": "评论预处理完成，开始 AI 分析…",
                **client_payload,
                "sample_meta": sess["comment_sample_meta"],
            })
        except Exception as e:
            print(f"[comment-preprocess] failed sid={session_id} error={e!r}", flush=True)
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


def _comment_sample_note_md(sess: dict) -> str:
    meta = sess.get("comment_sample_meta") or {}
    if not meta:
        return ""
    warning = str(meta.get("warning") or "").strip()
    capped = bool(meta.get("scan_capped"))
    note = (
        "> 样本口径：本次评论分析基于抽样评论统计；"
        f"扫描 {meta.get('scan_rows', 0)} 行，"
        f"非空评论 {meta.get('nonempty_count', 0)} 条，"
        f"有效评论 {meta.get('valid_count', 0)} 条，"
        f"初始抽样 {meta.get('sample_count', 0)} 条"
    )
    relevance_sample = meta.get("relevance_sample_count")
    if relevance_sample is not None:
        note += f"，实际用于相关性筛选 {relevance_sample} 条"
    related = meta.get("topic_related_count")
    off_topic = meta.get("off_topic_count")
    if related is not None:
        note += f"，其中与帖子主题/正文相关 {related} 条"
        if off_topic is not None:
            note += f"，剔除无关评论 {off_topic} 条"
    note += "。"
    if capped:
        note += " 文件超过扫描上限，结果仅代表已扫描部分。"
    if warning:
        note += f" {warning}"
    return note


@app.get("/api/comment-analysis/run/{session_id}")
async def comment_analysis_run(session_id: str, request: Request):
    """SSE：执行评论舆情分析流水线，流式回传进度与最终结果。"""
    sess = get_session(session_id)
    if sess.get("kind") != "comment" or not sess.get("comment_preprocess_done") or not sess.get("comment_sample"):
        raise HTTPException(status_code=400, detail="该会话尚未完成评论预处理，请重新上传并等待预处理完成")

    async def generate():
        quote_task: asyncio.Task | None = None
        try:
            if sess.get("comment_long_candidates"):
                quote_task = asyncio.create_task(_select_comment_raw_quotes(sess))
                yield sse_event({"type": "progress", "message": "已启动玩家评论原文精选，舆情报告将先生成…"})
            result = None
            async for item in _comment_analysis_pipeline(sess):
                if item[0] == "progress":
                    yield sse_event({"type": "progress", "message": item[1]})
                elif item[0] == "result":
                    result = item[1]
            if result is None:
                yield sse_event({"type": "error", "message": "分析未产出结果"})
                return
            sample_note = _comment_sample_note_md(sess)
            if sample_note:
                result["report_md"] = sample_note + "\n\n" + (result.get("report_md") or "*（未生成简报）*").strip()
            report_title = _comment_report_title(sess)
            sess["comment_report_title"] = report_title
            result["title"] = report_title
            result["post_title"] = sess.get("comment_post_title", "")
            sess["comment_result"] = result
            # 同步写入 report_md，复用现有 PDF / 飞书导出端点
            if result.get("report_md"):
                sess["report_md"] = result["report_md"]
            save_session(session_id, sess)
            save_to_history(session_id, sess)
            yield sse_event({"type": "comment_done", **result, "sample_meta": sess.get("comment_sample_meta", {})})
            if quote_task:
                try:
                    selected_raw_comments = await quote_task
                    current_report = sess.get("report_md") or result.get("report_md") or ""
                    if selected_raw_comments:
                        full_report = _comment_append_selected_raw_comments(current_report, selected_raw_comments)
                    else:
                        full_report = current_report
                    sess["comment_selected_raw_comments"] = selected_raw_comments
                    sess["report_md"] = full_report
                    result["selected_raw_comments"] = selected_raw_comments
                    result["report_md"] = full_report
                    result["title"] = sess.get("comment_report_title") or _comment_report_title(sess)
                    result["post_title"] = sess.get("comment_post_title", "")
                    sess["comment_result"] = result
                    save_session(session_id, sess)
                    save_to_history(session_id, sess)
                    if selected_raw_comments:
                        yield sse_event({"type": "progress", "message": f"玩家评论原文精选已生成：{len(selected_raw_comments)} 条"})
                    yield sse_event({
                        "type": "comment_quotes_done",
                        "selected_raw_comments": selected_raw_comments,
                        "report_md": full_report,
                    })
                except Exception as exc:  # noqa: BLE001 - 精选失败不影响已生成主报告
                    print(f"[comment quote_select] skipped due to error: {exc!r}", flush=True)
                    yield sse_event({
                        "type": "comment_quotes_error",
                        "message": "玩家评论原文精选生成失败，舆情报告已完成",
                    })
            else:
                yield sse_event({
                    "type": "comment_quotes_done",
                    "selected_raw_comments": [],
                    "report_md": result.get("report_md") or "",
                })
        except Exception as e:
            if quote_task and not quote_task.done():
                quote_task.cancel()
                await asyncio.gather(quote_task, return_exceptions=True)
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── 跑数表模式上传（问卷 + 回答数据 + 跑数统计表）──────────────────────

_Q_TITLE_RE = re.compile(r"^Q(\d+)\[")
_HTML_TAG_RE = re.compile(r"<[^>]+>")


def _questionnaire_title_map(q_rows: list[list]) -> dict[int, str]:
    """从问卷行里抽 Q号 → 题目文本(用于给清数列起可读名)。"""
    m: dict[int, str] = {}
    for r in q_rows:
        if not r:
            continue
        c0 = str(r[0]).strip() if len(r) > 0 else ""
        mt = _Q_TITLE_RE.match(c0)
        if mt:
            text = str(r[1]).strip() if len(r) > 1 else ""
            text = _HTML_TAG_RE.sub("", text).strip()
            if text:
                m[int(mt.group(1))] = text
    return m


def _build_crosstab_columns(headers: list[str], q_title_map: dict[int, str]) -> list[dict]:
    """跑数表模式：确定性构建列元数据（无需 AI 题型识别）。

    只关心三类：open_text(供聚类)、id/profile(原文署名)，其余 ignore
    （数字来自跑数表，清数闭合列不参与统计）。
    """
    cols: list[dict] = []
    for i, h in enumerate(headers):
        hs = str(h).strip()
        low = hs.lower()
        role = "ignore"
        name = hs
        if hs.endswith("__open"):
            role = "open_text"
            mt = re.match(r"Q(\d+)", hs)
            if mt and int(mt.group(1)) in q_title_map:
                name = q_title_map[int(mt.group(1))]
        elif "zone_id" in low or "zoneid" in low:
            role = "ignore"
        elif "role_id" in low or "roleid" in low:
            role = "id"
        elif low in ("response id", "responseid"):
            role = "id"
        elif hs in ("段位", "等级", "性别", "年龄", "区服", "国家", "地区", "服务器"):
            role = "profile_dim"
        cols.append({
            "index": i,
            "name": name,
            "role": role,
            "source": "crosstab",
            "column_indexes": [i],
        })
    return cols


@app.post("/api/upload/crosstab")
async def upload_crosstab(
    request: Request,
    survey_file: UploadFile = File(...),    # 问卷（题目意图上下文）
    data_file: UploadFile = File(...),      # 清数：清洗后的问卷回答（用于主观题原文）
    crosstab_file: UploadFile = File(...),  # 跑数表：倍市得交叉统计表（数字来源）
):
    """倍市得「跑数表模式」上传：平台不再自算统计，数字直接取自跑数表，
    只对主观题做聚类。三文件均必需。"""
    # 1) 清数（回答数据）→ rows
    data_content = await data_file.read()
    try:
        rows = _parse_file(data_file.filename or "data.xlsx", data_content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=f"回答数据解析失败：{e}")
    if not rows or len(rows) <= 1:
        raise HTTPException(status_code=400, detail="回答数据为空或只有表头")

    # 2) 跑数表 → 结构化 → stats markdown
    ct_content = await crosstab_file.read()
    try:
        parsed = crosstab_parser.parse(ct_content)
        crosstab_md = crosstab_parser.render_to_markdown(parsed)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"跑数表解析失败：{e}")

    # 3) 问卷 → 纯文本（题目意图上下文）+ Q号→题名映射（给清数列起可读名）
    survey_content = await survey_file.read()
    try:
        q_rows = _parse_file(survey_file.filename or "survey.xlsx", survey_content)
        questionnaire_text = "\n".join(
            " | ".join(str(c) for c in r if str(c).strip())
            for r in q_rows if any(str(c).strip() for c in r)
        )
        q_title_map = _questionnaire_title_map(q_rows)
    except Exception:
        questionnaire_text = ""
        q_title_map = {}

    # 4) 确定性建列（跳过 AI 题型识别）+ 跑数表题目清单（给 planner）
    columns = _build_crosstab_columns(rows[0], q_title_map)
    crosstab_questions = crosstab_parser.question_names(parsed)

    sid = new_session()
    sess = get_session(sid)
    sess["rows"] = rows
    sess["filename"] = data_file.filename or "data.xlsx"
    sess["mode"] = "crosstab"
    sess["crosstab_md"] = crosstab_md
    sess["questionnaire_text"] = questionnaire_text
    sess["crosstab_questions"] = crosstab_questions
    # 列已确定性建好，直接作为"已确认"，跳过题型确认步骤
    sess["columns_detected"] = columns
    sess["confirmed_columns"] = columns
    _assign_session_owner(sess, await _current_login(request))
    save_session(sid, sess)

    result = {
        "session_id": sid,
        "filename": data_file.filename,
        "total_rows": len(rows) - 1,
        "headers": rows[0],
        "preview": rows[1: min(6, len(rows))],
        "mode": "crosstab",
        "crosstab_questions": len(parsed.get("questions", [])),
        "crosstab_segments": [s["label"] for s in parsed.get("segments", [])],
    }
    await audit_log(
        request,
        "survey",
        "上传跑数表数据",
        f"问卷：{survey_file.filename or '?'}；数据：{data_file.filename or '?'}；"
        f"跑数表：{crosstab_file.filename or '?'}；样本行数：{len(rows) - 1}",
        metadata={"session_id": sid, "rows": len(rows) - 1, "mode": "crosstab"},
    )
    return result


# ── 题型识别（Step 2，LLM，SSE）──────────────────────

@app.get("/api/columns/{session_id}")
async def get_columns(session_id: str, request: Request):
    """LLM 识别列题型（含 Google Form 矩阵题分组、中文题名、多选选项清单）。

    流式返回；最终发 `columns_ready`，columns 为「逻辑题」列表（矩阵题跨多列）。
    LLM 解析失败时回退本地启发式。
    """
    sess = get_session(session_id)
    rows = sess.get("rows")
    if not rows:
        raise HTTPException(status_code=400, detail="会话中没有数据")
    if not DIFY_COLUMN_KEY:
        raise HTTPException(status_code=500, detail="未配置 DIFY_COLUMN_KEY（题型识别应用）")

    async def generate():
        try:
            groups = _group_googleform_matrix(rows[0])
            query = _build_column_detect_query(rows, groups)
            header_count = len(rows[0])

            answer_chunks: list[str] = []
            final_conv = ""
            async for chunk, conv_id in sse_dify_stream(query, session_id, "", DIFY_COLUMN_KEY):
                if chunk:
                    answer_chunks.append(chunk)
                    yield sse_event({"type": "chunk", "content": chunk})
                if conv_id:
                    final_conv = conv_id

            questions, err = survey_plan.parse_columns_from_llm("".join(answer_chunks), header_count)

            if not questions:
                retry_q = (
                    f"上次输出无法解析: {err}。请严格按 schema 用 ```json``` 围栏重新输出，"
                    "不要附加任何解释文字。"
                )
                retry_chunks: list[str] = []
                async for chunk, conv_id in sse_dify_stream(retry_q, session_id, final_conv, DIFY_COLUMN_KEY):
                    if chunk:
                        retry_chunks.append(chunk)
                questions, err = survey_plan.parse_columns_from_llm("".join(retry_chunks), header_count)

            if not questions:
                print(f"[columns] LLM 解析失败，回退本地启发式：{err}")
                questions = _heuristic_questions(rows, groups)
                yield sse_event({"type": "chunk", "content": "\n（题型识别解析失败，已回退本地推断，请仔细核对）\n"})

            questions = _enrich_questions(questions, rows[0], groups)
            questions = _sanitize_choice_options(rows, questions)
            # 跑数表模式安全网：倍市得清数中以 __open 结尾的列必是开放题，
            # 强制 role=open_text，避免 AI 误判导致主观题漏聚类。
            if sess.get("mode") == "crosstab":
                hdrs = rows[0]
                for q in questions:
                    idx = q.get("index")
                    if isinstance(idx, int) and 0 <= idx < len(hdrs) \
                            and str(hdrs[idx]).strip().endswith("__open"):
                        q["role"] = "open_text"
            sess["columns_detected"] = questions
            save_session(session_id, sess)
            await audit_log(
                request,
                "survey",
                "识别题型",
                f"会话：{session_id}；识别列数：{len(questions)}",
                metadata={"session_id": session_id, "columns": len(questions)},
            )
            yield sse_event({"type": "columns_ready", "columns": questions})
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/columns/{session_id}/confirm")
async def confirm_columns(session_id: str, req: ColumnConfirmRequest, request: Request):
    """存储用户确认（或修改）后的列题型。"""
    sess = get_session(session_id)
    sess["confirmed_columns"] = req.columns
    save_session(session_id, sess)
    await audit_log(
        request,
        "survey",
        "确认数据列",
        f"会话：{session_id}；确认列数：{len(req.columns)}",
        metadata={"session_id": session_id, "columns": len(req.columns)},
    )
    return {"ok": True}


# ── Plan（SSE）────────────────────────────────────────

@app.get("/api/plan/{session_id}")
async def get_plan(session_id: str, request: Request):
    sess = get_session(session_id)
    rows = sess.get("rows")
    confirmed_columns = sess.get("confirmed_columns")

    is_crosstab = sess.get("mode") == "crosstab"

    if not rows:
        raise HTTPException(status_code=400, detail="会话中没有数据，请先上传文件")
    if is_crosstab:
        if not DIFY_CROSSTAB_PLANNER_KEY:
            raise HTTPException(status_code=500, detail="未配置 DIFY_CROSSTAB_PLANNER_KEY")
    elif not DIFY_PLANNER_KEY:
        raise HTTPException(status_code=500, detail="未配置 DIFY_PLANNER_KEY")

    async def generate():
        try:
            if is_crosstab:
                # ── 跑数表模式：AI 读问卷原文规划章节大纲 ──────────────────
                cols = confirmed_columns or []
                open_names = [c["name"] for c in cols if c.get("role") == "open_text"]
                avail = sess.get("crosstab_questions", [])
                query = _build_crosstab_planner_query(
                    sess.get("questionnaire_text", ""), avail, open_names
                )
                ans_chunks: list[str] = []
                conv = ""
                async for chunk, cid in sse_dify_stream(query, session_id, "", DIFY_CROSSTAB_PLANNER_KEY):
                    if chunk:
                        ans_chunks.append(chunk)
                        yield sse_event({"type": "chunk", "content": chunk})
                    if cid:
                        conv = cid
                ctp, err = survey_plan.parse_crosstab_plan("".join(ans_chunks))
                if not ctp:
                    yield sse_event({"type": "progress", "message": "方案格式校验中，正在修订输出…"})
                    retry_q = (
                        f"上次输出无法解析: {err}。请只输出一个 JSON 对象"
                        "(含 parts 和 open_questions)，用 ```json``` 围栏，不要解释文字。"
                    )
                    retry_chunks: list[str] = []
                    async for chunk, cid in sse_dify_stream(retry_q, session_id, conv, DIFY_CROSSTAB_PLANNER_KEY):
                        if chunk:
                            retry_chunks.append(chunk)
                        if cid:
                            conv = cid
                    ctp, err = survey_plan.parse_crosstab_plan("".join(retry_chunks))
                if not ctp:
                    yield sse_event({"type": "error", "message": f"章节大纲解析失败：{err}"}); return

                plan = {
                    "mode": "crosstab",
                    "columns": cols,
                    "parts": ctp["parts"],
                    "open_questions": ctp["open_questions"],
                    "cross_tabs": [],
                }
                sess["plan"] = plan
                sess["planner_conv_id"] = conv
                save_session(session_id, sess)
                card_text = _render_crosstab_plan_card(plan)
                await audit_log(
                    request, "survey", "生成章节大纲",
                    f"会话：{session_id}；章节数：{len(plan['parts'])}",
                    metadata={"session_id": session_id, "parts": len(plan["parts"]), "mode": "crosstab"},
                )
                yield sse_event({"type": "plan_ready", "plan": plan, "card_text": card_text, "headers": rows[0]})
                return

            if confirmed_columns:
                planner_query = _build_planner_query_with_confirmed(rows, confirmed_columns)
            else:
                planner_query = (
                    _build_planner_sample(rows)
                    + "\n\n" + _get_planner_extra()
                )

            answer_chunks: list[str] = []
            final_conv_id = ""

            async for chunk, conv_id in sse_dify_stream(planner_query, session_id, "", DIFY_PLANNER_KEY):
                if chunk:
                    answer_chunks.append(chunk)
                    yield sse_event({"type": "chunk", "content": chunk})
                if conv_id:
                    final_conv_id = conv_id

            full_answer = "".join(answer_chunks)
            headers = rows[0]
            plan, err = survey_plan.parse_plan_from_llm(full_answer, len(headers))

            if not plan:
                yield sse_event({"type": "progress", "message": "方案格式校验中，正在修订输出…"})
                retry_q = (
                    f"上次输出无法解析: {err}。请严格按 JSON schema 重新输出，"
                    "用 ```json ``` 围栏，不要附加解释文字。"
                )
                retry_chunks: list[str] = []
                async for chunk, conv_id in sse_dify_stream(retry_q, session_id, final_conv_id, DIFY_PLANNER_KEY):
                    if chunk: retry_chunks.append(chunk)
                    if conv_id: final_conv_id = conv_id
                plan, err = survey_plan.parse_plan_from_llm("".join(retry_chunks), len(headers))

            if not plan:
                yield sse_event({"type": "error", "message": f"方案解析失败：{err}"}); return

            # 用用户确认的题型覆盖 planner 的列分类（权威），并修补 parts（矩阵题整体入同一 part）
            if confirmed_columns:
                plan = survey_plan.merge_confirmed_into_plan(plan, confirmed_columns)

            sess["plan"] = plan
            sess["planner_conv_id"] = final_conv_id
            save_session(session_id, sess)
            card_text = survey_plan.render_plan_for_user(plan, headers)
            await audit_log(
                request,
                "survey",
                "生成分析方案",
                f"会话：{session_id}；Part 数：{len(plan.get('parts', []))}",
                metadata={"session_id": session_id, "parts": len(plan.get("parts", []))},
            )
            yield sse_event({"type": "plan_ready", "plan": plan, "card_text": card_text, "headers": headers})

        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── Plan 确认/修改 ──────────────────────────────────

@app.post("/api/plan/confirm")
async def confirm_plan(req: PlanConfirmRequest, request: Request):
    sess = get_session(req.session_id)
    plan = sess.get("plan")
    rows = sess.get("rows")
    planner_conv_id = sess.get("planner_conv_id", "")

    if not plan or not rows:
        raise HTTPException(status_code=400, detail="会话状态丢失，请重新上传文件")

    if survey_plan.is_user_approval(req.user_text):
        await audit_log(
            request,
            "survey",
            "确认分析方案",
            f"会话：{req.session_id}",
            metadata={"session_id": req.session_id},
        )
        return JSONResponse({"approved": True})

    async def generate():
        try:
            # ── 跑数表模式：章节大纲修订 ──────────────────────────────
            if sess.get("mode") == "crosstab":
                conv = planner_conv_id
                rev_q = _build_crosstab_plan_revision_query(
                    sess.get("questionnaire_text", ""), plan.get("parts", []), req.user_text
                )
                rchunks: list[str] = []
                async for chunk, cid in sse_dify_stream(rev_q, req.session_id, "", DIFY_CROSSTAB_PLANNER_KEY):
                    if chunk:
                        rchunks.append(chunk)
                        yield sse_event({"type": "chunk", "content": chunk})
                    if cid:
                        conv = cid
                ctp, err = survey_plan.parse_crosstab_plan("".join(rchunks))
                if not ctp:
                    yield sse_event({"type": "error", "message": f"修订章节大纲解析失败：{err}"}); return
                new_plan = dict(plan)
                new_plan["parts"] = ctp["parts"]
                new_plan["open_questions"] = ctp["open_questions"]
                sess["plan"] = new_plan
                sess["planner_conv_id"] = conv
                save_session(req.session_id, sess)
                card_text = _render_crosstab_plan_card(new_plan)
                await audit_log(
                    request, "survey", "修订章节大纲",
                    f"会话：{req.session_id}；修改意见：{_short_text(req.user_text)}",
                    metadata={"session_id": req.session_id, "mode": "crosstab"},
                )
                yield sse_event({"type": "plan_ready", "plan": new_plan, "card_text": card_text, "headers": rows[0]})
                return

            new_conv_id = planner_conv_id
            answer_chunks: list[str] = []
            headers = rows[0]
            revision_query = _build_plan_revision_query(
                plan,
                headers,
                sess.get("confirmed_columns", []),
                req.user_text,
            )
            async for chunk, conv_id in sse_dify_stream(
                revision_query, req.session_id, "", DIFY_PLANNER_KEY
            ):
                if chunk:
                    answer_chunks.append(chunk)
                    yield sse_event({"type": "chunk", "content": chunk})
                if conv_id:
                    new_conv_id = conv_id

            full_answer = "".join(answer_chunks)
            new_plan, err = survey_plan.parse_plan_from_llm(
                full_answer,
                len(headers)
            )
            if not new_plan:
                retry_query = (
                    f"{revision_query}\n\n"
                    "上一次输出无法解析为 JSON。请修正并只返回完整 JSON 对象。\n"
                    f"解析错误：{err}\n"
                    f"<previous_output>\n{full_answer[:4000]}\n</previous_output>"
                )
                retry_chunks: list[str] = []
                retry_conv_id = new_conv_id
                yield sse_event({"type": "progress", "message": "方案格式校验中，正在修订输出…"})
                yield sse_event({"type": "chunk", "content": "\n\n正在按严格 JSON 格式重新修订方案...\n"})
                async for chunk, conv_id in sse_dify_stream(
                    retry_query, req.session_id, "", DIFY_PLANNER_KEY
                ):
                    if chunk:
                        retry_chunks.append(chunk)
                        yield sse_event({"type": "chunk", "content": chunk})
                    if conv_id:
                        retry_conv_id = conv_id
                new_plan, err = survey_plan.parse_plan_from_llm("".join(retry_chunks), len(headers))
                if new_plan:
                    new_conv_id = retry_conv_id
            if not new_plan:
                yield sse_event({"type": "error", "message": f"修订方案解析失败：{err}"}); return

            if sess.get("confirmed_columns"):
                new_plan = survey_plan.merge_confirmed_into_plan(new_plan, sess["confirmed_columns"])

            sess["plan"] = new_plan
            sess["planner_conv_id"] = new_conv_id
            save_session(req.session_id, sess)
            card_text = survey_plan.render_plan_for_user(new_plan, headers)
            await audit_log(
                request,
                "survey",
                "修订分析方案",
                f"会话：{req.session_id}；修改意见：{_short_text(req.user_text)}",
                metadata={"session_id": req.session_id},
            )
            yield sse_event({"type": "plan_ready", "plan": new_plan, "card_text": card_text, "headers": headers})
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── 统计 ────────────────────────────────────────────

@app.post("/api/stats/{session_id}")
async def compute_stats(session_id: str, request: Request):
    sess = get_session(session_id)
    plan = sess.get("plan")
    rows = sess.get("rows")
    if not plan or not rows:
        raise HTTPException(status_code=400, detail="会话状态丢失")
    loop = asyncio.get_event_loop()
    if sess.get("mode") == "crosstab":
        # 跑数表模式：数字直接取自跑数表，平台不自算统计；只收集开放题原文供聚类。
        stats_md = sess.get("crosstab_md", "")
        open_text = await loop.run_in_executor(
            None, survey_stats.collect_open_text, rows, plan
        )
    else:
        stats_md, open_text = await loop.run_in_executor(
            None, survey_stats.compute, rows, plan
        )
    sess["stats_md"] = stats_md
    sess["open_text"] = open_text
    sess["rows_fed"] = False
    save_session(session_id, sess)
    await audit_log(
        request,
        "survey",
        "计算统计",
        f"会话：{session_id}；样本行数：{max(0, len(rows) - 1)}",
        metadata={"session_id": session_id, "rows": max(0, len(rows) - 1)},
    )
    return {"stats_md": stats_md}


# ── 报告（SSE）──────────────────────────────────────

@app.get("/api/report/{session_id}")
async def generate_report(session_id: str, request: Request):
    sess = get_session(session_id)
    _assign_session_owner(sess, await _current_login(request))
    plan = sess.get("plan")
    rows = sess.get("rows")
    stats_md = sess.get("stats_md")
    open_text = sess.get("open_text", {})

    if not all([plan, rows, stats_md]):
        raise HTTPException(status_code=400, detail="请先完成统计计算")

    is_crosstab = sess.get("mode") == "crosstab"
    # 跑数表模式恒定走聚类；定性模式以单列最大回答数判断（避免多开放题列加总误触发）
    use_large_mode = is_crosstab or any(len(v) > LARGE_SAMPLE_THRESHOLD for v in open_text.values())

    if use_large_mode:
        if not DIFY_LARGE_ANALYST_KEY:
            raise HTTPException(status_code=500, detail="未配置 DIFY_LARGE_ANALYST_KEY")
    else:
        if not DIFY_ANALYST_KEY:
            raise HTTPException(status_code=500, detail="未配置 DIFY_ANALYST_KEY")

    async def generate():
        try:
            if use_large_mode:
                # ── 大样本 / 跑数表模式：主观题四阶段批处理聚类 ──────────────
                total_open_text = sum(len(v) for v in open_text.values())
                start_msg = (
                    f"跑数表模式：数字取自跑数表，开始对 {total_open_text} 条主观题回复做聚类"
                    if is_crosstab
                    else "检测到超过500条回复，启用批量处理模式"
                )
                yield sse_event({"type": "progress", "message": start_msg})
                clustered_themes: dict = {}
                cluster_diagnostics: dict = {}
                async for item in _batch_qualitative_analysis(open_text, plan, rows[0], session_id):
                    if item[0] == "progress":
                        yield sse_event({"type": "progress", "message": item[1]})
                    elif item[0] == "diagnostics":
                        cluster_diagnostics = item[1]
                    elif item[0] == "result":
                        clustered_themes = item[1]

                failed_cols = [
                    d.get("col_name", f"列{k}") for k, d in (cluster_diagnostics or {}).items()
                    if d.get("status") != "ok"
                ]
                sess["open_text_cluster_diagnostics"] = cluster_diagnostics
                save_session(session_id, sess)
                if failed_cols:
                    msg = "部分主观题聚类未完成，报告将使用原文兜底：" + "、".join(failed_cols[:4])
                    if len(failed_cols) > 4:
                        msg += f"等 {len(failed_cols)} 列"
                    yield sse_event({"type": "progress", "message": msg})

                yield sse_event({"type": "progress", "message": "主题分析完成，开始生成报告..."})
                writer_query = _build_large_sample_writer_query(stats_md, clustered_themes, plan, rows[0], open_text)
                # 跑数表模式：把问卷原文作为题目意图上下文一并喂给 Writer
                if is_crosstab:
                    q_text = (sess.get("questionnaire_text") or "").strip()
                    if q_text:
                        if len(q_text) > 8000:
                            q_text = q_text[:8000] + "\n…（问卷过长，已截断）"
                        writer_query = (
                            f"<questionnaire>\n以下是问卷原文（仅供理解题目意图与背景，"
                            f"不要直接搬运）：\n{q_text}\n</questionnaire>\n\n" + writer_query
                        )
                analyst_key = DIFY_LARGE_ANALYST_KEY

                answer_chunks: list[str] = []
                final_conv_id = ""
                async for chunk, conv_id in sse_dify_stream(writer_query, session_id, "", analyst_key):
                    if chunk:
                        answer_chunks.append(chunk)
                        yield sse_event({"type": "chunk", "content": chunk})
                    if conv_id:
                        final_conv_id = conv_id
                full_report = "".join(answer_chunks)
            else:
                # ── 标准模式：同会话分章多轮生成（规避 Dify 单次 10 分钟超时）──
                # 第 1 轮喂全部原文+要求、只写标题并建立会话；之后复用 conversation_id，
                # 每轮只写一个 Part / Bug 模块 / 核心结论，原文不再重发。每轮输出短，单次不超时。
                analyst_key = DIFY_ANALYST_KEY
                parts_meta = _writer_parts_meta(plan, rows[0])
                final_conv_id = ""

                async def _round(query: str, conv_id: str):
                    """跑一轮：流式 yield chunk 事件，结束后把整段文本与 conv_id 存入 _round.out。"""
                    buf: list[str] = []
                    cid = conv_id
                    async for ch, c in sse_dify_stream(query, session_id, conv_id, analyst_key):
                        if ch:
                            buf.append(ch)
                            yield sse_event({"type": "chunk", "content": ch})
                        if c:
                            cid = c
                    _round.out = ("".join(buf), cid)

                total_rounds = len(parts_meta) + 3  # 标题 + N 个 Part + Bug + 核心结论

                # 第 1 轮：全部上下文 + 只写标题
                yield sse_event({"type": "progress",
                                 "message": f"分章生成 1/{total_rounds}：准备数据并生成标题…"})
                first_q = _build_writer_first_query(stats_md, open_text, plan, rows[0])
                async for ev in _round(first_q, ""):
                    yield ev
                title_text, final_conv_id = _round.out
                # 只保留标题段（首个 `## ` 之前），防止模型抢跑后续章节
                title_lines = []
                for ln in title_text.split("\n"):
                    if ln.lstrip().startswith("## "):
                        break
                    title_lines.append(ln)
                title_block = "\n".join(title_lines).strip() or title_text.strip()

                # 第 2..N+1 轮：逐 Part
                part_sections: list[str] = []
                for m in parts_meta:
                    rnd = m["i"] + 1
                    yield sse_event({"type": "progress",
                                     "message": f"分章生成 {rnd}/{total_rounds}：Part {m['i']} {m['name']}…"})
                    yield sse_event({"type": "chunk", "content": "\n\n"})
                    async for ev in _round(_build_writer_part_query(m), final_conv_id):
                        yield ev
                    sec, final_conv_id = _round.out
                    part_sections.append(sec.strip())

                # 倒数第 2 轮：Bug 模块（需要才写，否则模型回 NONE）
                yield sse_event({"type": "progress",
                                 "message": f"分章生成 {total_rounds - 1}/{total_rounds}：核查待确认问题…"})
                async for ev in _round(_build_writer_bug_query(), final_conv_id):
                    yield ev
                bug_text, final_conv_id = _round.out
                bug_clean = bug_text.strip()
                has_bug = bool(bug_clean) and bug_clean.upper().strip(" .。`*") != "NONE" and "## Bug" in bug_clean
                bug_section = bug_clean if has_bug else ""

                # 最后一轮：核心结论（汇总全部章节，放报告顶部）
                yield sse_event({"type": "progress",
                                 "message": f"分章生成 {total_rounds}/{total_rounds}：汇总核心结论…"})
                yield sse_event({"type": "chunk", "content": "\n\n"})
                async for ev in _round(_build_writer_core_query(parts_meta, has_bug), final_conv_id):
                    yield ev
                core_text, final_conv_id = _round.out
                core_block = core_text.strip()

                # 组装最终报告：标题 → 核心结论 → 各 Part → Bug 模块
                assembled = [title_block, core_block, *part_sections]
                if bug_section:
                    assembled.append(bug_section)
                full_report = "\n\n".join(b for b in assembled if b)
            drifted = survey_stats.find_numbers_not_in_stats(full_report, stats_md)
            if drifted:
                print(f"[stats] WARN drifted numbers: {drifted[:20]}")

            full_report = _inject_disclaimer(full_report, mode=sess.get("mode") or "")
            sess["report_md"] = full_report
            sess["analyst_conv_id"] = final_conv_id
            sess["analyst_app"] = "large" if use_large_mode else "standard"
            sess["rows_fed"] = False

            save_session(session_id, sess)
            save_to_history(session_id, sess)
            await audit_log(
                request,
                "survey",
                "生成报告",
                f"会话：{session_id}；文件：{sess.get('filename', 'unknown')}；模式：{'大样本' if use_large_mode else '标准'}",
                metadata={"session_id": session_id, "filename": sess.get("filename", "unknown"), "large_mode": use_large_mode},
            )

            yield sse_event({"type": "report_done", "report_md": full_report})
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── QA（SSE）────────────────────────────────────────

def _analyst_key_for_report(obj: dict) -> tuple[str, str]:
    """Return the Dify key that owns this report conversation."""
    analyst_app = obj.get("analyst_app") or ""
    mode = obj.get("mode") or obj.get("plan", {}).get("mode", "")
    if analyst_app == "large" or mode == "crosstab":
        return DIFY_LARGE_ANALYST_KEY, "DIFY_LARGE_ANALYST_KEY"
    return DIFY_ANALYST_KEY, "DIFY_ANALYST_KEY"


@app.post("/api/qa")
async def qa(req: QARequest, request: Request):
    sess = get_session(req.session_id)
    _assign_session_owner(sess, await _current_login(request))
    analyst_conv_id = sess.get("analyst_conv_id", "")
    rows = sess.get("rows", [])
    plan = sess.get("plan", {})
    rows_fed = sess.get("rows_fed", False)

    if not analyst_conv_id:
        raise HTTPException(status_code=400, detail="请先生成报告")
    analyst_key, analyst_key_name = _analyst_key_for_report(sess)
    if not analyst_key:
        raise HTTPException(status_code=500, detail=f"未配置 {analyst_key_name}")

    async def generate():
        try:
            if not rows_fed and rows:
                rows_block = _format_rows_for_qa(rows, plan)
                qa_query = f"<rows>\n{rows_block}\n</rows>\n\n用户问题: {req.question}"
            else:
                qa_query = req.question

            answer_chunks: list[str] = []
            new_conv_id = analyst_conv_id

            async for chunk, conv_id in sse_dify_stream(
                qa_query, req.session_id, analyst_conv_id, analyst_key
            ):
                if chunk:
                    answer_chunks.append(chunk)
                    yield sse_event({"type": "chunk", "content": chunk})
                if conv_id:
                    new_conv_id = conv_id

            answer_text = "".join(answer_chunks)
            sess["analyst_conv_id"] = new_conv_id or analyst_conv_id
            sess["analyst_app"] = "large" if analyst_key_name == "DIFY_LARGE_ANALYST_KEY" else "standard"
            sess["rows_fed"] = True
            sess.setdefault("qa_messages", []).extend([
                {"role": "user", "content": req.question, "ts": datetime.now().isoformat()},
                {"role": "ai", "content": answer_text, "ts": datetime.now().isoformat()},
            ])
            save_session(req.session_id, sess)
            save_to_history(req.session_id, sess)
            await audit_log(
                request,
                "report",
                "追问当前报告",
                f"会话：{req.session_id}；问题：{_short_text(req.question)}",
                metadata={"session_id": req.session_id},
            )
            yield sse_event({"type": "qa_done", "answer": answer_text})
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── 历史 QA（从历史条目恢复会话）────────────────────

@app.post("/api/history-qa")
async def history_qa(req: HistoryQARequest, request: Request):
    """从历史记录中续聊 QA（无行数据，直接使用 analyst conv_id）。"""
    login = await _current_login(request)
    history = _load_history()
    entry = _find_history_for_login(history, req.history_id, login)
    if not entry:
        raise HTTPException(status_code=404, detail="历史记录不存在")

    analyst_conv_id = entry.get("analyst_conv_id", "")
    if not analyst_conv_id:
        raise HTTPException(status_code=400, detail="该历史记录没有可续聊的对话")
    analyst_key, analyst_key_name = _analyst_key_for_report(entry)
    if not analyst_key:
        raise HTTPException(status_code=500, detail=f"未配置 {analyst_key_name}")

    async def generate():
        try:
            answer_chunks: list[str] = []
            new_conv_id = analyst_conv_id
            async for chunk, conv_id in sse_dify_stream(
                req.question, req.history_id, analyst_conv_id, analyst_key
            ):
                if chunk:
                    answer_chunks.append(chunk)
                    yield sse_event({"type": "chunk", "content": chunk})
                if conv_id:
                    new_conv_id = conv_id

            # 更新历史中的 conv_id
            answer_text = "".join(answer_chunks)
            for h in history:
                if h["id"] == req.history_id:
                    h["analyst_conv_id"] = new_conv_id or analyst_conv_id
                    h["analyst_app"] = "large" if analyst_key_name == "DIFY_LARGE_ANALYST_KEY" else "standard"
                    h.setdefault("qa_messages", []).extend([
                        {"role": "user", "content": req.question, "ts": datetime.now().isoformat()},
                        {"role": "ai", "content": answer_text, "ts": datetime.now().isoformat()},
                    ])
                    break
            _save_history(history)
            await audit_log(
                request,
                "report",
                "追问历史报告",
                f"历史报告：{entry.get('title', req.history_id)}；问题：{_short_text(req.question)}",
                metadata={"history_id": req.history_id},
            )

            yield sse_event({"type": "qa_done", "answer": answer_text})
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── 导出 ────────────────────────────────────────────

def _make_download_response(data: bytes, mime: str, filename: str) -> StreamingResponse:
    """构建带 RFC 5987 编码文件名的下载响应。"""
    encoded = quote(filename)
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
        "Content-Length": str(len(data)),
    }
    return StreamingResponse(io.BytesIO(data), media_type=mime, headers=headers)


PDF_PAGE_RULE = "@page { size: 10in 14in; margin: 0; }"
PDF_CSS = PDF_PAGE_RULE + """
* { box-sizing: border-box; }
html {
  margin: 0;
  background: #fff;
}
body {
  font-family: "Noto Sans CJK SC", "Noto Sans CJK", "WenQuanYi Micro Hei", "Microsoft YaHei", "PingFang SC", "Source Han Sans SC", Arial, sans-serif;
  font-size: 13px;
  line-height: 1.75;
  color: #222;
  background: #fff;
  width: auto;
  max-width: 900px;
  margin: 0 auto;
  padding: 36px 44px 48px;
}
h1 {
  font-size: 22px;
  font-weight: 700;
  color: #1a1a1a;
  border-bottom: 2px solid #7c3aed;
  padding-bottom: 8px;
  margin: 0 0 18px;
}
h2 {
  font-size: 17px;
  font-weight: 600;
  color: #2d2d2d;
  margin: 24px 0 10px;
  border-left: 4px solid #7c3aed;
  padding-left: 10px;
  page-break-after: avoid;
}
h3 { font-size: 14.5px; color: #444; margin: 18px 0 8px; page-break-after: avoid; }
h4 { font-size: 13.5px; color: #555; margin: 14px 0 6px; page-break-after: avoid; }
p { margin: 0 0 9px; }
ul, ol { margin: 6px 0 12px 22px; padding: 0; }
li { margin-bottom: 4px; }
blockquote {
  border-left: 3px solid #999;
  padding: 6px 14px;
  margin: 10px 0;
  color: #666;
  font-style: italic;
  background: #f9f9f9;
}
table {
  border-collapse: collapse;
  width: 100%;
  max-width: 100%;
  margin: 14px 0;
  font-size: 11.5px;
  line-height: 1.55;
  page-break-inside: avoid;
  table-layout: auto;
}
thead { display: table-header-group; }
tr { page-break-inside: avoid; }
th {
  background: #f0ebff;
  color: #3d1d8a;
  font-weight: 600;
  padding: 7px 10px;
  border: 1px solid #d6c8ff;
  text-align: left;
  white-space: normal;
  overflow-wrap: anywhere;
  word-break: normal;
}
th:first-child, td:first-child { min-width: 72px; }
td {
  padding: 6px 10px;
  border: 1px solid #e0e0e0;
  vertical-align: top;
  white-space: normal;
  word-break: normal;
  overflow-wrap: anywhere;
}
tr:nth-child(even) td { background: #fafafa; }
img { max-width: 100%; height: auto; }
code { background: #f3f0ff; color: #5b21b6; padding: 1px 5px; border-radius: 3px; font-size: 12px; }
pre { background: #f5f5f5; padding: 12px; border-radius: 4px; overflow-wrap: break-word; white-space: pre-wrap; }
code, pre { font-family: "Noto Sans Mono CJK SC", "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Microsoft YaHei", monospace; }
strong { color: #111; font-weight: 600; }
hr { border: none; border-top: 1px solid #e0e0e0; margin: 20px 0; }
.core-highlight-box {
  background: rgba(139, 92, 246, 0.07);
  border: 1.5px solid rgba(139, 92, 246, 0.22);
  border-radius: 10px;
  padding: 14px 18px 10px;
  margin: 10px 0 16px;
  page-break-inside: avoid;
}
.core-highlight-box h2,
.core-highlight-box h3,
.core-highlight-box h4 { margin-top: 0; }
.core-highlight-box h3 {
  margin: 14px 0 8px;
  padding: 5px 9px;
  border-radius: 7px;
  background: rgba(139, 92, 246, 0.10);
  color: #5b21b6;
  font-size: 14px;
}
.core-highlight-box li { margin-bottom: 8px; }
"""


def _set_pdf_page_height(doc: str, height_in: float) -> str:
    height_in = max(6.0, min(float(height_in), 500.0))
    return doc.replace(PDF_PAGE_RULE, f"@page {{ size: 10in {height_in:.2f}in; margin: 0; }}", 1)


def _wrap_report_highlights(html_text: str) -> str:
    html_text = re.sub(r"<!--CORE_START-->\s*", '<div class="core-highlight-box">', html_text)
    html_text = re.sub(r"\s*<!--CORE_END-->", "</div>", html_text)

    summary_title = r"(?:本章总结|本节总结|章节总结|本部分总结)"
    heading_pat = re.compile(
        rf"(<h[34][^>]*>\s*{summary_title}\s*[:：]?\s*</h[34]>)(.*?)(?=<h[1-6][^>]*>|$)",
        re.S,
    )
    html_text = heading_pat.sub(r'<div class="core-highlight-box">\1\2</div>', html_text)

    paragraph_pat = re.compile(
        rf"(<p>\s*{summary_title}\s*[:：].*?</p>(?:(?!<h[1-6][^>]*>|<table|<pre|<div).)*?)"
        rf"(?=<h[1-6][^>]*>|$)",
        re.S,
    )
    return paragraph_pat.sub(r'<div class="core-highlight-box">\1</div>', html_text)


def report_markdown_to_pdf(md_text: str, mode: str = "") -> bytes:
    md_text = _inject_disclaimer(md_text or "", mode=mode)
    body = markdown_lib.markdown(
        md_text,
        extensions=["extra", "sane_lists", "nl2br"],
        output_format="html5",
    )
    body = _wrap_report_highlights(body)
    title_m = re.search(r"^#\s+(.+?)$", md_text, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>{PDF_CSS}</style>
</head>
<body>{body}</body>
</html>"""
    return html_to_pdf_bytes(doc)


def _find_pdf_browser() -> str:
    candidates = [
        os.getenv("PDF_BROWSER_PATH", "").strip(),
        shutil.which("google-chrome"),
        shutil.which("google-chrome-stable"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        shutil.which("microsoft-edge"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
        "/snap/bin/chromium",
        "/usr/lib/chromium/chromium",
        "/usr/lib/chromium-browser/chromium-browser",
        "/opt/google/chrome/chrome",
        "/usr/bin/microsoft-edge",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return path
    raise RuntimeError("未找到可用于生成 PDF 的 Chrome/Edge/Chromium，请安装浏览器或设置 PDF_BROWSER_PATH")


def _free_local_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return int(s.getsockname()[1])


def _wait_for_cdp_target(port: int, timeout_seconds: float = 10.0) -> str:
    deadline = time.time() + timeout_seconds
    last_err = ""
    while time.time() < deadline:
        try:
            with urlopen(f"http://127.0.0.1:{port}/json", timeout=0.5) as resp:
                targets = json.loads(resp.read().decode("utf-8"))
            for target in targets:
                if target.get("type") == "page" and target.get("webSocketDebuggerUrl"):
                    return target["webSocketDebuggerUrl"]
        except Exception as exc:
            last_err = str(exc)
        time.sleep(0.1)
    raise RuntimeError(f"等待浏览器调试端口超时：{last_err}")


def _html_to_pdf_with_browser(doc: str) -> bytes:
    browser = _find_pdf_browser()
    last_err: Exception | None = None
    for headless_arg in ("--headless=new", "--headless"):
        try:
            return _html_to_pdf_with_browser_cmd(doc, browser, headless_arg)
        except Exception as exc:
            last_err = exc
            print(f"[pdf] browser render failed with {headless_arg}: {exc}", flush=True)
    raise RuntimeError(f"浏览器 PDF 生成失败：{last_err}")


def _html_to_pdf_with_browser_cmd(doc: str, browser: str, headless_arg: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="survey_pdf_") as tmp:
        tmp_path = Path(tmp)
        html_path = tmp_path / "report.html"
        profile_path = tmp_path / "profile"
        html_path.write_text(doc, encoding="utf-8")

        port = _free_local_port()
        cmd = [
            browser,
            headless_arg,
            f"--remote-debugging-port={port}",
            f"--user-data-dir={profile_path}",
            "--disable-gpu",
            "--no-sandbox",
            "--disable-dev-shm-usage",
            "--hide-scrollbars",
            "--window-size=960,800",
            html_path.as_uri(),
        ]
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        try:
            ws_url = _wait_for_cdp_target(port)
            from websockets.sync.client import connect

            with connect(ws_url, open_timeout=10, max_size=None) as ws:
                msg_id = 0

                def cdp(method: str, params: dict | None = None) -> dict:
                    nonlocal msg_id
                    msg_id += 1
                    ws.send(json.dumps({"id": msg_id, "method": method, "params": params or {}}))
                    while True:
                        data = json.loads(ws.recv())
                        if data.get("id") != msg_id:
                            continue
                        if "error" in data:
                            raise RuntimeError(f"{method} failed: {data['error']}")
                        return data.get("result", {})

                cdp("Page.enable")
                cdp("Runtime.evaluate", {
                    "expression": (
                        "new Promise(resolve => {"
                        " if (document.readyState === 'complete') resolve(true);"
                        " else window.addEventListener('load', () => resolve(true), {once:true});"
                        "})"
                    ),
                    "awaitPromise": True,
                    "returnByValue": True,
                })
                height_result = cdp("Runtime.evaluate", {
                    "expression": (
                        "Math.ceil(Math.max("
                        "document.body.scrollHeight,"
                        "document.documentElement.scrollHeight,"
                        "document.body.offsetHeight,"
                        "document.documentElement.offsetHeight"
                        "))"
                    ),
                    "returnByValue": True,
                })
                height_px = int(height_result.get("result", {}).get("value") or 1200)
                width_in = 10.0
                height_in = max(6.0, min((height_px + 24) / 96, 500.0))
                pdf_params = {
                    "printBackground": True,
                    "paperWidth": width_in,
                    "paperHeight": height_in,
                    "marginTop": 0,
                    "marginBottom": 0,
                    "marginLeft": 0,
                    "marginRight": 0,
                    "preferCSSPageSize": False,
                    "scale": 1,
                    "generateDocumentOutline": True,
                }
                try:
                    pdf_result = cdp("Page.printToPDF", pdf_params)
                except RuntimeError as exc:
                    if "generateDocumentOutline" not in str(exc):
                        raise
                    pdf_params.pop("generateDocumentOutline", None)
                    pdf_result = cdp("Page.printToPDF", pdf_params)
                return base64.b64decode(pdf_result["data"])
        finally:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()


def html_to_pdf_bytes(doc: str) -> bytes:
    renderer = os.getenv("PDF_RENDERER", "").strip().lower()
    if renderer == "browser":
        try:
            return _html_to_pdf_with_browser(doc)
        except Exception as exc:
            print(f"[pdf] requested browser renderer failed, falling back to WeasyPrint: {exc}", flush=True)

    try:
        return _html_to_pdf_with_weasyprint(doc)
    except Exception as exc:
        print(f"[pdf] WeasyPrint renderer failed, retrying browser renderer: {exc}", flush=True)

    try:
        return _html_to_pdf_with_browser(doc)
    except Exception as exc:
        print(f"[pdf] browser renderer failed after WeasyPrint fallback: {exc}", flush=True)
        raise


def _html_to_pdf_with_weasyprint(doc: str) -> bytes:
    from weasyprint import HTML  # type: ignore

    html_obj = HTML(string=doc)
    rendered = html_obj.render()
    pages = getattr(rendered, "pages", []) or []
    if len(pages) <= 1:
        print("[pdf] rendered with WeasyPrint on one page", flush=True)
        return rendered.write_pdf()

    total_height_px = sum(float(getattr(page, "height", 14 * 96) or 14 * 96) for page in pages)
    total_height_in = min(max(total_height_px / 96 + 1.0, 14.0), 500.0)
    print(
        f"[pdf] WeasyPrint first pass produced {len(pages)} pages; rerendering as {total_height_in:.2f}in single page",
        flush=True,
    )
    return HTML(string=_set_pdf_page_height(doc, total_height_in)).write_pdf()


@app.get("/api/export/word/{session_id}")
async def export_word(session_id: str, request: Request):
    sess = get_session(session_id)
    report_md = sess.get("report_md", "")
    if not report_md:
        raise HTTPException(status_code=400, detail="还没有生成报告")

    report_md = _prep_export_md(report_md, mode=sess.get("mode") or "")
    title_m = re.search(r"^#\s+(.+?)$", report_md, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    safe = re.sub(r'[\\/:*?"<>|]', "_", title)

    loop = asyncio.get_event_loop()
    docx_bytes = await loop.run_in_executor(None, markdown_to_docx, report_md)
    await audit_log(request, "report", "下载 Word", f"报告：{title}", metadata={"session_id": session_id})
    return _make_download_response(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe}.docx",
    )


@app.get("/api/export/markdown/{session_id}")
async def export_markdown(session_id: str, request: Request):
    sess = get_session(session_id)
    report_md = sess.get("report_md", "")
    if not report_md:
        raise HTTPException(status_code=400, detail="还没有生成报告")

    report_md = _prep_export_md(report_md, mode=sess.get("mode") or "")
    title_m = re.search(r"^#\s+(.+?)$", report_md, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    safe = re.sub(r'[\\/:*?"<>|]', "_", title)
    await audit_log(request, "report", "下载 Markdown", f"报告：{title}", metadata={"session_id": session_id})
    return _make_download_response(report_md.encode("utf-8"), "text/markdown; charset=utf-8", f"{safe}.md")


@app.get("/api/export/pdf/{session_id}")
async def export_pdf(session_id: str, request: Request):
    sess = get_session(session_id)
    report_md = sess.get("report_md", "")
    if not report_md:
        raise HTTPException(status_code=400, detail="还没有生成报告")

    title_m = re.search(r"^#\s+(.+?)$", report_md, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    safe = re.sub(r'[\\/:*?"<>|]', "_", title)

    loop = asyncio.get_event_loop()
    pdf_bytes = await loop.run_in_executor(None, report_markdown_to_pdf, report_md, sess.get("mode") or "")
    await audit_log(request, "report", "下载 PDF", f"报告：{title}", metadata={"session_id": session_id})
    return _make_download_response(pdf_bytes, "application/pdf", f"{safe}.pdf")


@app.get("/api/export/word-history/{history_id}")
async def export_word_history(history_id: str, request: Request):
    login = await _current_login(request)
    history = _load_history()
    entry = _find_history_for_login(history, history_id, login)
    if not entry:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    entry_mode = entry.get("mode") or entry.get("plan", {}).get("mode", "")
    report_md = _prep_export_md(entry.get("report_md", ""), mode=entry_mode)
    safe = re.sub(r'[\\/:*?"<>|]', "_", entry.get("title", "调研报告"))
    loop = asyncio.get_event_loop()
    docx_bytes = await loop.run_in_executor(None, markdown_to_docx, report_md)
    await audit_log(request, "report", "下载历史 Word", f"报告：{entry.get('title', history_id)}", metadata={"history_id": history_id})
    return _make_download_response(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe}.docx",
    )


@app.get("/api/export/pdf-history/{history_id}")
async def export_pdf_history(history_id: str, request: Request):
    login = await _current_login(request)
    history = _load_history()
    entry = _find_history_for_login(history, history_id, login)
    if not entry:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    report_md = entry.get("report_md", "")
    if not report_md:
        raise HTTPException(status_code=400, detail="该历史记录没有报告内容")
    entry_mode = entry.get("mode") or entry.get("plan", {}).get("mode", "")
    safe = re.sub(r'[\\/:*?"<>|]', "_", entry.get("title", "调研报告"))
    loop = asyncio.get_event_loop()
    pdf_bytes = await loop.run_in_executor(
        None, report_markdown_to_pdf, report_md, entry_mode
    )
    await audit_log(request, "report", "下载历史 PDF", f"报告：{entry.get('title', history_id)}", metadata={"history_id": history_id})
    return _make_download_response(pdf_bytes, "application/pdf", f"{safe}.pdf")


# ── 飞书 OAuth 登录 + 文档导出 ────────────────────────

# Web 登录态（内存 + 本地文件；服务热更新/重启后仍保留 7 天会话）
def _extract_core_lines(md: str) -> list[str]:
    """取出 <!--CORE_START-->..<!--CORE_END--> 之间的内容行（供飞书高亮块用）。"""
    if not md or CORE_START not in md:
        return []
    try:
        seg = md.split(CORE_START, 1)[1].split(CORE_END, 1)[0]
    except IndexError:
        return []
    return [ln for ln in seg.split("\n") if ln.strip()]


def _drop_first_h1(md: str) -> str:
    """飞书文档标题栏已展示 title，正文导入前去掉首个 H1，避免重复标题。"""
    lines = md.split("\n")
    for i, line in enumerate(lines):
        if line.startswith("# ") and not line.startswith("## "):
            del lines[i]
            while i < len(lines) and not lines[i].strip():
                del lines[i]
            return "\n".join(lines).lstrip()
    return md


def _extract_feishu_callout_sections(md: str) -> list[dict]:
    sections: list[dict] = []
    core_lines = _extract_core_lines(md)
    if core_lines:
        sections.append({"title": "核心结论", "lines": core_lines, "occurrence": 1})

    occurrence_counts: dict[str, int] = {}
    lines = md.split("\n")
    summary_re = re.compile(r"^(#{3,4})\s+(本章总结|本节总结|章节总结|本部分总结)\s*[:：]?\s*$")
    i = 0
    while i < len(lines):
        match = summary_re.match(lines[i].strip())
        if not match:
            i += 1
            continue
        title = match.group(2)
        level = len(match.group(1))
        body: list[str] = []
        i += 1
        while i < len(lines):
            heading = re.match(r"^(#{1,6})\s+", lines[i].strip())
            if heading and len(heading.group(1)) <= level:
                break
            body.append(lines[i])
            i += 1
        occurrence_counts[title] = occurrence_counts.get(title, 0) + 1
        sections.append({"title": title, "lines": body, "occurrence": occurrence_counts[title]})
    return sections


async def _export_to_feishu(report_md: str, login: dict, mode: str = "") -> str:
    """将报告上传为飞书文档（docx），文档归登录用户所有，并通过机器人发消息通知。"""
    full = _prep_export_md(report_md, mode=mode)  # 补免责声明 + 去掉 CORE_START/END 标记
    title_m = re.search(r"^#\s+(.+?)$", full, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    open_id = login.get("open_id", "") or None
    url, _, _ = await feishu_export.create_doc_via_bot(title, full, open_id)
    print(f"[feishu-export] created doc title={title!r} url={url}")
    if open_id:
        await feishu_export.send_message_to_user(
            open_id,
            f"您的调研报告《{title}》已创建为飞书文档，点击查看：{url}"
        )
    return url


def _feishu_export_error(e: Exception) -> HTTPException:
    msg = str(e)
    if (
        "99991679" in msg
        or "drive:file:upload" in msg
        or "Unauthorized" in msg
        or "创建飞书云文档导入任务失败" in msg
        or "查询飞书云文档导入任务失败" in msg
    ):
        return HTTPException(
            status_code=403,
            detail=(
                "飞书授权缺少云文档上传/导入权限。请确认飞书开放平台已开通文件上传和云文档导入任务权限，"
                "然后点击左下角飞书账号退出登录，"
                "然后重新登录授权后再导出。"
            ),
        )
    return HTTPException(status_code=502, detail=f"创建飞书文档失败：{e}")


@app.post("/api/export/feishu/{session_id}")
async def export_feishu(session_id: str, request: Request):
    if not feishu_export.is_configured():
        raise HTTPException(status_code=500, detail="未配置飞书应用")
    login = await _current_login(request)
    if not login:
        raise HTTPException(status_code=401, detail="请先登录飞书")
    sess = get_session(session_id)
    report_md = sess.get("report_md", "")
    if not report_md:
        raise HTTPException(status_code=400, detail="还没有生成报告")
    try:
        url = await _export_to_feishu(report_md, login, mode=sess.get("mode") or "")
    except Exception as e:
        print(f"[feishu-export][ERROR] {e!r}")
        raise _feishu_export_error(e)
    await audit_log(request, "report", "导出飞书文档", f"会话：{session_id}", metadata={"session_id": session_id})
    return {"url": url, "type": "doc"}


@app.post("/api/export/feishu-history/{history_id}")
async def export_feishu_history(history_id: str, request: Request):
    if not feishu_export.is_configured():
        raise HTTPException(status_code=500, detail="未配置飞书应用")
    login = await _current_login(request)
    if not login:
        raise HTTPException(status_code=401, detail="请先登录飞书")
    history = _load_history()
    entry = _find_history_for_login(history, history_id, login)
    if not entry:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    try:
        entry_mode = entry.get("mode") or entry.get("plan", {}).get("mode", "")
        url = await _export_to_feishu(
            entry.get("report_md", ""), login,
            mode=entry_mode,
        )
    except Exception as e:
        print(f"[feishu-export-history][ERROR] {e!r}")
        raise _feishu_export_error(e)
    await audit_log(request, "report", "导出历史飞书文档", f"报告：{entry.get('title', history_id)}", metadata={"history_id": history_id})
    return {"url": url, "type": "doc"}


# ── Prompts 管理 ─────────────────────────────────────

# ── 历史记录 ──────────────────────────────────────────

@app.get("/api/history")
async def get_history(request: Request, mode: str = ""):
    login = await _current_login(request)
    history = _load_history()
    history = _ensure_history_report_numbers(history)
    visible_history = [h for h in history if _visible_to_owner(h, login)]
    if mode:
        visible_history = [h for h in visible_history if (h.get("mode") or "") == mode]
    await audit_log(request, "report", "查看历史记录", f"历史报告数：{len(visible_history)}")
    # 列表视图不返回完整 report_md（节省带宽）
    return [
        {
            "id": h["id"],
            "report_no": h.get("report_no", ""),
            "filename": h["filename"],
            "title": h["title"],
            "created_at": h["created_at"],
            "has_qa": bool(h.get("analyst_conv_id")),
            "qa_count": _qa_user_count(h),
            "mode": h.get("mode", ""),
            "row_count": _history_effective_row_count(h),
            "comment_valid_count": h.get("comment_valid_count", 0),
            "comment_sample_count": h.get("comment_sample_count", 0),
            "annotate_ai_count": h.get("annotate_ai_count", 0),
            "annotate_confirmed_ai_count": h.get("annotate_confirmed_ai_count", 0),
            "annotate_quality_count": h.get("annotate_quality_count", 0),
            "annotate_has_download": bool(h.get("annotate_result_path")),
        }
        for h in visible_history
    ]


@app.get("/api/history/{hist_id}")
async def get_history_item(hist_id: str, request: Request):
    login = await _current_login(request)
    history = _load_history()
    history = _ensure_history_report_numbers(history)
    entry = _find_history_for_login(history, hist_id, login)
    if not entry:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    await audit_log(request, "report", "打开历史报告", f"报告：{entry.get('title', hist_id)}", metadata={"history_id": hist_id})
    return entry


def _update_history_title_by_id(hist_id: str, title: str, login: dict | None = None) -> dict:
    hist_id = str(hist_id or "").strip()
    new_title = _sanitize_report_title(title)
    history = _load_history()
    history = _ensure_history_report_numbers(history, save=False)
    entry = _find_history_for_login(history, hist_id, login)
    if not entry:
        try:
            sess = get_session(hist_id)
            if sess.get("report_md") and _visible_to_owner(sess, login):
                _assign_session_owner(sess, login)
                sess["report_md"] = _replace_report_h1(sess.get("report_md", ""), new_title)
                save_session(hist_id, sess)
                save_to_history(hist_id, sess)
                history = _load_history()
                history = _ensure_history_report_numbers(history, save=False)
                entry = _find_history_for_login(history, hist_id, login)
        except HTTPException:
            pass
    if not entry:
        print(f"[history-title] not found id={hist_id!r} existing={[h.get('id') for h in history]}")
        raise HTTPException(status_code=404, detail="未找到这份报告，请刷新历史记录后重试")

    entry["title"] = new_title
    entry["report_md"] = _replace_report_h1(entry.get("report_md", ""), new_title)

    try:
        sess = get_session(hist_id)
        if sess.get("report_md"):
            sess["report_md"] = _replace_report_h1(sess.get("report_md", ""), new_title)
            save_session(hist_id, sess)
    except HTTPException:
        pass  # session 已过期，只更新历史记录即可

    _save_history(history)
    return {
        "ok": True,
        "id": hist_id,
        "report_no": entry.get("report_no", ""),
        "title": new_title,
        "report_md": entry.get("report_md", ""),
    }


@app.patch("/api/history-title")
async def update_history_title_by_body(req: HistoryTitleUpdateByIdRequest, request: Request):
    login = await _current_login(request)
    result = _update_history_title_by_id(req.id, req.title, login)
    await audit_log(request, "report", "修改报告名称", f"{req.id} → {result.get('title', req.title)}", metadata={"history_id": req.id})
    return result


@app.post("/api/history-title")
async def update_history_title_by_body_post(req: HistoryTitleUpdateByIdRequest, request: Request):
    login = await _current_login(request)
    result = _update_history_title_by_id(req.id, req.title, login)
    await audit_log(request, "report", "修改报告名称", f"{req.id} → {result.get('title', req.title)}", metadata={"history_id": req.id})
    return result


@app.patch("/api/history/{hist_id}/title")
async def update_history_title(hist_id: str, req: HistoryTitleUpdateRequest, request: Request):
    login = await _current_login(request)
    result = _update_history_title_by_id(hist_id, req.title, login)
    await audit_log(request, "report", "修改报告名称", f"{hist_id} → {result.get('title', req.title)}", metadata={"history_id": hist_id})
    return result


@app.post("/api/history/{hist_id}/title")
async def update_history_title_post(hist_id: str, req: HistoryTitleUpdateRequest, request: Request):
    login = await _current_login(request)
    result = _update_history_title_by_id(hist_id, req.title, login)
    await audit_log(request, "report", "修改报告名称", f"{hist_id} → {result.get('title', req.title)}", metadata={"history_id": hist_id})
    return result


# ============================================================
# 数据标注模块
# ============================================================

annotate_sessions: dict[str, dict] = {}
_ANNOTATE_SESSION_TTL = 7200  # 2 hours


def _clean_annotate_sessions() -> None:
    cutoff = time.time() - _ANNOTATE_SESSION_TTL
    expired = [k for k, v in annotate_sessions.items() if v.get("ts", 0) < cutoff]
    for k in expired:
        annotate_sessions.pop(k, None)


def _new_annotate_session() -> str:
    _clean_annotate_sessions()
    sid = str(uuid.uuid4())
    annotate_sessions[sid] = {"ts": time.time()}
    return sid


def _get_annotate_session(sid: str) -> dict:
    sess = annotate_sessions.get(sid)
    if not sess:
        raise HTTPException(status_code=404, detail="标注会话不存在或已过期，请重新上传文件")
    sess["ts"] = time.time()
    return sess


# ── 上传（标注专用）─────────────────────────────────────

async def _translate_headers(headers: list) -> list:
    """将表头翻译为中文简体，失败时原样返回。"""
    if not DIFY_AI_DETECT_KEY:
        return headers
    query = (
        "将以下问卷列名按顺序翻译为中文简体，只输出 JSON 数组，不加其他任何内容：\n"
        + json.dumps(headers, ensure_ascii=False)
    )
    full_text = ""
    try:
        async for chunk, _ in sse_dify_stream(query, "hdr-translate", "", DIFY_AI_DETECT_KEY):
            full_text += chunk
        result = _parse_string_array(full_text)
        if result and len(result) == len(headers):
            return result
    except Exception:
        pass
    return headers


def _parse_string_array(text: str) -> list[str] | None:
    """从 LLM 输出中提取字符串数组，容忍值内部的裸双引号。"""
    m = re.search(r'\[.*\]', text, re.DOTALL)
    if not m:
        return None
    raw = m.group()
    # 先尝试直接解析
    try:
        result = json.loads(raw)
        if isinstance(result, list):
            return [str(r) for r in result]
    except json.JSONDecodeError:
        pass
    # 兜底：逐行提取——取每行最外层引号之间的内容
    items = []
    for line in raw.splitlines():
        line = line.strip().rstrip(',')
        if line.startswith('"') and line.endswith('"') and len(line) >= 2:
            items.append(line[1:-1])
    return items if items else None


@app.post("/api/annotate/upload")
async def annotate_upload(request: Request, file: UploadFile = File(...)):
    content = await file.read()
    try:
        rows = _parse_file(file.filename or "upload.csv", content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not rows or len(rows) <= 1:
        raise HTTPException(status_code=400, detail="文件为空或只有表头")

    headers = rows[0]
    body    = rows[1:]

    # 自动检测列（用 server.py 现有逻辑，含矩阵题过滤）
    id_col         = annotate.detect_id_column(headers, rows)
    open_text_cols = _detect_open_text_cols(rows, headers)

    # 矩阵子列索引（前端隐藏，不在可选列表中显示）
    matrix_col_idxs: list[int] = []
    for g in _group_googleform_matrix(headers):
        if g["type"] == "matrix":
            matrix_col_idxs.extend(g["member_indexes"])

    # 翻译表头（用于前端展示，不影响处理逻辑）
    headers_zh = await _translate_headers(headers)

    sid = _new_annotate_session()
    sess = {
        "rows":            rows,
        "headers":         headers,
        "headers_zh":      headers_zh,
        "filename":        file.filename or "upload.csv",
        "id_col":          id_col,
        "open_text_cols":  open_text_cols,
    }
    _assign_session_owner(sess, await _current_login(request))
    annotate_sessions[sid].update(sess)

    result = {
        "session_id":       sid,
        "filename":         file.filename,
        "total_rows":       len(body),
        "headers":          headers,
        "headers_zh":       headers_zh,
        "id_col":           id_col,
        "open_text_cols":   open_text_cols,
        "matrix_col_idxs":  matrix_col_idxs,
        "preview":          rows[1: min(4, len(rows))],
    }
    await audit_log(
        request,
        "annotate",
        "上传标注数据",
        f"文件：{file.filename or 'unknown'}；样本行数：{len(body)}",
        metadata={"session_id": sid, "rows": len(body)},
    )
    return result


# ── 确认列 + 任务选择 ────────────────────────────────────

@app.post("/api/annotate/{sid}/confirm-columns")
async def annotate_confirm_columns(sid: str, req: AnnotateConfirmRequest, request: Request):
    sess = _get_annotate_session(sid)
    sess["id_col"]        = req.id_col
    sess["open_text_cols"] = req.open_text_cols
    sess["tasks"]         = req.tasks
    sess["background"]    = req.background
    sess["ai_results"]       = []
    sess["confirmed_ai_ids"] = []
    sess["quality_results"]  = []
    sess.pop("missing_ai_ids", None)
    sess.pop("missing_quality_ids", None)
    task_names = []
    if req.tasks.get("ai_detect"):
        task_names.append("AI 作答识别")
    if req.tasks.get("quality"):
        task_names.append("回答质量打标")
    await audit_log(
        request,
        "annotate",
        "确认标注任务",
        f"会话：{sid}；主观题列数：{len(req.open_text_cols)}；任务：{', '.join(task_names) or '未选择'}",
        metadata={"session_id": sid, "open_text_cols": len(req.open_text_cols), "tasks": req.tasks},
    )
    return {"ok": True}


# ── AI 作答识别（SSE）───────────────────────────────────

@app.get("/api/annotate/{sid}/run-ai-detect")
async def annotate_run_ai_detect(sid: str, request: Request):
    sess = _get_annotate_session(sid)
    rows           = sess.get("rows", [])
    headers        = sess.get("headers", [])
    id_col         = sess.get("id_col", 1)
    open_text_cols = sess.get("open_text_cols", [])
    background     = sess.get("background", "")

    if not rows or not open_text_cols:
        raise HTTPException(status_code=400, detail="会话状态不完整，请重新上传")
    if not DIFY_AI_DETECT_KEY:
        raise HTTPException(status_code=500, detail="未配置 DIFY_AI_DETECT_KEY")

    body = rows[1:]
    batch_size = annotate.AI_DETECT_BATCH
    batches = [body[i: i + batch_size] for i in range(0, len(body), batch_size)]
    total_batches = len(batches)

    def _has_open_text(row: list) -> bool:
        return any((str(row[c]) if c < len(row) else "").strip() for c in open_text_cols)

    def _row_id(row: list) -> str:
        return str(row[id_col]).strip() if id_col < len(row) else ""

    def _empty_ai_result(row: list, reason: str) -> dict:
        return {
            "id": _row_id(row),
            "ai_prob": 0,
            "is_polished": "low",
            "reason": reason,
            "evidence": "",
            "originals": _open_text_originals(row),
            "translations": {},
        }

    def _chunks(items: list, size: int) -> list[list]:
        return [items[i: i + size] for i in range(0, len(items), size)]

    def _open_text_originals(row: list) -> dict[str, str]:
        return {
            f"col_{c}": str(row[c]).strip()
            for c in open_text_cols
            if c < len(row) and str(row[c]).strip()
        }

    def _attach_originals(results: list[dict], batch_rows: list[list]) -> list[dict]:
        rows_by_id = {_row_id(row): row for row in batch_rows if _row_id(row)}
        for result in results:
            row = rows_by_id.get(str(result.get("id", "")).strip())
            if row is not None:
                result["originals"] = _open_text_originals(row)
        return results

    async def _run_ai_direct(batch_rows: list[list], label: str) -> tuple[list[dict], str]:
        query = annotate.build_ai_detect_query(
            batch_rows, headers, open_text_cols, id_col, label, background
        )
        _annotate_ai_log("subbatch start", sid=sid, batch=label, rows=len(batch_rows), query_len=len(query))
        try:
            text, final_conv, mode, fallback_reason = await call_dify_compatible(
                query, f"{sid}-split-{label}", DIFY_AI_DETECT_KEY
            )
            _annotate_ai_log(
                "subbatch dify done",
                sid=sid,
                batch=label,
                mode=mode,
                answer_len=len(text or ""),
                fallback=bool(fallback_reason),
            )
            if not (text or "").strip():
                return [], "Dify 返回空内容"
            results, err = annotate.parse_ai_detect_result(text)
            if results:
                return _attach_originals(results, batch_rows), ""

            retry_q = (
                f"上次输出无法解析（{err}）。请重新处理下面这批数据，并严格按 schema 用 ```json``` 围栏输出，"
                "不要附加任何解释文字。\n\n"
                f"{query}"
            )
            retry_text, _, retry_mode, retry_fallback = await call_dify_compatible(
                retry_q, f"{sid}-split-retry-{label}", DIFY_AI_DETECT_KEY, final_conv
            )
            _annotate_ai_log(
                "subbatch retry done",
                sid=sid,
                batch=label,
                mode=retry_mode,
                answer_len=len(retry_text or ""),
                fallback=bool(retry_fallback),
            )
            results, retry_err = annotate.parse_ai_detect_result(retry_text)
            if results:
                return _attach_originals(results, batch_rows), ""
            return results, retry_err
        except Exception as exc:
            _annotate_ai_log("subbatch failed", sid=sid, batch=label, error=str(exc)[:1000])
            return [], str(exc)

    async def generate():
        all_results: list[dict] = []
        all_missing_ids: set[str] = set()
        try:
            yield sse_event({
                "type": "started",
                "rows": len(body),
                "total_batches": total_batches,
                "batch_size": batch_size,
                "msg": f"已连接，准备分析 {len(body)} 行，约 {total_batches} 批",
            })

            for batch_num, batch in enumerate(batches, 1):
                empty_rows = [r for r in batch if not _has_open_text(r)]
                active_batch = [r for r in batch if _has_open_text(r)]
                missing_ids = sum(1 for r in active_batch if not _row_id(r))
                yield sse_event({
                    "type": "batch_started",
                    "batch": batch_num,
                    "done": batch_num - 1,
                    "total": total_batches,
                    "rows": len(active_batch),
                    "skipped": len(empty_rows),
                    "missing_ids": missing_ids,
                    "msg": f"正在分析第 {batch_num}/{total_batches} 批（{len(active_batch)} 行，跳过 {len(empty_rows)} 行空主观题）",
                })
                if empty_rows:
                    all_results.extend(_empty_ai_result(r, "主观题为空，系统自动判定为非 AI 作答") for r in empty_rows)
                if missing_ids:
                    msg = f"第 {batch_num} 批有 {missing_ids} 行缺少玩家唯一 ID，结果可能无法正确回填"
                    _annotate_ai_log("missing ids", sid=sid, batch=batch_num, count=missing_ids)
                    yield sse_event({"type": "warn", "msg": msg})
                if not active_batch:
                    msg = f"第 {batch_num} 批没有可分析的主观题内容，已跳过 AI 调用"
                    _annotate_ai_log("skip empty batch", sid=sid, batch=batch_num, rows=len(batch))
                    yield sse_event({"type": "warn", "msg": msg})
                    yield sse_event({
                        "type": "batch_done",
                        "batch": batch_num,
                        "done": batch_num,
                        "total": total_batches,
                        "count": len(empty_rows),
                        "msg": f"第 {batch_num}/{total_batches} 批完成，空主观题 {len(empty_rows)} 行已自动跳过",
                    })
                    continue

                query = annotate.build_ai_detect_query(
                    active_batch, headers, open_text_cols, id_col, batch_num, background
                )
                _annotate_ai_log(
                    "batch start",
                    sid=sid,
                    batch=batch_num,
                    rows=len(active_batch),
                    skipped=len(empty_rows),
                    missing_ids=missing_ids,
                    query_len=len(query),
                )
                try:
                    dify_task = asyncio.create_task(call_dify_compatible(
                        query, sid, DIFY_AI_DETECT_KEY
                    ))
                    while not dify_task.done():
                        yield sse_event({
                            "type": "dify_waiting",
                            "batch": batch_num,
                            "total": total_batches,
                            "msg": "正在等待 AI 返回，请勿关闭页面",
                        })
                        await asyncio.sleep(12)
                    answer_text, final_conv, mode, fallback_reason = await dify_task
                    _annotate_ai_log(
                        "dify done",
                        sid=sid,
                        batch=batch_num,
                        mode=mode,
                        answer_len=len(answer_text or ""),
                        fallback=bool(fallback_reason),
                    )
                    yield sse_event({
                        "type": "dify_done",
                        "batch": batch_num,
                        "mode": mode,
                        "answer_len": len(answer_text or ""),
                        "msg": f"第 {batch_num} 批 AI 返回完成（{mode}，{len(answer_text or '')} 字符）",
                    })
                    if fallback_reason:
                        _annotate_ai_log(
                            "fallback",
                            sid=sid,
                            batch=batch_num,
                            reason=fallback_reason[:500],
                        )
                        yield sse_event({
                            "type": "warn",
                            "msg": f"第 {batch_num} 批 chat 调用不匹配，已自动改用 completion 调用：{fallback_reason[:240]}",
                        })
                except Exception as e:
                    _annotate_ai_log("dify failed", sid=sid, batch=batch_num, error=str(e)[:1000])
                    yield sse_event({"type": "warn", "msg": f"第 {batch_num} 批 Dify 调用失败：{e}"})
                    split_results: list[dict] = []
                    if len(active_batch) > 1:
                        yield sse_event({
                            "type": "warn",
                            "msg": f"第 {batch_num} 批已自动拆成更小子批次重试，避免 Dify 插件/模型超时",
                        })
                        for sub_idx, sub_batch in enumerate(_chunks(active_batch, 2), 1):
                            sub_label = f"{batch_num}.{sub_idx}"
                            sub_results, sub_err = await _run_ai_direct(sub_batch, sub_label)
                            if sub_results:
                                split_results.extend(sub_results)
                                all_results.extend(sub_results)
                                yield sse_event({
                                    "type": "warn",
                                    "msg": f"第 {sub_label} 子批次重试成功，获得 {len(sub_results)} 条结果",
                                })
                            else:
                                yield sse_event({
                                    "type": "warn",
                                    "msg": f"第 {sub_label} 子批次仍失败：{sub_err}",
                                })
                    yield sse_event({
                        "type": "batch_done",
                        "batch": batch_num,
                        "done": batch_num,
                        "total": total_batches,
                        "count": len(split_results),
                        "msg": f"第 {batch_num}/{total_batches} 批完成，拆分重试获得 {len(split_results)} 条结果，跳过 {len(empty_rows)} 行空主观题",
                    })
                    continue

                results, err = annotate.parse_ai_detect_result(answer_text)
                if not results:
                    snippet = (answer_text or "")[:500].replace("\n", " ")
                    _annotate_ai_log(
                        "parse failed",
                        sid=sid,
                        batch=batch_num,
                        answer_len=len(answer_text or ""),
                        error=err,
                        snippet=snippet,
                    )
                    retry_q = (
                        f"上次输出无法解析（{err}）。请重新处理下面这批数据，并严格按 schema 用 ```json``` 围栏输出，"
                        "不要附加任何解释文字。\n\n"
                        f"{query}"
                    )
                    yield sse_event({
                        "type": "warn",
                        "msg": f"第 {batch_num} 批首次解析失败，正在自动重试：{err}；返回长度 {len(answer_text or '')}；片段：{snippet[:180]}",
                    })
                    try:
                        retry_task = asyncio.create_task(call_dify_compatible(
                            retry_q, sid, DIFY_AI_DETECT_KEY, final_conv
                        ))
                        while not retry_task.done():
                            yield sse_event({
                                "type": "dify_waiting",
                                "batch": batch_num,
                                "total": total_batches,
                                "msg": "正在等待 AI 重试返回，请勿关闭页面",
                            })
                            await asyncio.sleep(12)
                        retry_text, _, retry_mode, retry_fallback = await retry_task
                        _annotate_ai_log(
                            "retry done",
                            sid=sid,
                            batch=batch_num,
                            mode=retry_mode,
                            answer_len=len(retry_text or ""),
                            fallback=bool(retry_fallback),
                        )
                        results, err = annotate.parse_ai_detect_result(retry_text)
                    except Exception as e:
                        _annotate_ai_log("retry failed", sid=sid, batch=batch_num, error=str(e)[:1000])
                        results, err = [], str(e)

                if results:
                    results = _attach_originals(results, active_batch)

                if not results:
                    _annotate_ai_log("batch no results", sid=sid, batch=batch_num, error=err)
                    yield sse_event({"type": "warn", "msg": f"第 {batch_num} 批解析失败：{err}"})
                    split_results = []
                    if len(active_batch) > 1:
                        yield sse_event({
                            "type": "warn",
                            "msg": f"第 {batch_num} 批将拆成更小子批次继续重试",
                        })
                        for sub_idx, sub_batch in enumerate(_chunks(active_batch, 2), 1):
                            sub_label = f"{batch_num}.{sub_idx}"
                            sub_results, sub_err = await _run_ai_direct(sub_batch, sub_label)
                            if sub_results:
                                split_results.extend(sub_results)
                                yield sse_event({
                                    "type": "warn",
                                    "msg": f"第 {sub_label} 子批次重试成功，获得 {len(sub_results)} 条结果",
                                })
                            else:
                                yield sse_event({
                                    "type": "warn",
                                    "msg": f"第 {sub_label} 子批次仍失败：{sub_err}",
                                })
                    if split_results:
                        all_results.extend(split_results)
                        results = split_results
                    else:
                        batch_ids = {_row_id(r) for r in active_batch if _row_id(r)}
                        if batch_ids:
                            all_missing_ids.update(batch_ids)
                            yield sse_event({
                                "type": "warn",
                                "msg": f"第 {batch_num} 批所有重试均失败，{len(batch_ids)} 行计入缺失，完成后将阻断下载",
                            })
                else:
                    all_results.extend(results)
                    _annotate_ai_log("batch parsed", sid=sid, batch=batch_num, count=len(results))

                if results:
                    expected_ids = {_row_id(r) for r in active_batch if _row_id(r)}
                    returned_ids = {r["id"] for r in results if r.get("id")}
                    missing = expected_ids - returned_ids
                    if missing:
                        yield sse_event({
                            "type": "warn",
                            "msg": f"第 {batch_num} 批漏返 {len(missing)} 行，正在补发重试…",
                        })
                        missing_rows = [r for r in active_batch if _row_id(r) in missing]
                        retry_results, _ = await _run_ai_direct(missing_rows, f"{batch_num}.miss")
                        if retry_results:
                            results = list(results) + retry_results
                            all_results.extend(retry_results)
                            missing -= {r["id"] for r in retry_results if r.get("id")}
                        if missing:
                            all_missing_ids.update(missing)
                            yield sse_event({
                                "type": "warn",
                                "msg": f"第 {batch_num} 批重试后仍漏返 {len(missing)} 行（ID：{', '.join(sorted(missing)[:5])}{'…' if len(missing) > 5 else ''}），完成后将阻断下载",
                            })

                yield sse_event({
                    "type": "batch_done",
                    "batch": batch_num,
                    "done": batch_num,
                    "total": total_batches,
                    "count": len(results),
                    "msg": f"第 {batch_num}/{total_batches} 批完成，获得 {len(results)} 条 AI 结果，跳过 {len(empty_rows)} 行空主观题",
                })

            sess["ai_results"] = all_results
            sess.pop("missing_ai_ids", None)
            if all_missing_ids:
                sess["missing_ai_ids"] = sorted(all_missing_ids)
            high_prob = [r for r in all_results if r.get("ai_prob", 0) >= 80]
            if not all_missing_ids and not high_prob and not (sess.get("tasks") or {}).get("quality"):
                await _save_annotate_result_history(sid, sess, request)
            await audit_log(
                request,
                "annotate",
                "完成 AI 作答识别",
                f"会话：{sid}；结果数：{len(all_results)}；高概率数：{len(high_prob)}；未回填：{len(all_missing_ids)}",
                metadata={"session_id": sid, "results": len(all_results), "high_prob": len(high_prob), "missing": len(all_missing_ids)},
            )
            yield sse_event({
                "type":        "ai_detect_done",
                "results":     all_results,
                "high_prob":   high_prob,
                "missing_ids": sorted(all_missing_ids),
            })
        except Exception as e:
            import traceback; traceback.print_exc()
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# 用户确认 AI 结果
@app.post("/api/annotate/{sid}/confirm-ai")
async def annotate_confirm_ai(sid: str, req: AnnotateConfirmAIRequest, request: Request):
    sess = _get_annotate_session(sid)
    sess["confirmed_ai_ids"] = req.confirmed_ai_ids
    if not (sess.get("tasks") or {}).get("quality"):
        await _save_annotate_result_history(sid, sess, request)
    await audit_log(
        request,
        "annotate",
        "确认 AI 作答结果",
        f"会话：{sid}；确认 AI 作答数：{len(req.confirmed_ai_ids)}",
        metadata={"session_id": sid, "confirmed_count": len(req.confirmed_ai_ids)},
    )
    return {"ok": True, "confirmed_count": len(req.confirmed_ai_ids)}


# ── 质量打标（SSE）──────────────────────────────────────

@app.get("/api/annotate/{sid}/run-quality")
async def annotate_run_quality(sid: str, request: Request):
    sess = _get_annotate_session(sid)
    rows              = sess.get("rows", [])
    headers           = sess.get("headers", [])
    id_col            = sess.get("id_col", 1)
    open_text_cols    = sess.get("open_text_cols", [])
    confirmed_ai_ids  = set(sess.get("confirmed_ai_ids", []))

    if not rows or not open_text_cols:
        raise HTTPException(status_code=400, detail="会话状态不完整，请重新上传")
    if not DIFY_QUALITY_KEY:
        raise HTTPException(status_code=500, detail="未配置 DIFY_QUALITY_KEY")

    # 质量打标仅处理非 AI 行
    body = rows[1:]
    non_ai_body = [r for r in body if str(r[id_col]).strip() not in confirmed_ai_ids] if body and id_col < len(headers) else body
    batch_size = annotate.QUALITY_BATCH
    batches = [non_ai_body[i: i + batch_size] for i in range(0, len(non_ai_body), batch_size)]
    total_batches = len(batches)

    QUALITY_CONCURRENCY = 3

    async def _run_one_batch(batch_num: int, batch: list) -> tuple[int, list[dict], set[str], str]:
        """运行单批质量打标，含解析重试和覆盖重试。永远返回结构化结果，不向外抛异常。"""
        batch_ids = {str(r[id_col]).strip() if id_col < len(r) else "" for r in batch}
        batch_ids.discard("")
        try:
            query = annotate.build_quality_label_query(batch, headers, open_text_cols, id_col, batch_num)
            chunks: list[str] = []
            final_conv = ""
            async for chunk, conv_id in sse_dify_stream(query, sid, "", DIFY_QUALITY_KEY):
                if chunk:
                    chunks.append(chunk)
                if conv_id:
                    final_conv = conv_id

            results, err = annotate.parse_quality_result("".join(chunks))
            if not results:
                retry_q = (
                    f"上次输出无法解析（{err}）。请严格按 schema 用 ```json``` 围栏重新输出，"
                    "不要附加任何解释文字。"
                )
                retry_chunks: list[str] = []
                async for chunk, _ in sse_dify_stream(retry_q, sid, final_conv, DIFY_QUALITY_KEY):
                    if chunk:
                        retry_chunks.append(chunk)
                results, err = annotate.parse_quality_result("".join(retry_chunks))

            if not results:
                return batch_num, [], batch_ids, err

            # 覆盖检查 + 漏返重试
            expected = set(batch_ids)
            still_missing: set[str] = set()
            if expected:
                returned = {r["id"] for r in results if r.get("id")}
                missing = expected - returned
                if missing:
                    missing_rows = [r for r in batch if (str(r[id_col]).strip() if id_col < len(r) else "") in missing]
                    retry_miss_q = annotate.build_quality_label_query(missing_rows, headers, open_text_cols, id_col, f"{batch_num}.miss")
                    miss_chunks: list[str] = []
                    async for chunk, _ in sse_dify_stream(retry_miss_q, sid, "", DIFY_QUALITY_KEY):
                        if chunk:
                            miss_chunks.append(chunk)
                    miss_results, _ = annotate.parse_quality_result("".join(miss_chunks))
                    if miss_results:
                        results.extend(miss_results)
                        missing -= {r["id"] for r in miss_results if r.get("id")}
                    still_missing = missing

            return batch_num, results, still_missing, ""
        except Exception as exc:
            return batch_num, [], batch_ids, str(exc)

    async def generate():
        all_results: list[dict] = []
        all_missing_ids_q: set[str] = set()
        pending: set[asyncio.Task] = set()
        try:
            sem = asyncio.Semaphore(QUALITY_CONCURRENCY)

            async def run_with_sem(batch_num: int, batch: list):
                async with sem:
                    return await _run_one_batch(batch_num, batch)

            pending = {
                asyncio.create_task(run_with_sem(i, b))
                for i, b in enumerate(batches, 1)
            }
            done_count = 0
            yield sse_event({
                "type": "progress", "done": 0, "total": total_batches,
                "msg": f"已连接，{len(non_ai_body)} 行分 {total_batches} 批，最多 {QUALITY_CONCURRENCY} 批并行",
            })

            while pending:
                finished, pending = await asyncio.wait(pending, return_when=asyncio.FIRST_COMPLETED)
                for task in finished:
                    batch_num, results, still_missing, err = await task
                    done_count += 1
                    if not results:
                        all_missing_ids_q.update(still_missing)
                        yield sse_event({
                            "type": "warn",
                            "msg": f"第 {batch_num} 批解析失败：{err}；{len(still_missing)} 行计入缺失，完成后将阻断下载",
                        })
                    else:
                        all_results.extend(results)
                        if still_missing:
                            all_missing_ids_q.update(still_missing)
                            yield sse_event({
                                "type": "warn",
                                "msg": f"第 {batch_num} 批重试后仍漏返 {len(still_missing)} 行（ID：{', '.join(sorted(still_missing)[:5])}{'…' if len(still_missing) > 5 else ''}），完成后将阻断下载",
                            })
                    yield sse_event({
                        "type": "progress", "done": done_count, "total": total_batches,
                        "msg": f"第 {batch_num} 批完成，获得 {len(results)} 条结果（{done_count}/{total_batches} 批已完成）",
                    })

            sess["quality_results"] = all_results
            sess.pop("missing_quality_ids", None)
            if all_missing_ids_q:
                sess["missing_quality_ids"] = sorted(all_missing_ids_q)
            if not all_missing_ids_q:
                await _save_annotate_result_history(sid, sess, request)
            await audit_log(
                request,
                "annotate",
                "完成回答质量打标",
                f"会话：{sid}；结果数：{len(all_results)}；未回填：{len(all_missing_ids_q)}",
                metadata={"session_id": sid, "results": len(all_results), "missing": len(all_missing_ids_q)},
            )
            yield sse_event({"type": "quality_done", "count": len(all_results), "missing_ids": sorted(all_missing_ids_q)})
        except Exception as e:
            import traceback; traceback.print_exc()
            for t in pending:
                t.cancel()
            if pending:
                await asyncio.gather(*pending, return_exceptions=True)
            yield sse_event({"type": "error", "message": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── 下载标注 Excel ───────────────────────────────────────

@app.get("/api/annotate/{sid}/download")
async def annotate_download(sid: str, request: Request):
    sess = _get_annotate_session(sid)
    filename = sess.get("filename", "annotated")
    loop = asyncio.get_event_loop()
    excel_bytes, download_name = await loop.run_in_executor(
        None,
        _build_annotate_excel_from_session,
        sess,
    )
    await _save_annotate_result_history(sid, sess, request)
    await audit_log(request, "annotate", "下载标注结果", f"会话：{sid}；文件：{filename}", metadata={"session_id": sid})
    return _make_download_response(
        excel_bytes,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        download_name,
    )


@app.get("/api/annotate-history/{history_id}/download")
async def annotate_history_download(history_id: str, request: Request):
    login = await _current_login(request)
    history = _load_history()
    history = _ensure_history_report_numbers(history)
    entry = _find_history_for_login(history, history_id, login)
    if not entry or entry.get("mode") != "annotate":
        raise HTTPException(status_code=404, detail="标注历史记录不存在")

    raw_path = str(entry.get("annotate_result_path") or "")
    if not raw_path:
        raise HTTPException(status_code=404, detail="这条历史记录没有可下载的标注文件")
    result_path = Path(raw_path)
    if not result_path.is_absolute():
        result_path = ANNOTATE_RESULT_DIR / result_path.name
    try:
        result_resolved = result_path.resolve(strict=True)
        root_resolved = ANNOTATE_RESULT_DIR.resolve()
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="标注结果文件已不存在")
    if root_resolved not in result_resolved.parents and result_resolved != root_resolved:
        raise HTTPException(status_code=400, detail="标注结果路径无效")

    download_name = entry.get("annotate_download_name") or _annotate_download_filename(entry.get("filename", "annotated"))
    await audit_log(
        request,
        "annotate",
        "下载历史标注结果",
        f"历史记录：{history_id}；文件：{entry.get('filename', '')}",
        metadata={"history_id": history_id},
    )
    return _make_download_response(
        result_resolved.read_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        download_name,
    )


# ── 启动 ────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=True)
