"""services/report_render:报告渲染与导出前处理。

Markdown 清洗(免责声明注入、核心标记处理)、Markdown→Word、Markdown→PDF、
飞书 callout 区块提取。被 export / 报告生成 等共同调用。
"""
import base64
import html
import io
import json
import os
import re
import shutil
import socket
import subprocess
import tempfile
import time
from pathlib import Path
from urllib.request import urlopen

import markdown as markdown_lib
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import (
    COMMENT_DISCLAIMER,
    CORE_END,
    CORE_START,
    QUALITATIVE_DISCLAIMER,
    REPORT_DISCLAIMER,
)


def _inject_disclaimer(md: str, mode: str = "") -> str:
    """在第一行 `# 标题` 之后插入免责声明引用行；幂等；无 H1 则插到最前。

    通用声明 REPORT_DISCLAIMER 三种模式都插。第二条声明按 mode 选择：
      - mode == "crosstab"（倍市得）：不插第二条，并清除历史报告里残留的定性/评论声明
      - mode == "comment"（评论舆情）：插 COMMENT_DISCLAIMER，清除残留的定性声明
      - 其它（定性，默认）：插 QUALITATIVE_DISCLAIMER，清除残留的评论声明
    """
    if not md:
        return md

    # 选定本模式应有的“第二条声明”
    if mode == "crosstab":
        second = None
    elif mode == "comment":
        second = COMMENT_DISCLAIMER
    else:
        second = QUALITATIVE_DISCLAIMER

    # 清除不属于当前模式的历史第二条声明（兼容历史报告 / 模式切换）
    for stale in (QUALITATIVE_DISCLAIMER, COMMENT_DISCLAIMER):
        if stale is not second and stale in md:
            md = md.replace("\n" + stale, "").replace(stale + "\n", "")

    has_report = REPORT_DISCLAIMER in md
    has_second = bool(second) and second in md

    # 已经齐备就不再插
    if has_report and (second is None or has_second):
        return md

    lines = md.split("\n")

    # 已有 REPORT_DISCLAIMER 但缺第二条声明
    if has_report and second and not has_second:
        for i, ln in enumerate(lines):
            if ln.strip() == REPORT_DISCLAIMER:
                lines.insert(i + 1, second)
                return "\n".join(lines)

    for i, ln in enumerate(lines):
        if ln.startswith("# ") and not ln.startswith("## "):
            lines.insert(i + 1, "")
            insert_at = i + 2
            if not has_report:
                lines.insert(insert_at, REPORT_DISCLAIMER)
                insert_at += 1
            if second and not has_second:
                lines.insert(insert_at, second)
            return "\n".join(lines)

    # 没有 H1：插到最前
    prefix = []
    if not has_report:
        prefix.append(REPORT_DISCLAIMER)
    if second and not has_second:
        prefix.append(second)
    return "\n".join(prefix) + "\n\n" + md


def _strip_core_markers(md: str) -> str:
    """移除核心结论包裹标记行（仅飞书导出用于定位高亮块，其它导出不应出现）。"""
    if not md:
        return md
    return "\n".join(ln for ln in md.split("\n") if ln.strip() not in (CORE_START, CORE_END))


def _prep_export_md(md: str, mode: str = "") -> str:
    """通用导出前处理：补免责声明（幂等）+ 去掉核心结论标记。"""
    return _strip_core_markers(_inject_disclaimer(md, mode=mode))


def _extract_core_lines(md: str) -> list[str]:
    """取出 <!--CORE_START-->..<!--CORE_END--> 之间的内容行（供飞书高亮块用）。"""
    if not md or CORE_START not in md:
        return []
    try:
        seg = md.split(CORE_START, 1)[1].split(CORE_END, 1)[0]
    except IndexError:
        return []
    return [ln for ln in seg.split("\n") if ln.strip()]


def _drop_first_h1(md: str) -> str:
    """飞书文档标题栏已展示 title，正文导入前去掉首个 H1，避免重复标题。"""
    lines = md.split("\n")
    for i, line in enumerate(lines):
        if line.startswith("# ") and not line.startswith("## "):
            del lines[i]
            while i < len(lines) and not lines[i].strip():
                del lines[i]
            return "\n".join(lines).lstrip()
    return md


def _extract_feishu_callout_sections(md: str) -> list[dict]:
    sections: list[dict] = []
    core_lines = _extract_core_lines(md)
    if core_lines:
        sections.append({"title": "核心结论", "lines": core_lines, "occurrence": 1})

    occurrence_counts: dict[str, int] = {}
    lines = md.split("\n")
    summary_re = re.compile(r"^(#{3,4})\s+(本章总结|本节总结|章节总结|本部分总结)\s*[:：]?\s*$")
    i = 0
    while i < len(lines):
        match = summary_re.match(lines[i].strip())
        if not match:
            i += 1
            continue
        title = match.group(2)
        level = len(match.group(1))
        body: list[str] = []
        i += 1
        while i < len(lines):
            heading = re.match(r"^(#{1,6})\s+", lines[i].strip())
            if heading and len(heading.group(1)) <= level:
                break
            body.append(lines[i])
            i += 1
        occurrence_counts[title] = occurrence_counts.get(title, 0) + 1
        sections.append({"title": title, "lines": body, "occurrence": occurrence_counts[title]})
    return sections


def _add_formatted_run(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    result = []
    for line in lines:
        if re.match(r'^\|[\s\-|:]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        result.append(cells)
    return result


def markdown_to_docx(md_text: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(level=1); h.clear()
            _add_formatted_run(h, line[2:].strip())
        elif line.startswith('## ') and not line.startswith('### '):
            h = doc.add_heading(level=2); h.clear()
            _add_formatted_run(h, line[3:].strip())
        elif line.startswith('### '):
            h = doc.add_heading(level=3); h.clear()
            _add_formatted_run(h, line[4:].strip())
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i]); i += 1
            table_data = _parse_md_table(table_lines)
            if table_data:
                num_cols = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=len(table_data), cols=num_cols)
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(table_data):
                    for ci in range(num_cols):
                        ct = row_data[ci].replace('\\|', '|') if ci < len(row_data) else ""
                        cell = tbl.cell(ri, ci); cell.text = ""
                        run = cell.paragraphs[0].add_run(ct)
                        if ri == 0: run.bold = True
            continue
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet'); p.clear()
            _add_formatted_run(p, line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number'); p.clear()
            _add_formatted_run(p, re.sub(r'^\d+\.\s', '', line).strip())
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_run(p, line[2:].strip())
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph()
        elif not line.strip():
            pass
        else:
            p = doc.add_paragraph(); _add_formatted_run(p, line.strip())

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def report_markdown_to_pdf(md_text: str, mode: str = "") -> bytes:
    md_text = _inject_disclaimer(md_text or "", mode=mode)
    body = markdown_lib.markdown(
        md_text,
        extensions=["extra", "sane_lists", "nl2br"],
        output_format="html5",
    )
    body = _wrap_report_highlights(body)
    title_m = re.search(r"^#\s+(.+?)$", md_text, re.MULTILINE)
    title = title_m.group(1).strip() if title_m else "调研报告"
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>{PDF_CSS}</style>
</head>
<body>{body}</body>
</html>"""
    return html_to_pdf_bytes(doc)


def _html_to_pdf_with_weasyprint(doc: str) -> bytes:
    from weasyprint import HTML  # type: ignore

    html_obj = HTML(string=doc)
    rendered = html_obj.render()
    pages = getattr(rendered, "pages", []) or []
    if len(pages) <= 1:
        print("[pdf] rendered with WeasyPrint on one page", flush=True)
        return rendered.write_pdf()

    total_height_px = sum(float(getattr(page, "height", 14 * 96) or 14 * 96) for page in pages)
    total_height_in = min(max(total_height_px / 96 + 1.0, 14.0), 500.0)
    print(
        f"[pdf] WeasyPrint first pass produced {len(pages)} pages; rerendering as {total_height_in:.2f}in single page",
        flush=True,
    )
    return HTML(string=_set_pdf_page_height(doc, total_height_in)).write_pdf()


def _find_pdf_browser() -> str:
    candidates = [
        os.getenv("PDF_BROWSER_PATH", "").strip(),
        shutil.which("google-chrome"),
        shutil.which("google-chrome-stable"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        shutil.which("microsoft-edge"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
        "/snap/bin/chromium",
        "/usr/lib/chromium/chromium",
        "/usr/lib/chromium-browser/chromium-browser",
        "/opt/google/chrome/chrome",
        "/usr/bin/microsoft-edge",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return path
    raise RuntimeError("未找到可用于生成 PDF 的 Chrome/Edge/Chromium，请安装浏览器或设置 PDF_BROWSER_PATH")


def _free_local_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return int(s.getsockname()[1])


def _wait_for_cdp_target(port: int, timeout_seconds: float = 10.0) -> str:
    deadline = time.time() + timeout_seconds
    last_err = ""
    while time.time() < deadline:
        try:
            with urlopen(f"http://127.0.0.1:{port}/json", timeout=0.5) as resp:
                targets = json.loads(resp.read().decode("utf-8"))
            for target in targets:
                if target.get("type") == "page" and target.get("webSocketDebuggerUrl"):
                    return target["webSocketDebuggerUrl"]
        except Exception as exc:
            last_err = str(exc)
        time.sleep(0.1)
    raise RuntimeError(f"等待浏览器调试端口超时：{last_err}")


def _html_to_pdf_with_browser(doc: str) -> bytes:
    browser = _find_pdf_browser()
    last_err: Exception | None = None
    for headless_arg in ("--headless=new", "--headless"):
        try:
            return _html_to_pdf_with_browser_cmd(doc, browser, headless_arg)
        except Exception as exc:
            last_err = exc
            print(f"[pdf] browser render failed with {headless_arg}: {exc}", flush=True)
    raise RuntimeError(f"浏览器 PDF 生成失败：{last_err}")


def _html_to_pdf_with_browser_cmd(doc: str, browser: str, headless_arg: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="survey_pdf_") as tmp:
        tmp_path = Path(tmp)
        html_path = tmp_path / "report.html"
        profile_path = tmp_path / "profile"
        html_path.write_text(doc, encoding="utf-8")

        port = _free_local_port()
        cmd = [
            browser,
            headless_arg,
            f"--remote-debugging-port={port}",
            f"--user-data-dir={profile_path}",
            "--disable-gpu",
            "--no-sandbox",
            "--disable-dev-shm-usage",
            "--hide-scrollbars",
            "--window-size=960,800",
            html_path.as_uri(),
        ]
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        try:
            ws_url = _wait_for_cdp_target(port)
            from websockets.sync.client import connect

            with connect(ws_url, open_timeout=10, max_size=None) as ws:
                msg_id = 0

                def cdp(method: str, params: dict | None = None) -> dict:
                    nonlocal msg_id
                    msg_id += 1
                    ws.send(json.dumps({"id": msg_id, "method": method, "params": params or {}}))
                    while True:
                        data = json.loads(ws.recv())
                        if data.get("id") != msg_id:
                            continue
                        if "error" in data:
                            raise RuntimeError(f"{method} failed: {data['error']}")
                        return data.get("result", {})

                cdp("Page.enable")
                cdp("Runtime.evaluate", {
                    "expression": (
                        "new Promise(resolve => {"
                        " if (document.readyState === 'complete') resolve(true);"
                        " else window.addEventListener('load', () => resolve(true), {once:true});"
                        "})"
                    ),
                    "awaitPromise": True,
                    "returnByValue": True,
                })
                height_result = cdp("Runtime.evaluate", {
                    "expression": (
                        "Math.ceil(Math.max("
                        "document.body.scrollHeight,"
                        "document.documentElement.scrollHeight,"
                        "document.body.offsetHeight,"
                        "document.documentElement.offsetHeight"
                        "))"
                    ),
                    "returnByValue": True,
                })
                height_px = int(height_result.get("result", {}).get("value") or 1200)
                width_in = 10.0
                height_in = max(6.0, min((height_px + 24) / 96, 500.0))
                pdf_params = {
                    "printBackground": True,
                    "paperWidth": width_in,
                    "paperHeight": height_in,
                    "marginTop": 0,
                    "marginBottom": 0,
                    "marginLeft": 0,
                    "marginRight": 0,
                    "preferCSSPageSize": False,
                    "scale": 1,
                    "generateDocumentOutline": True,
                }
                try:
                    pdf_result = cdp("Page.printToPDF", pdf_params)
                except RuntimeError as exc:
                    if "generateDocumentOutline" not in str(exc):
                        raise
                    pdf_params.pop("generateDocumentOutline", None)
                    pdf_result = cdp("Page.printToPDF", pdf_params)
                return base64.b64decode(pdf_result["data"])
        finally:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()


def html_to_pdf_bytes(doc: str) -> bytes:
    renderer = os.getenv("PDF_RENDERER", "").strip().lower()
    if renderer == "browser":
        try:
            return _html_to_pdf_with_browser(doc)
        except Exception as exc:
            print(f"[pdf] requested browser renderer failed, falling back to WeasyPrint: {exc}", flush=True)

    try:
        return _html_to_pdf_with_weasyprint(doc)
    except Exception as exc:
        print(f"[pdf] WeasyPrint renderer failed, retrying browser renderer: {exc}", flush=True)

    try:
        return _html_to_pdf_with_browser(doc)
    except Exception as exc:
        print(f"[pdf] browser renderer failed after WeasyPrint fallback: {exc}", flush=True)
        raise


PDF_PAGE_RULE = "@page { size: 10in 14in; margin: 0; }"


PDF_CSS = PDF_PAGE_RULE + """
* { box-sizing: border-box; }
html {
  margin: 0;
  background: #fff;
}
body {
  font-family: "Noto Sans CJK SC", "Noto Sans CJK", "WenQuanYi Micro Hei", "Microsoft YaHei", "PingFang SC", "Source Han Sans SC", Arial, sans-serif;
  font-size: 13px;
  line-height: 1.75;
  color: #222;
  background: #fff;
  width: auto;
  max-width: 900px;
  margin: 0 auto;
  padding: 36px 44px 48px;
}
h1 {
  font-size: 22px;
  font-weight: 700;
  color: #1a1a1a;
  border-bottom: 2px solid #7c3aed;
  padding-bottom: 8px;
  margin: 0 0 18px;
}
h2 {
  font-size: 17px;
  font-weight: 600;
  color: #2d2d2d;
  margin: 24px 0 10px;
  border-left: 4px solid #7c3aed;
  padding-left: 10px;
  page-break-after: avoid;
}
h3 { font-size: 14.5px; color: #444; margin: 18px 0 8px; page-break-after: avoid; }
h4 { font-size: 13.5px; color: #555; margin: 14px 0 6px; page-break-after: avoid; }
p { margin: 0 0 9px; }
ul, ol { margin: 6px 0 12px 22px; padding: 0; }
li { margin-bottom: 4px; }
blockquote {
  border-left: 3px solid #999;
  padding: 6px 14px;
  margin: 10px 0;
  color: #666;
  font-style: italic;
  background: #f9f9f9;
}
table {
  border-collapse: collapse;
  width: 100%;
  max-width: 100%;
  margin: 14px 0;
  font-size: 11.5px;
  line-height: 1.55;
  page-break-inside: avoid;
  table-layout: auto;
}
thead { display: table-header-group; }
tr { page-break-inside: avoid; }
th {
  background: #f0ebff;
  color: #3d1d8a;
  font-weight: 600;
  padding: 7px 10px;
  border: 1px solid #d6c8ff;
  text-align: left;
  white-space: normal;
  overflow-wrap: anywhere;
  word-break: normal;
}
th:first-child, td:first-child { min-width: 72px; }
td {
  padding: 6px 10px;
  border: 1px solid #e0e0e0;
  vertical-align: top;
  white-space: normal;
  word-break: normal;
  overflow-wrap: anywhere;
}
tr:nth-child(even) td { background: #fafafa; }
img { max-width: 100%; height: auto; }
code { background: #f3f0ff; color: #5b21b6; padding: 1px 5px; border-radius: 3px; font-size: 12px; }
pre { background: #f5f5f5; padding: 12px; border-radius: 4px; overflow-wrap: break-word; white-space: pre-wrap; }
code, pre { font-family: "Noto Sans Mono CJK SC", "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Microsoft YaHei", monospace; }
strong { color: #111; font-weight: 600; }
hr { border: none; border-top: 1px solid #e0e0e0; margin: 20px 0; }
.core-highlight-box {
  background: rgba(139, 92, 246, 0.07);
  border: 1.5px solid rgba(139, 92, 246, 0.22);
  border-radius: 10px;
  padding: 14px 18px 10px;
  margin: 10px 0 16px;
  page-break-inside: avoid;
}
.core-highlight-box h2,
.core-highlight-box h3,
.core-highlight-box h4 { margin-top: 0; }
.core-highlight-box h3 {
  margin: 14px 0 8px;
  padding: 5px 9px;
  border-radius: 7px;
  background: rgba(139, 92, 246, 0.10);
  color: #5b21b6;
  font-size: 14px;
}
.core-highlight-box li { margin-bottom: 8px; }
"""


def _set_pdf_page_height(doc: str, height_in: float) -> str:
    height_in = max(6.0, min(float(height_in), 500.0))
    return doc.replace(PDF_PAGE_RULE, f"@page {{ size: 10in {height_in:.2f}in; margin: 0; }}", 1)


def _wrap_report_highlights(html_text: str) -> str:
    html_text = re.sub(r"<!--CORE_START-->\s*", '<div class="core-highlight-box">', html_text)
    html_text = re.sub(r"\s*<!--CORE_END-->", "</div>", html_text)

    summary_title = r"(?:本章总结|本节总结|章节总结|本部分总结)"
    heading_pat = re.compile(
        rf"(<h[34][^>]*>\s*{summary_title}\s*[:：]?\s*</h[34]>)(.*?)(?=<h[1-6][^>]*>|$)",
        re.S,
    )
    html_text = heading_pat.sub(r'<div class="core-highlight-box">\1\2</div>', html_text)

    paragraph_pat = re.compile(
        rf"(<p>\s*{summary_title}\s*[:：].*?</p>(?:(?!<h[1-6][^>]*>|<table|<pre|<div).)*?)"
        rf"(?=<h[1-6][^>]*>|$)",
        re.S,
    )
    return paragraph_pat.sub(r'<div class="core-highlight-box">\1</div>', html_text)
